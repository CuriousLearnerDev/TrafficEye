"""
模块功能: GUI
作者: W啥都学
创建日期: 2025-02-25
修改时间：2025-07-04
"""

__author__ = "W啥都学"

import csv
import json
import asyncio
import os
import queue
import secrets
import socket
import subprocess
import sys
import threading
import time
from pathlib import Path
from urllib.parse import unquote, quote

# Agent 内核（vendor/agent_core）
_VENDOR_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "vendor")
if _VENDOR_DIR not in sys.path:
    sys.path.insert(0, _VENDOR_DIR)

import pdfkit
from PyQt6 import QtCore
from PyQt6 import uic

import ai_analysis_core
import ai_config
from traffic_agent_runner import AgentWorker, SECURITY_REVIEW_GOAL
from traffic_agent_tools import SessionData, collect_risks
import requests
import yaml
import module
from module import CountryFlagDelegate
from jinja2 import Environment, FileSystemLoader
from examine import parse_http_request, classify_input, SecurityScanner

import replay_request
import re
import session_utils
from urllib.parse import urlparse
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer, QSize, QMarginsF, QRegularExpression, QDateTime, \
    QAbstractItemModel, QModelIndex, QObject, QSortFilterProxyModel, QUrl
from PyQt6.QtCore import Qt as QtCoreQt
from PyQt6.QtWidgets import (QApplication, QMainWindow, QPushButton, QMessageBox, QTextEdit,
                             QLineEdit, QCheckBox, QFileDialog, QVBoxLayout, QHBoxLayout,
                             QGridLayout, QGroupBox, QWidget, QTabWidget, QProgressBar,
                             QLabel, QRadioButton, QSplitter, QFrame, QTableWidget,
                             QTableWidgetItem, QListWidget, QListWidgetItem, QToolBar,
                             QStatusBar, QToolButton, QMenu, QSizePolicy, QFormLayout, QProgressDialog, QComboBox,
                             QInputDialog, QHeaderView, QAbstractItemView, QGraphicsDropShadowEffect, QScrollArea,
                             QToolTip, QTreeWidget, QTreeWidgetItem, QStyledItemDelegate, QTreeView,
                             QStyleOptionViewItem, QDialog, QDoubleSpinBox, QSpinBox, QPlainTextEdit)
from PyQt6.QtGui import (QIcon, QPixmap, QTextCursor, QTextCharFormat, QColor,
                         QPainter, QFont, QAction, QPalette, QTextDocument, QPageLayout, QPageSize, QSyntaxHighlighter,
                         QBrush, QCursor, QStandardItem, QStandardItemModel, QIntValidator, QMovie, QDesktopServices)
from PyQt6.QtCharts import QChartView, QChart, QPieSeries, QBarSeries, QBarSet, QBarCategoryAxis, QValueAxis, \
    QLineSeries, QDateTimeAxis, QCategoryAxis

import core_processing
import multiprocessing
import output_filtering
from collections import defaultdict
import datetime
import webbrowser
from log_parsing import log_identification
from log_parsing import table_import
from log_parsing import python_log_fallback
from binary_extraction import load_signatures, extract_file

import websockets

version = "0.11.1.0"

last_updated = "2026-08-02"

issuance = "内部版"

# Go 后台服务固定端口（与 Go -port 保持一致）
GO_SERVER_PORT = 10632
GO_BASE_URL = f"http://127.0.0.1:{GO_SERVER_PORT}"

# 设置代理
proxies = None


def _is_native_go_binary(path: str) -> bool:
    """按文件头判断是否为本机可执行格式，避免 Linux 误选 .exe。"""
    try:
        with open(path, "rb") as f:
            magic = f.read(4)
    except OSError:
        return False
    if sys.platform.startswith("win"):
        return magic[:2] == b"MZ"
    # ELF: 0x7F 'E' 'L' 'F'
    return magic == b"\x7fELF"


def find_go_executable():
    """查找 Go 分析器：冻结包只认 exe 旁 lib/（外部调用，不进 PyInstaller 包内）。"""
    roots = []
    if getattr(sys, "frozen", False):
        # 故意不搜 _MEIPASS：log_identification 应旁挂在 exe 同级 lib/
        exe_dir = os.path.dirname(os.path.abspath(sys.executable))
        roots.append(exe_dir)
    try:
        roots.append(os.path.dirname(os.path.abspath(__file__)))
    except Exception:
        pass
    roots.append(os.getcwd())
    seen = set()
    uniq_roots = []
    for r in roots:
        r = os.path.abspath(r)
        if r not in seen:
            seen.add(r)
            uniq_roots.append(r)

    # Windows 优先 .exe；Linux/macOS 优先无后缀，避免误选对方平台二进制
    if sys.platform.startswith("win"):
        candidates = ("log_identification.exe", "log_identification")
    else:
        candidates = ("log_identification", "log_identification.exe")

    rel_dirs = (
        "lib",
        os.path.join("go_source_code", "log_parsing"),
        "go_source_code",
        "GO流量分析源码",
        os.path.join("GO流量分析源码", "log_parsing"),
    )
    names = [os.path.join(d, name) for d in rel_dirs for name in candidates]

    for root in uniq_roots:
        for rel in names:
            path = os.path.join(root, rel)
            if os.path.isfile(path) and _is_native_go_binary(path):
                return path
    return None


def get_output_dir() -> str:
    """分析结果目录：安装包写到 exe 旁的 output/，开发态写到当前目录 output/。"""
    env = os.environ.get("TRAFFICEYE_OUTPUT_DIR")
    if env:
        path = os.path.abspath(env)
        os.makedirs(path, exist_ok=True)
        return path
    if getattr(sys, "frozen", False):
        path = os.path.join(os.path.dirname(os.path.abspath(sys.executable)), "output")
    else:
        path = os.path.join(os.getcwd(), "output")
    os.makedirs(path, exist_ok=True)
    return path


def _pid_listening_on_port(port: int):
    """Windows: 查占用端口的 PID；其它平台返回 None。"""
    if sys.platform != "win32":
        return None
    try:
        out = subprocess.check_output(
            ["netstat", "-ano"],
            creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0,
            text=True,
            errors="ignore",
        )
    except Exception:
        return None
    needle = f":{port} "
    for line in out.splitlines():
        if "LISTENING" not in line.upper():
            continue
        if needle not in line and f":{port}\t" not in line:
            # 兼容 "0.0.0.0:10632" / "[::]:10632"
            if f":{port}" not in line:
                continue
        parts = line.split()
        if not parts:
            continue
        try:
            return int(parts[-1])
        except ValueError:
            continue
    return None


def _kill_pid(pid: int) -> bool:
    if not pid or pid <= 0:
        return False
    try:
        if sys.platform == "win32":
            subprocess.run(
                ["taskkill", "/PID", str(pid), "/F", "/T"],
                capture_output=True,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0,
            )
        else:
            os.kill(pid, 9)
        return True
    except Exception:
        return False


def wait_go_ready(base_url, timeout=15.0):
    """轮询 /status 直到 Go 服务就绪"""
    deadline = time.time() + timeout
    last_err = None
    while time.time() < deadline:
        try:
            resp = requests.get(f"{base_url}/status", timeout=1)
            if resp.status_code == 200:
                return True
        except requests.exceptions.RequestException as e:
            last_err = e
        time.sleep(0.3)
    raise TimeoutError(f"等待 Go 服务就绪超时 ({base_url}): {last_err}")


def prefer_go_log_analyzer() -> bool:
    """
    默认优先 Go。
    - TRAFFICEYE_USE_GO_LOG=0/false 强制 Python
    - TRAFFICEYE_USE_GO_LOG=1/true 强制尝试 Go
    - 未设置：有 exe 就用 Go
    """
    flag = os.environ.get("TRAFFICEYE_USE_GO_LOG", "").strip().lower()
    if flag in ("0", "false", "no", "python"):
        return False
    if flag in ("1", "true", "yes", "go"):
        return True
    return find_go_executable() is not None


def INITIALIZE_DEFAULTDICT():
    return defaultdict(lambda: {
        'count': 0,
        'danger_count': 0,
        'source_ips': defaultdict(lambda: {
            'count': 0,
            'methods': defaultdict(int),
            'status_codes': defaultdict(int),
            'UA': defaultdict(int),
            'sizes': defaultdict(int),
            'frontend': defaultdict(int),
            'request_time': defaultdict(int),
            'backend': defaultdict(int),
            'danger': list()
        }), })


def extract_url_stats_data(url_stats):
    """
    统一取出 URI 统计字典。
    - Go/表格结果: {"data": {uri: stats, ...}, "_global_stats": ...}
    - 远程实时解析: 扁平 {uri: stats, ...}（没有 data 层）
    """
    if not isinstance(url_stats, dict):
        return {}
    skip = {"_global_stats", "filter_applied", "_import_meta"}
    nested = url_stats.get("data") if "data" in url_stats else None
    flat = {
        k: v for k, v in url_stats.items()
        if k not in skip and k != "data" and isinstance(v, dict) and ("source_ips" in v or "count" in v)
    }
    if flat:
        return flat
    if isinstance(nested, dict):
        return {
            k: v for k, v in nested.items()
            if k not in skip and isinstance(v, dict)
        }
    return {}


def compute_global_stats(url_stats) -> dict:
    """从 URI 统计计算仪表盘用的全局数字（远程扁平结构也能用）。"""
    data = extract_url_stats_data(url_stats)
    if data:
        total_ips = set()
        total_status = set()
        request_total = 0
        danger_total = 0
        for _uri, stats in data.items():
            if not isinstance(stats, dict):
                continue
            request_total += int(stats.get("count") or 0)
            danger_total += int(stats.get("danger_count") or 0)
            for ip, ip_stats in (stats.get("source_ips") or {}).items():
                if not isinstance(ip_stats, dict):
                    continue
                total_ips.add(ip)
                for sc in (ip_stats.get("status_codes") or {}):
                    total_status.add(str(sc))
                danger_total += len(ip_stats.get("danger") or [])
        return {
            "request_total": request_total,
            "danger_total": danger_total,
            "total_unique_ips": len(total_ips),
            "total_unique_uris": len(data),
            "total_unique_status_code": len(total_status),
        }
    if isinstance(url_stats, dict):
        existing = url_stats.get("_global_stats")
        if isinstance(existing, dict):
            return {
                "request_total": int(existing.get("request_total") or 0),
                "danger_total": int(existing.get("danger_total") or 0),
                "total_unique_ips": int(existing.get("total_unique_ips") or 0),
                "total_unique_uris": int(existing.get("total_unique_uris") or 0),
                "total_unique_status_code": int(existing.get("total_unique_status_code") or 0),
            }
    return {
        "request_total": 0,
        "danger_total": 0,
        "total_unique_ips": 0,
        "total_unique_uris": 0,
        "total_unique_status_code": 0,
    }


FULL_DATA = []


class RemoteLogReceiver(QThread):
    """模式 A：本机主动连接服务器 Agent（ws://host:port）。"""
    log_received = pyqtSignal(str)
    log_batch = pyqtSignal(list)
    log_output = pyqtSignal(str)
    connection_status = pyqtSignal(bool)

    def __init__(self, server_url, token: str = "", line_sink=None, opts_line: str = "OPTS mode=live"):
        super().__init__()
        self.server_url = server_url
        self.token = token or ""
        self.running = True
        # 直接写入解析线程队列，不经主线程
        self.line_sink = line_sink
        self.opts_line = opts_line or "OPTS mode=live"

    def _build_url(self) -> str:
        url = self.server_url
        if self.token and "token=" not in url:
            sep = "&" if "?" in url else "?"
            url = f"{url}{sep}token={quote(self.token, safe='')}"
        return url

    def _push(self, batch):
        sink = self.line_sink
        if sink is None or not batch:
            return
        try:
            sink(batch)
        except Exception:
            pass

    async def listen_logs(self):
        url = self._build_url()
        while self.running:
            try:
                async with websockets.connect(url, max_size=2 ** 20) as ws:
                    if self.token:
                        await ws.send(f"AUTH {self.token}")
                    # 告知 Agent：默认只要实时；可选历史
                    opts = getattr(self, "opts_line", None) or "OPTS mode=live"
                    try:
                        await ws.send(opts)
                    except Exception:
                        pass
                    self.connection_status.emit(True)
                    self.log_output.emit(f"✅ 已连接到 WebSocket: {self.server_url}")
                    self.log_output.emit(f"⏳ {opts} ；后台解析中…")

                    n = 0
                    last_ui = 0.0
                    last_preview = ""
                    batch = []
                    last_flush = time.monotonic()
                    while self.running:
                        try:
                            msg = await ws.recv()
                            if isinstance(msg, bytes):
                                msg = msg.decode("utf-8", errors="ignore")
                            n += 1
                            batch.append(msg)
                            last_preview = (msg or "").strip()
                            now = time.monotonic()
                            if len(batch) >= 120 or (now - last_flush) >= 0.35:
                                self._push(batch)
                                batch = []
                                last_flush = now
                            # 监控窗口极低频提示
                            if n == 1 or (now - last_ui) >= 2.0:
                                last_ui = now
                                preview = last_preview[:120] + ("…" if len(last_preview) > 120 else "")
                                self.log_output.emit(f"[收包 {n}] {preview}")
                        except websockets.ConnectionClosed:
                            if batch:
                                self._push(batch)
                            self.log_output.emit("⚠️ WebSocket连接关闭，尝试重连...")
                            self.connection_status.emit(False)
                            await asyncio.sleep(2)
                            break
            except Exception as e:
                self.log_output.emit(f"❌ WebSocket连接失败: {e} 3秒后尝试")
                self.connection_status.emit(False)
                await asyncio.sleep(3)

    def run(self):
        asyncio.run(self.listen_logs())

    def stop(self):
        self.running = False
        self.connection_status.emit(False)


class RemoteLogListener(QThread):
    """模式 B：本机监听，等待服务器 Agent 主动连入并推送日志。"""
    log_received = pyqtSignal(str)
    log_batch = pyqtSignal(list)
    log_output = pyqtSignal(str)
    connection_status = pyqtSignal(bool)

    def __init__(self, host: str = "0.0.0.0", port: int = 8765, token: str = "", line_sink=None):
        super().__init__()
        self.host = host or "0.0.0.0"
        self.port = int(port)
        self.token = token or ""
        self.running = True
        self._server = None
        self._loop = None
        self.line_sink = line_sink

    async def _authenticate(self, ws) -> bool:
        if not self.token:
            return True
        path = ""
        req = getattr(ws, "request", None)
        if req is not None:
            path = getattr(req, "path", "") or ""
        elif hasattr(ws, "path"):
            path = ws.path or ""
        from urllib.parse import parse_qs, urlparse as _urlparse
        qs = parse_qs(_urlparse(path).query)
        q_token = (qs.get("token") or qs.get("auth") or [None])[0]
        if q_token and q_token == self.token:
            return True
        try:
            msg = await asyncio.wait_for(ws.recv(), timeout=8.0)
        except Exception:
            return False
        if isinstance(msg, bytes):
            msg = msg.decode("utf-8", errors="ignore")
        msg = (msg or "").strip()
        if msg.startswith("AUTH "):
            return msg[5:].strip() == self.token
        return msg == self.token

    async def _handle(self, ws):
        peer = getattr(ws, "remote_address", None)
        self.log_output.emit(f"📡 Agent 连入: {peer}")
        if not await self._authenticate(ws):
            self.log_output.emit(f"❌ Agent 鉴权失败: {peer}")
            try:
                await ws.close(code=4001, reason="auth failed")
            except Exception:
                pass
            return

        self.connection_status.emit(True)
        self.log_output.emit(f"✅ Agent 已认证，开始接收日志")
        n = 0
        last_ui = 0.0
        batch = []
        last_flush = time.monotonic()
        try:
            async for msg in ws:
                if not self.running:
                    break
                if isinstance(msg, bytes):
                    msg = msg.decode("utf-8", errors="ignore")
                n += 1
                batch.append(msg)
                now = time.monotonic()
                if len(batch) >= 120 or (now - last_flush) >= 0.35:
                    sink = self.line_sink
                    if sink:
                        try:
                            sink(batch)
                        except Exception:
                            pass
                    else:
                        self.log_batch.emit(batch)
                    batch = []
                    last_flush = now
                if n == 1 or (now - last_ui) >= 2.0:
                    last_ui = now
                    preview = (msg or "").strip()[:120]
                    self.log_output.emit(f"[收包 {n}] {preview}")
        except websockets.ConnectionClosed:
            self.log_output.emit("⚠️ Agent 连接已关闭")
        finally:
            if batch:
                sink = self.line_sink
                if sink:
                    try:
                        sink(batch)
                    except Exception:
                        pass
            self.connection_status.emit(False)

    async def _serve(self):
        self.log_output.emit(f"🎧 本机监听 ws://{self.host}:{self.port} （等待服务器 Agent dial）")
        async with websockets.serve(self._handle, self.host, self.port, max_size=2 ** 20) as server:
            self._server = server
            while self.running:
                await asyncio.sleep(0.3)

    def run(self):
        self._loop = asyncio.new_event_loop()
        asyncio.set_event_loop(self._loop)
        try:
            self._loop.run_until_complete(self._serve())
        except Exception as e:
            self.log_output.emit(f"❌ 本地监听失败: {e}")
            self.connection_status.emit(False)
        finally:
            try:
                self._loop.close()
            except Exception:
                pass

    def stop(self):
        self.running = False
        self.connection_status.emit(False)
        if self._loop and self._loop.is_running():
            self._loop.call_soon_threadsafe(lambda: None)


class TreeItem:
    """树形数据结构中的单个项"""

    def __init__(self, data, parent=None):
        self.parent_item = parent
        self.item_data = data  # 这一行的数据
        self.child_items = []  # 子项列表

    def append_child(self, child):
        """添加子项"""
        self.child_items.append(child)

    def child(self, row):
        """获取指定行的子项"""
        if row < 0 or row >= len(self.child_items):
            return None
        return self.child_items[row]

    def child_count(self):
        """子项数量"""
        return len(self.child_items)

    def column_count(self):
        """列数"""
        return len(self.item_data)

    def data(self, column):
        """获取指定列的数据"""
        if column < 0 or column >= len(self.item_data):
            return None
        return self.item_data[column]

    def set_data(self, column, value):
        """设置指定列的数据"""
        if column < 0 or column >= len(self.item_data):
            return False
        self.item_data[column] = value
        return True

    def parent(self):
        """获取父项"""
        return self.parent_item

    def row(self):
        """获取自己在父项中的行号"""
        if self.parent_item:
            return self.parent_item.child_items.index(self)
        return 0


class GoPcapProcessingThread(QThread):
    """
    使用Go后台服务进行全流量(PCAP)分析的QThread。
    接口: POST /analyze_pcap -> GET /pcap_status
    """
    finished = pyqtSignal(object, str, str, str)  # (ai, result_json, full_traffic, dns_json)
    progress_updated = pyqtSignal(int)  # 进度百分比
    error = pyqtSignal(str)  # 错误信息
    status_msg = pyqtSignal(str)  # 状态文字更新

    def __init__(self, file_path, ai_analysis_starts, optional_parameters, server_url=None):
        super().__init__()
        self.file_path = file_path
        self.ai_analysis_starts = ai_analysis_starts
        self.optional_parameters = optional_parameters or {}
        self.server_url = server_url or GO_BASE_URL
        self._stop = False

    def stop(self):
        """请求结束分析（轮询循环会退出）。"""
        self._stop = True
        self.requestInterruption()

    def run(self):
        try:
            self._stop = False
            # --- 1. 构造请求 Payload ---
            data_section = self.optional_parameters.get("Data_section_detection", {}) or {}

            enable_scan = (
                    self.optional_parameters.get("URL_Security_Check", False) or
                    self.optional_parameters.get("Request_Head_Security_Check", False) or
                    data_section.get("enabled", False)
            )

            payload = {
                "use_ek_format": True,
                "file_path": self.file_path,
                "enable_scan": enable_scan,
                "url_security_check": self.optional_parameters.get("URL_Security_Check", False),
                "request_head_security_check": self.optional_parameters.get("Request_Head_Security_Check", False),
                "data_section_detection": {
                    "enabled": data_section.get("enabled", False),
                    "binary": data_section.get("binary", False),
                    "forms": data_section.get("forms", False),
                    "json": data_section.get("json", False),
                    "xml": data_section.get("xml", False),
                    "multipart": data_section.get("multipart", False)
                },
                # 触发 Go 侧写出 HTTP 明文；data_flow=false 用可读正文
                "filter": {
                    "data_flow": bool(self.optional_parameters.get("Show_Body_Hex", False)),
                },
            }

            # --- 2. 发送开始分析请求 ---
            try:
                response = requests.post(
                    f"{self.server_url}/analyze_pcap",
                    json=payload,
                    timeout=5
                )
                if response.status_code == 409:
                    self.error.emit(response.json().get("error", "已有流量分析任务正在进行中"))
                    return
                response.raise_for_status()
                if response.json().get("status") != "started":
                    self.error.emit(f"Go服务未能启动流量分析任务: {response.text}")
                    return
            except requests.exceptions.RequestException as e:
                self.error.emit(f"无法连接到Go流量分析服务: {e}")
                return

            # --- 3. 轮询状态 ---
            consecutive_failures = 0
            max_failures = 10
            while not self._stop and not self.isInterruptionRequested():
                try:
                    status_response = requests.get(f"{self.server_url}/pcap_status", timeout=2)
                    status_response.raise_for_status()
                    status_data = status_response.json()
                    consecutive_failures = 0

                    status = status_data.get("status")
                    progress = status_data.get("progress", 0)
                    stage = status_data.get("stage") or ""

                    pct = int(progress * 100)
                    self.progress_updated.emit(pct)
                    if stage:
                        self.status_msg.emit(f"{stage}... {pct}%")
                    else:
                        self.status_msg.emit(f"正在分析... {pct}%")

                    if status == "done":
                        if self._stop or self.isInterruptionRequested():
                            return
                        result_path = status_data.get("result_path") or ""
                        full_traffic_path = status_data.get("full_traffic_path") or ""

                        if not result_path or not os.path.exists(result_path):
                            self.error.emit(f"Go服务报告完成，但结果文件不存在: {result_path}")
                            return

                        dns_result_path = status_data.get("dns_result_path") or ""
                        self.finished.emit(self.ai_analysis_starts, result_path, full_traffic_path, dns_result_path)
                        return

                    elif status == "error":
                        err = status_data.get("error_msg") or status_data.get("error") or "未知错误"
                        self.error.emit(f"Go服务分析出错: {err}")
                        return

                except requests.exceptions.RequestException as e:
                    consecutive_failures += 1
                    if consecutive_failures >= max_failures:
                        self.error.emit(f"轮询Go流量分析状态失败: {e}")
                        return

                time.sleep(0.5)

        except Exception as e:
            if self._stop or self.isInterruptionRequested():
                return
            self.error.emit(f"流量分析线程发生异常: {str(e)}")


class GoLogProcessingThread(QThread):
    """
    使用Go后台服务进行日志分析的QThread。
    它通过HTTP POST启动任务，然后轮询 /status 接口获取进度，
    最后在完成后读取结果JSON文件。
    """
    finished = pyqtSignal(object, str)  # (ai_analysis_starts, result_file_path)
    progress_updated = pyqtSignal(int)  # 进度百分比
    error = pyqtSignal(str)  # 错误信息

    def __init__(self, file_path, log_type, ai_analysis_starts, optional_parameters_log, server_url=None):
        super().__init__()
        self.file_path = file_path
        self.log_type = log_type
        self.ai_analysis_starts = ai_analysis_starts
        self.optional_parameters_log = optional_parameters_log or {}
        self.server_url = server_url or GO_BASE_URL
        self._stop = False

    def stop(self):
        """请求结束分析（轮询循环会退出）。"""
        self._stop = True
        self.requestInterruption()

    def run(self):
        try:
            self._stop = False
            enable_scan = bool(self.optional_parameters_log.get("URI_Security_Check", False))
            try:
                response = requests.post(
                    f"{self.server_url}/analyze",
                    json={"file_path": self.file_path, "enable_scan": enable_scan},
                    timeout=5
                )
                if response.status_code == 409:
                    self.error.emit(response.json().get("error", "已有日志分析任务正在进行中"))
                    return
                response.raise_for_status()
                if response.json().get("status") != "started":
                    self.error.emit(f"Go服务未能启动任务: {response.text}")
                    return
            except requests.exceptions.RequestException as e:
                self.error.emit(f"无法连接到Go分析器: {e}。请确保服务正在运行。")
                return

            consecutive_failures = 0
            max_failures = 10
            while not self._stop and not self.isInterruptionRequested():
                try:
                    status_response = requests.get(f"{self.server_url}/status", timeout=2)
                    status_response.raise_for_status()
                    status_data = status_response.json()
                    consecutive_failures = 0

                    status = status_data.get("status")
                    progress = status_data.get("progress", 0)
                    self.progress_updated.emit(int(progress * 100))

                    if status == "done":
                        if self._stop or self.isInterruptionRequested():
                            return
                        result_path = status_data.get("result_path") or ""
                        if not result_path:
                            self.error.emit("Go分析器完成了任务，但未返回结果路径。")
                            return

                        if not os.path.exists(result_path):
                            self.error.emit(f"Go分析器完成了任务，但结果文件 '{result_path}' 不存在。")
                            return

                        self.finished.emit(self.ai_analysis_starts, result_path)
                        return

                    elif status == "error":
                        error_msg = status_data.get("error_msg") or status_data.get("error") or "Go分析器报告了未知错误。"
                        self.error.emit(error_msg)
                        return

                except requests.exceptions.RequestException as e:
                    consecutive_failures += 1
                    if consecutive_failures >= max_failures:
                        self.error.emit(f"轮询Go分析状态失败: {e}")
                        return

                time.sleep(0.5)

        except Exception as e:
            if self._stop or self.isInterruptionRequested():
                return
            self.error.emit(f"Go日志处理线程发生意外错误: {e}")


# 在LanguageManager类中添加多语言支持
class LanguageManager:
    def __init__(self, default_lang="zh"):
        self.current_lang = default_lang
        self.translations = {}
        self.load_language(default_lang)

    def load_language(self, lang_code):
        """加载语言文件"""
        lang_file = f"lib/lang_{lang_code}.json"
        try:
            if os.path.exists(lang_file):
                with open(lang_file, "r", encoding="utf-8") as f:
                    self.translations = json.load(f)
                self.current_lang = lang_code
                return True
            else:
                print(f"Language file not found: {lang_file}")
                return False
        except Exception as e:
            print(f"Error loading language file: {e}")
            return False

    def tr(self, key, default=None):
        """获取翻译文本，支持默认值"""
        return self.translations.get(key, default or key)

    def get_current_language(self):
        """获取当前语言"""
        return self.current_lang


class FileWriter(QObject):
    finished = pyqtSignal()

    def __init__(self):
        super().__init__()
        self.queue = queue.Queue()
        self.running = True
        self.buffer = []
        self.buffer_size = 100  # 缓冲100行再写入
        self.thread = threading.Thread(target=self._write_loop, daemon=True)
        self.thread.start()

    def _write_loop(self):
        while self.running:
            try:
                # 从队列获取数据
                file_name, content = self.queue.get(timeout=0.1)

                # 添加到缓冲区
                self.buffer.append(content)

                # 缓冲区满或队列为空时写入文件
                if len(self.buffer) >= self.buffer_size or self.queue.empty():
                    self._flush_buffer(file_name)

            except queue.Empty:
                # 定期检查缓冲区
                if self.buffer:
                    self._flush_buffer(file_name)
                continue

        # 退出前确保所有数据写入
        if self.buffer:
            self._flush_buffer(file_name)
        self.finished.emit()

    def _flush_buffer(self, file_name):
        if not self.buffer:
            return

        try:
            with open(file_name, 'a', encoding='utf-8') as file:
                file.write('\n'.join(self.buffer) + '\n')
            self.buffer.clear()
        except Exception as e:
            print(f"Error writing to file: {e}")

    def add_content(self, file_name, content):
        self.queue.put((file_name, content))

    def stop(self):
        self.running = False
        self.thread.join()


class OverlayLoadingWidget(QWidget):
    """
    半透明全屏覆盖层 + 动画指示 + 提示文字
    作为父控件子窗口铺满并居中显示卡片。
    """

    def __init__(self, parent=None, lang_manager=None):
        super().__init__(parent)
        self._host = parent
        self.lang_manager = lang_manager
        self.setAttribute(Qt.WidgetAttribute.WA_StyledBackground, True)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents, False)
        self.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        # 不用布局居中：父控件未撑满时布局会把卡片挤到左上角
        self.card = QWidget(self)
        self.card.setFixedSize(320, 200)
        self.card.setStyleSheet("""
            background-color: rgba(60, 60, 60, 230);
            border-radius: 10px;
            border: 1px solid #444;
        """)

        card_layout = QVBoxLayout(self.card)
        card_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        card_layout.setContentsMargins(20, 20, 20, 20)
        card_layout.setSpacing(16)

        self.loading_movie = QMovie("ico/loading.gif")
        self.loading_label = QLabel()
        self.loading_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.loading_label.setMovie(self.loading_movie)
        card_layout.addWidget(self.loading_label)

        default_msg = "正在处理数据，请稍候..."
        if self.lang_manager:
            default_msg = self.lang_manager.tr("processing_data", default_msg)
        self.message_label = QLabel(default_msg)
        self.message_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.message_label.setWordWrap(True)
        self.message_label.setStyleSheet("""
            color: #DDD;
            font-size: 16px;
            font-weight: bold;
        """)
        card_layout.addWidget(self.message_label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(6)
        self.progress_bar.hide()

        self.hide()

    def _center_card(self):
        """把提示卡片放到覆盖层正中央"""
        if not hasattr(self, "card"):
            return
        x = max(0, (self.width() - self.card.width()) // 2)
        y = max(0, (self.height() - self.card.height()) // 2)
        self.card.move(x, y)
        self.card.raise_()

    def _cover_parent(self):
        host = self.parentWidget() or self._host
        if host is None:
            return
        # 铺满父控件客户区
        self.setGeometry(0, 0, max(1, host.width()), max(1, host.height()))
        self._center_card()
        self.raise_()

    def show(self, message=None):
        """显示加载覆盖层"""
        if message is None:
            message = "正在处理数据，请稍候..."
            if self.lang_manager:
                message = self.lang_manager.tr("processing_data", message)
        self.message_label.setText(message)
        self._cover_parent()
        if self.loading_movie.state() == QMovie.MovieState.NotRunning:
            self.loading_movie.start()
        super().show()
        self._cover_parent()
        self._center_card()
        self.raise_()
        self.repaint()
        QApplication.processEvents()
        # 再补一次：部分情况下首帧父控件尺寸尚未稳定
        QTimer.singleShot(0, self._cover_parent)

    def hide(self):
        """隐藏加载覆盖层"""
        if self.loading_movie.state() != QMovie.MovieState.NotRunning:
            self.loading_movie.stop()
        super().hide()
        QApplication.processEvents()

    def update_progress(self, value):
        QApplication.processEvents()

    def set_message(self, message):
        self.message_label.setText(message)
        self._center_card()
        QApplication.processEvents()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.fillRect(self.rect(), QColor(0, 0, 0, 150))

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._center_card()


class RiskTreeModel(QAbstractItemModel):
    """按严重级别分组的风险树，便于快速研判"""

    SEVERITY_ORDER = ("严重", "高危", "中危", "低危", "信息", "中", "未知")
    SEVERITY_RANK = {name: i for i, name in enumerate(SEVERITY_ORDER)}

    def __init__(self, url_stats, parent=None, lang_manager=None):
        super().__init__(parent)
        self.lang_manager = lang_manager
        self.stats = {
            "total": 0,
            "by_severity": defaultdict(int),
            "by_type": defaultdict(int),
            "unique_ips": set(),
            "unique_urls": set(),
        }
        self.root_item = TreeItem([
            self.lang_manager.tr("risk_group", "分组 / 来源"),
            self.lang_manager.tr("risk_url", "风险URL"),
            self.lang_manager.tr("risk_type", "风险类型"),
            self.lang_manager.tr("risk_level", "风险等级"),
            self.lang_manager.tr("matched_rule", "匹配规则"),
            self.lang_manager.tr("matched_content", "匹配摘要"),
        ])
        self.setup_model_data(url_stats)

    @staticmethod
    def _normalize_severity(level):
        text = str(level or "").strip()
        if not text:
            return "未知"
        lower = text.lower()
        mapping = [
            (("严重", "critical"), "严重"),
            (("高危", "high"), "高危"),
            (("中危", "medium", "中"), "中危"),
            (("低危", "low"), "低危"),
            (("信息", "info"), "信息"),
        ]
        for keys, norm in mapping:
            if any(k in lower or k in text for k in keys):
                return norm
        return text

    @staticmethod
    def _short_ip(ip_key):
        if not ip_key:
            return ""
        for sep in ("：", ":"):
            if sep in ip_key:
                return ip_key.split(sep, 1)[0].strip()
        return ip_key

    def collect_risks(self, url_stats):
        risks = []
        seen = set()
        for url, stats in url_stats.items():
            if url == "_global_stats" or "source_ips" not in stats:
                continue
            for src_ip, ip_stats in stats.get("source_ips", {}).items():
                for danger in ip_stats.get("danger") or []:
                    if isinstance(danger, dict):
                        level = self._normalize_severity(danger.get("severity", "中危"))
                        risk = {
                            "ip": src_ip,
                            "ip_short": self._short_ip(src_ip),
                            "url": url,
                            "type": danger.get("rule_type", self.lang_manager.tr("unknown_type", "未知")),
                            "level": level,
                            "rule": danger.get("rule_name", self.lang_manager.tr("unknown_rule", "未知规则")),
                            "content": danger.get("context", "") or danger.get("matched", "") or "",
                            "position": danger.get("position", (0, 0)),
                        }
                    else:
                        risk = {
                            "ip": src_ip,
                            "ip_short": self._short_ip(src_ip),
                            "url": url,
                            "type": self.lang_manager.tr("unknown_type", "未知"),
                            "level": "中危",
                            "rule": str(danger),
                            "content": "",
                            "position": (0, 0),
                        }

                    # 去重：同一 IP+URL+规则+等级 只保留一条，避免刷屏
                    key = (risk["ip_short"], risk["url"], risk["rule"], risk["level"])
                    if key in seen:
                        continue
                    seen.add(key)
                    risks.append(risk)

                    self.stats["total"] += 1
                    self.stats["by_severity"][risk["level"]] += 1
                    self.stats["by_type"][risk["type"]] += 1
                    self.stats["unique_ips"].add(risk["ip_short"])
                    self.stats["unique_urls"].add(risk["url"])
        return risks

    def setup_model_data(self, url_stats):
        risks = self.collect_risks(url_stats)

        # 按严重级别分组，便于别人快速抓住重点
        severity_groups = defaultdict(list)
        for risk in risks:
            severity_groups[risk["level"]].append(risk)

        ordered = sorted(
            severity_groups.items(),
            key=lambda kv: self.SEVERITY_RANK.get(kv[0], 99)
        )

        for level, risk_list in ordered:
            # 组内再按类型聚合
            type_groups = defaultdict(list)
            for risk in risk_list:
                type_groups[risk["type"]].append(risk)

            advice = {
                "严重": "建议立即处置",
                "高危": "建议优先排查",
                "中危": "建议安排核查",
                "低危": "可纳入监控",
                "信息": "仅供参考",
            }.get(level, "建议复核")

            level_item = TreeItem([
                f"{level}（{len(risk_list)}）",
                advice,
                f"{len(type_groups)} 类规则",
                level,
                "",
                "",
            ], self.root_item)
            level_item.is_group = True
            level_item.group_level = level
            self.root_item.append_child(level_item)

            for rtype, items in sorted(type_groups.items(), key=lambda x: (-len(x[1]), x[0])):
                type_item = TreeItem([
                    f"  {rtype}（{len(items)}）",
                    "",
                    rtype,
                    level,
                    "",
                    "",
                ], level_item)
                type_item.is_group = True
                type_item.group_level = level
                level_item.append_child(type_item)

                # 同类型内按 IP 排序，方便对照来源
                items.sort(key=lambda r: (r["ip_short"], r["url"]))
                for risk in items:
                    content = risk["content"]
                    if content and len(content) > 60:
                        content = content[:60] + "…"
                    pos = risk.get("position") or (0, 0)
                    if isinstance(pos, (list, tuple)) and len(pos) >= 2:
                        pos_hint = f" @{pos[0]}-{pos[1]}"
                    else:
                        pos_hint = ""
                    leaf = TreeItem([
                        risk["ip_short"],
                        unquote(risk["url"]),
                        risk["type"],
                        risk["level"],
                        risk["rule"],
                        f"{content}{pos_hint}" if content else pos_hint.strip(),
                    ], type_item)
                    leaf.risk_data = risk
                    type_item.append_child(leaf)

    def index(self, row, column, parent=QModelIndex()):
        if not self.hasIndex(row, column, parent):
            return QModelIndex()
        parent_item = self.root_item if not parent.isValid() else parent.internalPointer()
        child_item = parent_item.child(row)
        if child_item:
            return self.createIndex(row, column, child_item)
        return QModelIndex()

    def parent(self, index):
        if not index.isValid():
            return QModelIndex()
        child_item = index.internalPointer()
        parent_item = child_item.parent()
        if parent_item == self.root_item or not parent_item:
            return QModelIndex()
        return self.createIndex(parent_item.row(), 0, parent_item)

    def rowCount(self, parent=QModelIndex()):
        parent_item = self.root_item if not parent.isValid() else parent.internalPointer()
        return parent_item.child_count()

    def columnCount(self, parent=QModelIndex()):
        return self.root_item.column_count()

    def data(self, index, role=QtCore.Qt.ItemDataRole.DisplayRole):
        if not index.isValid():
            return None
        item = index.internalPointer()

        if role == QtCore.Qt.ItemDataRole.DisplayRole:
            return item.data(index.column())
        if role == QtCore.Qt.ItemDataRole.ToolTipRole:
            if hasattr(item, "risk_data"):
                r = item.risk_data
                return (
                    f"IP: {r.get('ip', '')}\n"
                    f"URL: {unquote(r.get('url', ''))}\n"
                    f"类型: {r.get('type', '')}\n"
                    f"等级: {r.get('level', '')}\n"
                    f"规则: {r.get('rule', '')}\n"
                    f"内容: {r.get('content', '')}"
                )
            return item.data(index.column())
        if role == QtCore.Qt.ItemDataRole.TextAlignmentRole:
            return int(QtCore.Qt.AlignmentFlag.AlignVCenter | QtCore.Qt.AlignmentFlag.AlignLeft)
        if role == QtCore.Qt.ItemDataRole.ForegroundRole:
            level = getattr(item, "group_level", None) or item.data(3) or ""
            colors = {
                "严重": QColor("#D32F2F"),
                "高危": QColor("#F57C00"),
                "中危": QColor("#F9A825"),
                "低危": QColor("#388E3C"),
                "信息": QColor("#757575"),
            }
            color = colors.get(self._normalize_severity(level))
            if color and (index.column() in (0, 3) or getattr(item, "is_group", False)):
                return color
        if role == QtCore.Qt.ItemDataRole.FontRole:
            if getattr(item, "is_group", False):
                font = QFont()
                font.setBold(True)
                return font
        return None

    def headerData(self, section, orientation, role=QtCore.Qt.ItemDataRole.DisplayRole):
        if orientation == QtCore.Qt.Orientation.Horizontal and role == QtCore.Qt.ItemDataRole.DisplayRole:
            return self.root_item.data(section)
        return None

    def flags(self, index):
        if not index.isValid():
            return QtCore.Qt.ItemFlag.NoItemFlags
        return super().flags(index) | QtCore.Qt.ItemFlag.ItemIsEnabled | QtCore.Qt.ItemFlag.ItemIsSelectable


class StatsTreeModel(QAbstractItemModel):
    # 树表只渲染 Top N，避免海量 URL 一次性塞进 UI 导致卡顿
    MAX_URLS = 500
    MAX_IPS_PER_URL = 50

    def __init__(self, url_stats, parent=None, lang_manager=None, max_urls=None):
        super().__init__(parent)
        self.lang_manager = lang_manager
        self.max_urls = max_urls if max_urls is not None else self.MAX_URLS
        self.total_urls = 0
        self.displayed_urls = 0
        self.root_item = TreeItem([
            self.lang_manager.tr("url", "URL"),
            self.lang_manager.tr("visit_count", "访问次数"),
            self.lang_manager.tr("source_ip", "来源IP"),
            self.lang_manager.tr("status_code", "状态码"),
            self.lang_manager.tr("method", "方法"),
            self.lang_manager.tr("user_agent", "UA"),
        ])

        self.setup_model_data(url_stats)

    @staticmethod
    def _parse_ip_key(ip_key):
        """解析 'IP：归属地' 或纯 IP；返回 (ip, location)"""
        if not ip_key:
            return "", "未知"
        for sep in ("：", ":"):
            if sep in ip_key:
                ip_addr, loc = ip_key.split(sep, 1)
                return ip_addr.strip(), (loc.strip() or "未知")
        return ip_key.strip(), "未知"

    @staticmethod
    def _compact_counts(data, limit=5):
        """紧凑展示: 200×10, 404×2"""
        if not data:
            return "-"
        items = sorted(data.items(), key=lambda x: x[1], reverse=True)[:limit]
        return ", ".join(f"{k}×{v}" for k, v in items)

    @staticmethod
    def _short_ua(ua):
        """压缩 UA 为可读短名"""
        if not ua or ua == "-":
            return "-"
        rules = [
            (r"Edg(?:e)?/([\d.]+)", "Edge/{0}"),
            (r"Chrome/([\d.]+)", "Chrome/{0}"),
            (r"Firefox/([\d.]+)", "Firefox/{0}"),
            (r"Version/([\d.]+).*Safari", "Safari/{0}"),
            (r"MSIE ([\d.]+)", "IE/{0}"),
            (r"Trident/.*rv:([\d.]+)", "IE/{0}"),
            (r"curl/([\d.]+)", "curl/{0}"),
            (r"python-requests/([\d.]+)", "requests/{0}"),
            (r"Go-http-client/([\d.]+)", "Go/{0}"),
            (r"okhttp/([\d.]+)", "okhttp/{0}"),
            (r"Java/([\d.]+)", "Java/{0}"),
        ]
        for pattern, fmt in rules:
            m = re.search(pattern, ua, re.I)
            if m:
                return fmt.format(m.group(1))
        return (ua[:36] + "…") if len(ua) > 36 else ua

    @staticmethod
    def _short_location(location, max_len=28):
        if not location or location == "未知":
            return "未知"
        # 去掉全 0 段，缩短显示
        parts = [p for p in location.split("-") if p and p != "0"]
        text = "-".join(parts) if parts else location
        return (text[:max_len] + "…") if len(text) > max_len else text

    @classmethod
    def _format_ua_counts(cls, ua_data, limit=3):
        if not ua_data:
            return "-"
        items = sorted(ua_data.items(), key=lambda x: x[1], reverse=True)[:limit]
        return ", ".join(f"{cls._short_ua(k)}×{v}" for k, v in items)

    @classmethod
    def _aggregate_from_ips(cls, source_ips):
        """父行汇总：IP 数、状态码、方法、UA"""
        status = defaultdict(int)
        methods = defaultdict(int)
        uas = defaultdict(int)
        for ip_data in source_ips.values():
            for k, v in ip_data.get("status_codes", {}).items():
                status[k] += int(v)
            for k, v in ip_data.get("methods", {}).items():
                methods[k] += int(v)
            for k, v in ip_data.get("UA", {}).items():
                uas[k] += int(v)
        return status, methods, uas

    def setup_model_data(self, url_stats):
        sorted_stats = sorted(
            ((u, s) for u, s in url_stats.items() if u != "_global_stats"),
            key=lambda item: item[1].get("count", 0),
            reverse=True
        )
        self.total_urls = len(sorted_stats)
        limited_stats = sorted_stats[:self.max_urls]
        self.displayed_urls = len(limited_stats)

        for url, stats in limited_stats:
            source_ips = stats.get("source_ips", {}) or {}
            status_all, methods_all, uas_all = self._aggregate_from_ips(source_ips)
            ip_count = len(source_ips)
            decoded_url = unquote(url)

            # 父行：URL + 汇总信息（单行紧凑，避免空列）
            url_item = TreeItem([
                decoded_url,
                f"{stats.get('count', 0):,}",
                f"{ip_count} IP" if ip_count else "-",
                self._compact_counts(status_all),
                self._compact_counts(methods_all),
                self._format_ua_counts(uas_all),
            ], self.root_item)
            # 悬浮看完整信息
            url_item.tooltip = [
                decoded_url,
                str(stats.get("count", 0)),
                f"来源 IP 共 {ip_count} 个",
                self._compact_counts(status_all, limit=20),
                self._compact_counts(methods_all, limit=20),
                self._format_ua_counts(uas_all, limit=10),
            ]
            url_item.sort_count = int(stats.get("count", 0) or 0)
            self.root_item.append_child(url_item)

            ip_items = sorted(
                source_ips.items(),
                key=lambda x: x[1].get("count", 0),
                reverse=True
            )[:self.MAX_IPS_PER_URL]

            for ip_key, ip_data in ip_items:
                ip_addr, location = self._parse_ip_key(ip_key)
                short_loc = self._short_location(location)
                # 国旗委托期望格式: "IP\n(归属地)"
                ip_display = f"{ip_addr}\n({short_loc})"
                status_text = self._compact_counts(ip_data.get("status_codes", {}))
                method_text = self._compact_counts(ip_data.get("methods", {}))
                ua_text = self._format_ua_counts(ip_data.get("UA", {}))

                ip_item = TreeItem([
                    "",
                    f"{ip_data.get('count', 0):,}",
                    ip_display,
                    status_text,
                    method_text,
                    ua_text,
                ], url_item)
                ip_item.tooltip = [
                    ip_addr,
                    str(ip_data.get("count", 0)),
                    f"{ip_addr}\n{location}",
                    self._compact_counts(ip_data.get("status_codes", {}), limit=20),
                    self._compact_counts(ip_data.get("methods", {}), limit=20),
                    "\n".join(
                        f"{k} × {v}" for k, v in sorted(
                            (ip_data.get("UA") or {}).items(),
                            key=lambda x: x[1], reverse=True
                        )[:8]
                    ) or "-",
                ]
                ip_item.sort_count = int(ip_data.get("count", 0) or 0)
                url_item.append_child(ip_item)

            # 还有更多 IP 时提示
            remain = ip_count - len(ip_items)
            if remain > 0:
                more_item = TreeItem([
                    "",
                    "",
                    f"… 另有 {remain} 个 IP（导出可看全部）",
                    "", "", "",
                ], url_item)
                more_item.tooltip = ["", "", f"该 URL 共 {ip_count} 个来源 IP", "", "", ""]
                more_item.sort_count = -1
                url_item.append_child(more_item)

    def data(self, index, role=QtCoreQt.ItemDataRole.DisplayRole):
        if not index.isValid():
            return None

        item = index.internalPointer()

        if role == QtCoreQt.ItemDataRole.DisplayRole:
            return item.data(index.column())
        if role == QtCoreQt.ItemDataRole.ToolTipRole:
            tips = getattr(item, "tooltip", None)
            if tips and index.column() < len(tips):
                tip = tips[index.column()]
                return tip if tip else None
            return item.data(index.column()) or None
        if role == QtCoreQt.ItemDataRole.UserRole:
            return getattr(item, "sort_count", 0)
        if role == QtCoreQt.ItemDataRole.TextAlignmentRole:
            if index.column() == 1:
                return int(QtCoreQt.AlignmentFlag.AlignRight | QtCoreQt.AlignmentFlag.AlignVCenter)
            if index.column() in (3, 4):
                return int(QtCoreQt.AlignmentFlag.AlignCenter | QtCoreQt.AlignmentFlag.AlignVCenter)
        if role == QtCoreQt.ItemDataRole.ForegroundRole and index.column() == 2:
            # 父行「N IP」用灰色区分
            text = item.data(2) or ""
            if text.endswith(" IP") or text.startswith("…"):
                return QBrush(QColor("#888888"))

        return None

    def headerData(self, section, orientation, role=QtCoreQt.ItemDataRole.DisplayRole):
        if orientation == QtCoreQt.Orientation.Horizontal and role == QtCoreQt.ItemDataRole.DisplayRole:
            return self.root_item.data(section)
        return None

    def flags(self, index):
        if not index.isValid():
            return QtCoreQt.ItemFlag.NoItemFlags
        return super().flags(index) | QtCoreQt.ItemFlag.ItemIsEnabled | QtCoreQt.ItemFlag.ItemIsSelectable

    def index(self, row, column, parent=QModelIndex()):
        if not self.hasIndex(row, column, parent):
            return QModelIndex()

        parent_item = self.root_item if not parent.isValid() else parent.internalPointer()
        child_item = parent_item.child(row)

        if child_item:
            return self.createIndex(row, column, child_item)
        return QModelIndex()

    def parent(self, index):
        if not index.isValid():
            return QModelIndex()

        child_item = index.internalPointer()
        parent_item = child_item.parent()

        if parent_item == self.root_item or not parent_item:
            return QModelIndex()

        return self.createIndex(parent_item.row(), 0, parent_item)

    def rowCount(self, parent=QModelIndex()):
        parent_item = self.root_item if not parent.isValid() else parent.internalPointer()
        return parent_item.child_count()

    def columnCount(self, parent=QModelIndex()):
        return self.root_item.column_count()



class remoteDynamicsLogProcessingThread(QThread):
    """后台解析远程日志：主线程只收轻量统计快照，避免卡顿。"""
    finished = pyqtSignal(object)
    dynamic_updates = pyqtSignal(object)  # 兼容旧连接；实际发轻量 dict
    progress_updated = pyqtSignal(int)
    error = pyqtSignal(str)
    stats_tick = pyqtSignal(dict)  # {global, lines, top_rows}

    def __init__(self, url_stats, optional_parameters_log):
        super().__init__()
        self.url_stats = url_stats
        self.optional_parameters_log = optional_parameters_log or {}
        self.optional_parameters_log.setdefault("Skip_IP_Geo", True)
        # 恶意分析可选：勾选后开启 URI 规则（会更耗 CPU）
        self.optional_parameters_log.setdefault("URI_Security_Check", False)
        self.lines_processed = 0
        self._q = queue.Queue(maxsize=80000)
        self._stop = threading.Event()
        self._lock = threading.Lock()
        self._fixed_type = "nginx_access"

    def stop(self):
        self._stop.set()

    def enqueue(self, line: str):
        try:
            self._q.put_nowait(line)
        except queue.Full:
            # 积压时丢弃，保 UI 流畅
            pass

    def enqueue_batch(self, lines):
        if not lines:
            return
        for line in lines:
            self.enqueue(line)

    # 兼容旧信号槽名
    def feed_line(self, line, notify=True):
        self.enqueue(line)

    def feed_lines(self, lines):
        self.enqueue_batch(lines)

    def _parse_one(self, line: str):
        line = (line or "").strip()
        if not line or line == "0" or line.startswith("AUTH ") or line.startswith("[agent]"):
            return
        if line.upper().startswith("OPTS"):
            return
        line = line.replace("＂", '"')
        # 固定走通用 access，避免每行 guess 全格式
        log_identification.parse_access_line(
            line, self.url_stats, self._fixed_type, self.optional_parameters_log
        )
        self.lines_processed += 1

    def _snapshot(self) -> dict:
        with self._lock:
            g = compute_global_stats(self.url_stats)
            data = extract_url_stats_data(self.url_stats)
            top = sorted(
                ((u, s) for u, s in data.items() if isinstance(s, dict)),
                key=lambda kv: int(kv[1].get("count") or 0),
                reverse=True,
            )[:80]
            # 只拷贝展示需要的浅字段，避免拖垮 GUI
            top_rows = []
            for u, s in top:
                ips = s.get("source_ips") or {}
                codes = {}
                methods = {}
                for ip_s in list(ips.values())[:20]:
                    if not isinstance(ip_s, dict):
                        continue
                    for sc, c in (ip_s.get("status_codes") or {}).items():
                        codes[sc] = codes.get(sc, 0) + c
                    for m, c in (ip_s.get("methods") or {}).items():
                        methods[m] = methods.get(m, 0) + c
                top_rows.append({
                    "url": u,
                    "count": int(s.get("count") or 0),
                    "ips": list(ips.keys())[:5],
                    "status": list(codes.keys())[:6],
                    "methods": list(methods.keys())[:4],
                })
            return {
                "global": g,
                "lines": self.lines_processed,
                "top_rows": top_rows,
                "qsize": self._q.qsize(),
            }

    def run(self):
        last_emit = 0.0
        while not self._stop.is_set():
            got = False
            try:
                line = self._q.get(timeout=0.25)
                got = True
            except queue.Empty:
                line = None
            if got:
                try:
                    with self._lock:
                        self._parse_one(line)
                        # 同批尽量多消化，减少锁切换
                        for _ in range(200):
                            try:
                                line = self._q.get_nowait()
                            except queue.Empty:
                                break
                            self._parse_one(line)
                except Exception as e:
                    print(f"remote parse error: {e}")
            now = time.monotonic()
            if self.lines_processed and (now - last_emit) >= 2.5:
                last_emit = now
                try:
                    snap = self._snapshot()
                    self.stats_tick.emit(snap)
                    self.dynamic_updates.emit(self.url_stats)
                except Exception as e:
                    print(f"remote stats emit error: {e}")
        # 退出前再刷一次
        if self.lines_processed:
            try:
                self.stats_tick.emit(self._snapshot())
            except Exception:
                pass


class AIAnalysisThread(QThread):
    result_signal = pyqtSignal(str)
    update_content = pyqtSignal(str)
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)

    def __init__(self, model_type, analysis_data, config, traffic_type):
        super().__init__()
        self.model_type = model_type
        self.analysis_data = analysis_data
        self.traffic_type = traffic_type
        self.config = config
        self._stop_event = threading.Event()
        self._request_session = requests.Session()

    def run(self):
        try:
            if self.model_type.startswith("本地模型"):
                self._analyze_with_local_model()
            elif self.model_type == "DeepSeek":
                self._analyze_with_deepseek()
            elif self.model_type == "OpenAI":
                self._analyze_with_openai()
            elif self.model_type == "Gemini":
                self._analyze_with_gemini()

            self.finished_signal.emit("分析完成")

        except Exception as e:
            self.error_signal.emit(f"分析过程中发生错误: {str(e)}")
        finally:
            self._request_session.close()

    def _analyze_with_local_model(self):
        """使用本地Ollama模型分析"""
        ollama_config = self.config.get('ollama', {})
        api_url = f"{ollama_config.get('url')}/api/generate"
        model_name = ollama_config.get('model_name')

        if not api_url or not model_name:
            self.error_signal.emit("本地模型配置不完整")
            return

        for prompt in self.analysis_data:
            if self._stop_event.is_set():
                break

            self.update_content.emit(prompt)

            try:
                response = self._request_session.post(
                    api_url,
                    json={
                        "model": model_name,
                        "prompt": prompt,
                        "stream": True
                    },
                    stream=True,
                    timeout=60
                )
                response.raise_for_status()

                for line in response.iter_lines():
                    if self._stop_event.is_set():
                        break

                    if line:
                        try:
                            data = json.loads(line.decode())
                            if "response" in data:
                                self.result_signal.emit(data["response"])
                        except json.JSONDecodeError:
                            continue

            except requests.exceptions.RequestException as e:
                self.error_signal.emit(f"连接本地模型失败: {str(e)}")
                return

    def stop(self):
        """停止分析"""
        self._stop_event.set()


class FullscreenWindow(QMainWindow):
    def __init__(self, widget, title, manager, lang_manager=None):
        super().__init__(manager.parent)
        self.manager = manager
        self.widget = widget
        self.lang_manager = lang_manager
        self.setWindowTitle(title)
        self.setWindowFlags(Qt.WindowType.Window)
        # 工具栏退出按钮
        toolbar = QToolBar(self.lang_manager.tr("fullscreen_toolbar", "全屏工具栏"), self)
        exit_action = QAction(QIcon("ico/exit_fullscreen.png"), self.lang_manager.tr("exit_fullscreen", "退出全屏"),
                              self)
        exit_action.triggered.connect(self.manager.exit_fullscreen)
        toolbar.addAction(exit_action)
        self.addToolBar(toolbar)
        # 中央显示区域
        central = QWidget()
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.widget)
        self.setCentralWidget(central)

    def closeEvent(self, event):
        # 确保点击×也退出全屏并回归原位
        self.manager.exit_fullscreen()
        super().closeEvent(event)


class FullscreenManager:
    def __init__(self, parent, lang_manager=None):
        self.parent = parent
        self.fullscreen_widget = None
        self.fullscreen_window = None
        self.original_layout = None
        self.lang_manager = lang_manager

    def enter_fullscreen(self, widget, title):
        """进入全屏模式"""
        if self.fullscreen_widget:
            return

        self.fullscreen_widget = widget
        # 保存原父布局
        self.original_layout = widget.parent().layout()
        self.original_layout.removeWidget(widget)

        # 创建并显示全屏窗口
        self.fullscreen_window = FullscreenWindow(widget, title, self, lang_manager=self.lang_manager)
        self.fullscreen_window.showMaximized()

    def exit_fullscreen(self):
        """退出全屏模式"""
        if not self.fullscreen_widget:
            return

        # 从全屏窗口移除控件
        self.fullscreen_window.centralWidget().layout().removeWidget(self.fullscreen_widget)

        # 将控件返回原布局
        self.original_layout.addWidget(self.fullscreen_widget)

        # 关闭全屏窗口
        self.fullscreen_window.close()

        # 清理状态
        self.fullscreen_widget = None
        self.fullscreen_window = None
        self.original_layout = None


class RegexHighlighter(QSyntaxHighlighter):
    """正则表达式语法高亮"""

    def __init__(self, document):
        super().__init__(document)

        self.highlightingRules = []

        # 元字符
        meta_format = QTextCharFormat()
        meta_format.setForeground(QColor("#FF6600"))
        meta_format.setFontWeight(QFont.Weight.Bold)
        meta_chars = r"[\.\^\$\*\+\?\{\}\[\]\\\|\(\)]"
        self.highlightingRules.append((QRegularExpression(meta_chars), meta_format))

        # 字符类
        char_class_format = QTextCharFormat()
        char_class_format.setForeground(QColor("#0099FF"))
        self.highlightingRules.append((QRegularExpression(r"\[.*?\]"), char_class_format))

        # 量词
        quantifier_format = QTextCharFormat()
        quantifier_format.setForeground(QColor("#9900FF"))
        quantifiers = r"\*|\+|\?|\{\d+,?\d*\}"
        self.highlightingRules.append((QRegularExpression(quantifiers), quantifier_format))

        # 分组
        group_format = QTextCharFormat()
        group_format.setForeground(QColor("#00AA00"))
        self.highlightingRules.append((QRegularExpression(r"\(.*?\)"), group_format))

        # 转义序列
        escape_format = QTextCharFormat()
        escape_format.setForeground(QColor("#FF0000"))
        self.highlightingRules.append((QRegularExpression(r"\\."), escape_format))

        # 注释
        comment_format = QTextCharFormat()
        comment_format.setForeground(QColor("#999999"))
        comment_format.setFontItalic(True)
        self.highlightingRules.append((QRegularExpression(r"#.*$"), comment_format))

    def highlightBlock(self, text):
        for pattern, format in self.highlightingRules:
            iterator = pattern.globalMatch(text)
            while iterator.hasNext():
                match = iterator.next()
                self.setFormat(match.capturedStart(), match.capturedLength(), format)


def _ai_fmt(color: str, bold: bool = False, italic: bool = False) -> QTextCharFormat:
    fmt = QTextCharFormat()
    fmt.setForeground(QColor(color))
    if bold:
        fmt.setFontWeight(QFont.Weight.Bold)
    if italic:
        fmt.setFontItalic(True)
    return fmt


class AgentLogHighlighter(QSyntaxHighlighter):
    """Agent 步骤日志配色（深灰底上的代码风格）"""

    def __init__(self, document):
        super().__init__(document)
        self.rules = [
            (QRegularExpression(r"^——\s*你\s*——.*"), _ai_fmt("#CE9178", bold=True)),
            (QRegularExpression(r"^——\s*Agent\s*——.*"), _ai_fmt("#4EC9B0", bold=True)),
            (QRegularExpression(r"^\[step\s+\d+\].*"), _ai_fmt("#DCDCAA")),
            (QRegularExpression(r"🔧\s*\S+"), _ai_fmt("#4FC1FF", bold=True)),
            (QRegularExpression(r"^\s*→\s*.*"), _ai_fmt("#9CDCFE")),
            (QRegularExpression(r"^✅.*"), _ai_fmt("#6A9955", bold=True)),
            (QRegularExpression(r"^❌.*"), _ai_fmt("#F44747", bold=True)),
            (QRegularExpression(r"^….*|^⚠.*"), _ai_fmt("#D7BA7D")),
            (QRegularExpression(r"模型:.*|Agent 端点:.*|代理:.*|统计文件:.*|流量明文:.*"), _ai_fmt("#808080")),
            (QRegularExpression(r"https?://\S+"), _ai_fmt("#569CD6")),
            (QRegularExpression(r"\b(?:risk|stats|search|artifact)\.\w+\b"), _ai_fmt("#C586C0")),
            (QRegularExpression(r"\b(?:summary|list|get|overview|top_urls|top_ips|url_detail|query|read|info)\b"),
             _ai_fmt("#4EC9B0")),
        ]

    def highlightBlock(self, text):
        for pattern, fmt in self.rules:
            it = pattern.globalMatch(text)
            while it.hasNext():
                m = it.next()
                self.setFormat(m.capturedStart(), m.capturedLength(), fmt)


class AgentReportHighlighter(QSyntaxHighlighter):
    """研判结论配色"""

    def __init__(self, document):
        super().__init__(document)
        self.rules = [
            (QRegularExpression(r"^【[^】]+】.*"), _ai_fmt("#DCDCAA", bold=True)),
            (QRegularExpression(r"^\s*#{1,3}\s+.*"), _ai_fmt("#569CD6", bold=True)),
            (QRegularExpression(r"\*\*[^*]+\*\*"), _ai_fmt("#CE9178", bold=True)),
            (QRegularExpression(r"`[^`]+`"), _ai_fmt("#CE9178")),
            (QRegularExpression(r"\b(?:严重|高危)\b"), _ai_fmt("#F44747", bold=True)),
            (QRegularExpression(r"\b中危\b"), _ai_fmt("#D7BA7D", bold=True)),
            (QRegularExpression(r"\b(?:低危|信息)\b"), _ai_fmt("#6A9955", bold=True)),
            (QRegularExpression(r"\b(?:建议|依据|结论)\b"), _ai_fmt("#4EC9B0", bold=True)),
            (QRegularExpression(r"https?://\S+"), _ai_fmt("#569CD6")),
            (QRegularExpression(r"\b(?:\d{1,3}\.){3}\d{1,3}\b"), _ai_fmt("#B5CEA8")),
            (QRegularExpression(r"(?i)\b(?:sql|xss|rce|ssrf|lfi|csrf|traversal|injection)\b"),
             _ai_fmt("#C586C0")),
            (QRegularExpression(r"^[-*•]\s+.*"), _ai_fmt("#D4D4D4")),
        ]

    def highlightBlock(self, text):
        # 默认正文色
        self.setFormat(0, len(text), _ai_fmt("#D4D4D4"))
        for pattern, fmt in self.rules:
            it = pattern.globalMatch(text)
            while it.hasNext():
                m = it.next()
                self.setFormat(m.capturedStart(), m.capturedLength(), fmt)


AI_CONSOLE_QSS = """
QPlainTextEdit, QTextEdit {
    background-color: #1E1E1E;
    color: #D4D4D4;
    border: 1px solid #3C3C3C;
    border-radius: 6px;
    padding: 8px;
    selection-background-color: #264F78;
    selection-color: #FFFFFF;
    font-family: Consolas, "Cascadia Mono", "Microsoft YaHei UI", monospace;
    font-size: 13px;
}
"""


class AnalysisThread(QThread):
    result_signal = pyqtSignal(str)  # 结果信号
    result_signal_extract = pyqtSignal(dict)  # 文件提取
    handle_replay_response = pyqtSignal(str)  # 结果信号
    status_label = pyqtSignal(str)  # 结果信号
    finished_signal = pyqtSignal(object)  # 可以传如何参数
    request_found_signal = pyqtSignal()  # 任务完成信号
    progress_signal = pyqtSignal(object)  # 进度信号

    def __init__(self, file, uri, keyword, output, request_only, response_only, show_body, request_stream_id=None,
                 sslkeylogfile=None, fileextraction=None, ai_analysis_starts=None, optional_parameters=None,
                 lang_manager=None):
        super().__init__()
        self.result_cache = []
        self.file = file
        self.uri = uri
        self.keyword = keyword
        self.lang_manager = lang_manager
        self.output = output
        self.optional_parameters = optional_parameters
        self.request_only = request_only
        self.response_only = response_only
        self.show_body = show_body
        self.ai_analysis_storing_data = ai_analysis_starts  # AI 分析 会记录全部的请求响应数据
        self.sslkeylogfile = sslkeylogfile  # 追加的 SSL 密钥日志文件路径
        self.request_stream_id = request_stream_id
        self.fileextraction = fileextraction
        # if self.fileextraction:  # 判断是否文件读取
        self.result_queue = queue.Queue()
        self._send_thread = threading.Thread(target=self._emit_loop, daemon=True)
        self._send_thread.start()
        self.last_result = None
        self.analysis_similar = "pyshark"  # 默认使用pyshark

    def _emit_loop(self):
        """后台发送线程，专门负责发 result_signal"""
        while True:
            try:
                msg = self.result_queue.get(timeout=1)
                self.result_signal.emit(msg)
                # time.sleep(0.5)  # 控制一下发送频率，防止挤爆 UI
            except queue.Empty:
                continue

    def data_processing(self, result):
        if self.request_stream_id is not None and result['stream_id'] == self.request_stream_id and "Request" == result[
            'http_type']:
            self.result_queue.put(
                f"{self.lang_manager.tr('build_request_data', '构建发送请求数据')}: \n"
                + output_filtering.complete_data(result)
            )
            self.result_queue.put(replay_request.build_send(result, proxies))
        elif self.filter_result(result) and self.request_stream_id is None:
            self.result_queue.put(output_filtering.visual_output(result, self.show_body))

    def file_extraction(self, result):
        """ 文件提取模块 """
        config_path = "config.yaml"
        hex_data = result['file_data']
        os.makedirs(self.fileextraction['save_path'], exist_ok=True)
        signatures = load_signatures(config_path, self.fileextraction['file_filter'])
        result = extract_file(hex_data, signatures, self.fileextraction['save_path'], result['uri'])
        if result['filename']:
            self.result_signal_extract.emit(result)

    def run_tshark_analysis(self):
        """使用tshark进行分析"""
        session_data = []
        url_count = INITIALIZE_DEFAULTDICT()

        # 获取文件大小用于计算进度
        file_size = os.path.getsize(self.file)
        processed_size = 0

        # 创建并设置新的事件循环
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        process = core_processing.based_on_tshark(self.file, self.sslkeylogfile)

        for line in process.stdout:

            # 更新已处理大小
            processed_size += len(line.encode('utf-8'))

            # 计算进度百分比
            progress = int((processed_size / file_size) * 100)
            self.progress_signal.emit(progress)  # 发送给进度条

            result = core_processing.process_tshark_line(line, url_count, session_data=session_data,
                                                         optional_parameters=self.optional_parameters)

            if result['url'] and result['method'] and progress != 100:
                self.progress_signal.emit(
                    f'{self.lang_manager.tr("analysis_done", "完成请求分析/安全检测（请等待分析完成！）：")}：{result["method"]} | {result["url"]}'
                )

            # ======================== 修改开始 ========================
            # 修改：文件提取部分，只更新进度状态，不打印二进制乱码
            if self.fileextraction and result['file_data']:
                self.file_extraction(result)

                # 这里改为输出进度提示
                self.status_label.emit(
                    f"{self.lang_manager.tr('extracting_binary', '正在提取二进制文件...')} {progress}%"
                )
            # ======================== 修改结束 ========================

            if self.ai_analysis_storing_data and "Request" == result['http_type']:
                self.ai_analysis_storing_data["request"]['all'].append(result)

            self.data_processing(result)

        self.last_result = url_count

    def filter_result(self, result):
        """根据参数过滤结果"""
        if not result:
            return False
        # 请求/响应过滤
        if self.request_only and result['http_type'] != 'Request':
            return False
        if self.response_only and result['http_type'] != 'Response':
            return False

        # URI过滤
        if self.uri and self.uri not in result['url']:
            return False

        # 关键字过滤
        if self.keyword and self.keyword not in output_filtering.visual_output(result, self.show_body):
            return False

        return True

    def run(self):
        """运行流量分析"""
        # self.result_signal.emit(f"开始分析，使用方式: {self.analysis_similar}")
        start = time.perf_counter()  # 精度高，推荐

        # 初始化
        if self.ai_analysis_storing_data:
            self.ai_analysis_storing_data["request"]['all'] = []


        if self.analysis_similar == "tshark":
            self.run_tshark_analysis()

        elapsed = time.perf_counter() - start
        if self.ai_analysis_storing_data:  # AI分析判断
            self.finished_signal.emit(self.ai_analysis_storing_data)
        else:

            self.finished_signal.emit(None)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.go_process = None
        self.go_port = GO_SERVER_PORT
        self.go_base_url = GO_BASE_URL
        self.loading_overlay = None
        self._analysis_ui_running = False
        self._analysis_stop_requested = False

        self.gui_update_timer = QTimer(self)
        self.gui_update_timer.setSingleShot(True)  # 设置为单次触发
        self.gui_update_timer.setInterval(2000)  # 远程流式默认 2s 刷一次，避免卡顿
        self.gui_update_timer.timeout.connect(self._perform_gui_update)
        self._pending_gui_data = None  # 用来暂存数据
        self._remote_streaming = False
        self._remote_ui_light_ticks = 0

        # 确保输出目录存在（安装包：exe 旁 output/）
        self.output_dir = get_output_dir()
        self.file_writer = FileWriter()

        # 加载初始配置
        self.config = module.load_config()

        # ===========界面配置
        self.ui_config_path = "ui_config.yaml"
        self.ui_config = {}
        self.load_ui_settings()
        self.setMinimumSize(self.ui_config['Interface_height'], self.ui_config['Interface_width'])
        # ===========界面配置

        # 初始化语言管理器
        self.lang_manager = LanguageManager(self.ui_config.get("language", "zh"))
        self.setWindowTitle(self.lang_manager.tr("app_title"))

        self.ai_analysis_pending = False  # 添加状态标志

        self.ai_analysis_in_progress = False  # 添加进行中标志

        self.run_ai = False

        # 设置应用图标
        self.setWindowIcon(QIcon("ico/l.png"))

        # 初始化UI
        self.init_ui()

        # 初始化统计数据结构
        self.url_stats = INITIALIZE_DEFAULTDICT()
        # Agent 可读的最近分析落盘路径
        self.last_stats_json_path = ""
        self.last_traffic_path = ""
        # 风险/Agent 绑定的源文件（与 Import_box 同步，避免串结果）
        self.last_analysis_source = ""

        # 分析线程
        self.analysis_thread = None
        self.replay_thread = None

    def init_ui(self, *_):
        """初始化主界面"""
        # 设置主窗口样式
        # print(self.ui_config['Main_Window_Style'])
        self.setStyleSheet(self.ui_config['Main_Window_Style'])

        # 创建菜单栏
        self.create_menu_bar()

        # 创建工具栏
        self.create_tool_bar()

        # 创建主内容区域
        self.create_main_content()

        # 创建状态栏
        self.create_status_bar()

        # 创建报告
        self.create_report_tab()
        self.create_settings_tab()
        # 设置中央部件
        central_widget = QWidget()
        central_widget.setLayout(QVBoxLayout())
        central_widget.layout().setContentsMargins(0, 0, 0, 0)
        central_widget.layout().setSpacing(0)
        central_widget.layout().addWidget(self.main_content)
        self.setCentralWidget(central_widget)

        # 添加进度条变量
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)

        # 在状态栏中添加进度条
        status_bar = self.statusBar()
        status_bar.addPermanentWidget(self.progress_bar)
        status_bar.insertWidget(0, self.progress_bar)  # 插到最左端
        self.progress_bar.hide()  # 默认隐藏

    def create_menu_bar(self, *_):
        """创建菜单栏"""
        menubar = self.menuBar()

        # 视图菜单
        view_menu = menubar.addMenu(self.lang_manager.tr("menu_view"))

        toolbar_action = QAction(self.lang_manager.tr("toolbar_toggle"), self)
        toolbar_action.setCheckable(True)
        toolbar_action.setChecked(True)
        toolbar_action.triggered.connect(self.toggle_toolbar)
        view_menu.addAction(toolbar_action)

        statusbar_action = QAction(self.lang_manager.tr("statusbar_toggle"), self)
        statusbar_action.setCheckable(True)
        statusbar_action.setChecked(True)
        statusbar_action.triggered.connect(self.toggle_statusbar)
        view_menu.addAction(statusbar_action)

        # 帮助菜单
        help_menu = menubar.addMenu(self.lang_manager.tr("menu_help"))

        about_action = QAction(QIcon("ico/info.png"), self.lang_manager.tr("about"), self)
        about_action.triggered.connect(self.show_about_dialog)
        help_menu.addAction(about_action)

        docs_action = QAction(QIcon("ico/help.png"), self.lang_manager.tr("documentation"), self)
        docs_action.triggered.connect(self.open_documentation)
        help_menu.addAction(docs_action)

        # 添加语言切换菜单
        lang_menu = menubar.addMenu(self.lang_manager.tr("menu_language", "Language"))

        zh_action = QAction("中文", self)
        zh_action.triggered.connect(lambda: self.change_language("zh"))
        lang_menu.addAction(zh_action)

        en_action = QAction("English", self)
        en_action.triggered.connect(lambda: self.change_language("en"))
        lang_menu.addAction(en_action)

    def change_language(self, lang_code):
        """切换语言"""

        if self.lang_manager.get_current_language() != lang_code:
            if self.lang_manager.load_language(lang_code):
                # 将语言设置保存到配置
                self.ui_config['language'] = lang_code
                self.save_ui_settings()

                # --- 新增: 提示 + 优雅退出 ---
                QMessageBox.information(
                    self,
                    self.lang_manager.tr("prompt"),  # 如 “提示”
                    self.lang_manager.tr("msg_restart")  # 如 “修改成功，请重新进入应用”
                )
                # singleShot(0, ...) 让 quit() 排到消息队列尾部，
                # 等提示框关闭后再真正退出
                QTimer.singleShot(0, QApplication.instance().quit)

    def update_status_label(self, status, *_):

        self.status_label.setText(status)

    def Analysis_and_selection(self, dialog, optional_parameters, *_):
        """ 分析选择 """
        selected_option = self.file_filter_combo.currentText()
        # print(f"用户选择了: {selected_option}")
        self.add_recent_activity(self.lang_manager.tr("start_analysis", "开始分析"), self.Import_box.text(),
                                 f"{selected_option}")
        dialog.accept()  # 或者 dialog.close()
        if selected_option == "不识别":
            self.start_analysis(optional_parameters=optional_parameters)
        else:
            self.start_analysis(optional_parameters=optional_parameters)
            self.start_http_extraction()

    def start_log_analysis(self, dialog, optional_parameters_log):
        dialog.accept()  # 关闭对话框
        self.analyze_logs(log_type='auto', optional_parameters_log=optional_parameters_log)

    def _is_analysis_running(self) -> bool:
        if getattr(self, "_analysis_ui_running", False):
            return True
        for name in ("analysis_thread", "worker_thread"):
            th = getattr(self, name, None)
            try:
                if th is not None and th.isRunning():
                    return True
            except Exception:
                pass
        return False

    def _set_analyze_button_running(self, running: bool):
        """分析中：按钮变为「结束分析」并可点击；空闲时恢复「一键识别」。"""
        self._analysis_ui_running = bool(running)
        idle_text = self.lang_manager.tr("auto_analysis", "一键识别自动化分析")
        stop_text = self.lang_manager.tr("end_analysis", "结束分析")
        if hasattr(self, "btn_analyze") and self.btn_analyze is not None:
            self.btn_analyze.setEnabled(True)
            if running:
                self.btn_analyze.setText(stop_text)
                self.btn_analyze.setIcon(QIcon("ico/stop.png"))
                self.btn_analyze.setToolTip(stop_text)
            else:
                self.btn_analyze.setText(idle_text)
                self.btn_analyze.setIcon(QIcon("ico/analyze.png"))
                self.btn_analyze.setToolTip(idle_text)
        if hasattr(self, "analyze_btn") and self.analyze_btn is not None:
            self.analyze_btn.setEnabled(True)
            if running:
                self.analyze_btn.setIcon(QIcon("ico/stop.png"))
                self.analyze_btn.setToolTip(stop_text)
            else:
                self.analyze_btn.setIcon(QIcon("ico/analyze.png"))
                self.analyze_btn.setToolTip(self.lang_manager.tr("start_analysis", "开始分析"))
        stop_tb = getattr(self, "toolbar_stop_btn", None)
        if stop_tb is not None:
            stop_tb.setEnabled(bool(running))

    def _on_analyze_button_clicked(self, *_):
        """一键分析 / 结束分析 切换。"""
        if self._is_analysis_running():
            self.stop_analysis()
            return
        self.automatically_determine_the_analysis_type()

    def automatically_determine_the_analysis_type(self, *_):
        """ 一键自动化 """
        if self._is_analysis_running():
            self.stop_analysis()
            return

        file = self.Import_box.text()
        if not file:
            QMessageBox.warning(self, self.lang_manager.tr("prompt"),
                                self.lang_manager.tr("Please_select_the_traffic_file_first", "请先选择流量文件!"))
            self._set_analyze_button_running(False)
            self.status_label.setText(self.lang_manager.tr("Please_select_the_traffic_file_first", "请先选择流量文件!"))
            return
        # 设置AI分析标志：分析完成后自动跑 Agent
        self._pending_agent_review = bool(
            hasattr(self, "ai_auto_analyze_check") and self.ai_auto_analyze_check.isChecked()
        )


        if file.lower().endswith(('.csv', '.tsv')):
            self.start_table_analysis()
            return

        if file.lower().endswith(('.log', '.txt')):

            # 启动日志文件分析对话框
            optional_parameters_log = {
                "URI_Security_Check": False
            }

            dialog = QDialog(self)
            dialog.setWindowTitle(self.lang_manager.tr("log_uri_security_title", "日志文件 URI 安全检测"))
            dialog.setWindowIcon(QIcon("ico/extract.png"))
            dialog.resize(400, 150)

            layout = QVBoxLayout(dialog)
            layout.setContentsMargins(15, 15, 15, 15)
            layout.setSpacing(15)

            # 添加说明文字
            label = QLabel(self.lang_manager.tr(
                "log_uri_security_label",
                "检测到是日志文件，请选择是否启用 URI 安全检测："
            ))

            layout.addWidget(label)

            # 勾选框
            cb_uri = QCheckBox(self.lang_manager.tr(
                "enable_uri_security_check",
                "启用 URI 安全检测 （检测会消耗较长时间！）"
            ))
            cb_uri.setChecked(False)
            cb_uri.toggled.connect(lambda v: optional_parameters_log.__setitem__("URI_Security_Check", v))
            layout.addWidget(cb_uri)

            # 按钮区域
            btn_layout = QHBoxLayout()
            btn_layout.setSpacing(15)

            start_btn = QPushButton(self.lang_manager.tr("start_scan", "开始扫描"))
            start_btn.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 8px 16px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #45a049;
                }
                QPushButton:pressed {
                    background-color: #3e8e41;
                }
            """)
            start_btn.clicked.connect(lambda: self.start_log_analysis(dialog, optional_parameters_log))

            btn_layout.addStretch()
            btn_layout.addWidget(start_btn)
            layout.addLayout(btn_layout)
            if not dialog.exec() == QDialog.DialogCode.Accepted:
                self._set_analyze_button_running(False)
                self.status_label.setText(self.lang_manager.tr("canceled", "取消了"))

        elif file.lower().endswith(('.pcap', '.pcapng', '.cap')):
            # 创建文件提取选择对话框
            optional_parameters = {
                "URL_Security_Check": False,
                "Request_Head_Security_Check": False,
                "Data_section_detection": {
                    "enabled": False,  # ← 是否勾选“请求体安全检测”总开关
                    "binary": False,
                    "limit_size": 0.5,  # 设置最大检测请求体大小 最大1MB 默认是0.5MB
                    "forms": False,
                    "json": False,
                    "xml": False,
                    "multipart": False
                }
            }

            dialog = QDialog(self)
            dialog.setWindowTitle(self.lang_manager.tr("full_traffic_check_options", "全流量检测选项"))
            dialog.setWindowIcon(QIcon("ico/extract.png"))
            dialog.resize(400, 200)

            # 设置对话框样式
            dialog.setStyleSheet("""
                QDialog {
                    background-color: #2D2D2D;
                    color: #DDDDDD;
                }
                QLabel {
                    color: #EEEEEE;
                }
                QGroupBox {
                    border: 1px solid #444;
                    border-radius: 5px;
                    margin-top: 10px;
                    padding-top: 15px;
                    color: #AAAAAA;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 5px;
                }
            """)

            main_layout = QVBoxLayout(dialog)
            main_layout.setContentsMargins(15, 15, 15, 15)
            main_layout.setSpacing(15)

            # 添加说明标签
            label = QLabel(self.lang_manager.tr("detected_full_traffic_file", "检测到是全流量文件"))
            label.setStyleSheet("""
                QLabel {
                    font-size: 14px;
                    font-weight: bold;
                    color: #E0E0E0;
                }
            """)
            main_layout.addWidget(label)

            # 添加文件类型选择区域
            type_group = QGroupBox("选择要提取的文件类型")
            type_layout = QVBoxLayout(type_group)
            type_layout.setContentsMargins(10, 15, 10, 10)
            type_layout.setSpacing(10)

            # 读取config.yaml文件获取文件类型
            try:
                signatures = self.config.get('signatures', [])
                file_types = sorted(
                    {sig.get('type', '') for sig in signatures if
                     sig.get('enabled', False) and sig.get('type', '')})
            except Exception as e:
                print(f"Error loading config.yaml: {e}")
                file_types = []

            # 默认选项
            default_options = ["所有文件类型"]
            file_types = default_options + sorted(set(file_types) - set(default_options))

            # 文件类型过滤下拉框和标签
            filter_layout = QHBoxLayout()
            filter_layout.setSpacing(10)

            filter_label = QLabel("文件类型:")

            self.file_filter_combo = QComboBox()
            self.file_filter_combo.addItem("不识别")
            self.file_filter_combo.addItem("识别提取常见序列化二数据")
            if file_types:
                self.file_filter_combo.insertSeparator(1)
                for file_type in file_types:
                    self.file_filter_combo.addItem(f"{file_type} 文件")

            # filter_layout.addWidget(filter_label)
            # filter_layout.addWidget(self.file_filter_combo)
            # type_layout.addLayout(filter_layout)
            # main_layout.addWidget(type_group)
            check_group = QGroupBox(
                self.lang_manager.tr("select_check_items", "选择检测项---选择后分析可能会消耗较多时间"))
            check_layout = QVBoxLayout(check_group)
            check_layout.setContentsMargins(10, 15, 10, 10)
            check_layout.setSpacing(8)

            # — URL 安全检测 —
            cb_url = QCheckBox(self.lang_manager.tr("url_security_check", "URL 安全检测"))
            cb_url.setChecked(optional_parameters["URL_Security_Check"])
            cb_url.toggled.connect(
                lambda v: optional_parameters.__setitem__("URL_Security_Check", v)
            )
            check_layout.addWidget(cb_url)

            # — 请求头安全检测 —
            cb_head = QCheckBox(self.lang_manager.tr("request_header_security_check", "请求头安全检测"))
            cb_head.setChecked(optional_parameters["Request_Head_Security_Check"])
            cb_head.toggled.connect(
                lambda v: optional_parameters.__setitem__("Request_Head_Security_Check", v)
            )
            check_layout.addWidget(cb_head)

            # — 请求体总开关 —
            cb_body = QCheckBox(self.lang_manager.tr("request_body_security_check", "请求体安全检测"))
            cb_body.setChecked(False)
            check_layout.addWidget(cb_body)

            # 子项布局区域（缩进显示）
            body_sub_layout = QVBoxLayout()
            body_sub_layout.setContentsMargins(25, 0, 0, 0)  # 添加缩进效果
            body_sub_layout.setSpacing(4)

            # 子项勾选框集合
            body_checkboxes = {}

            for key, label in {
                "binary": self.lang_manager.tr("binary_detection", "二进制检测"),
                "forms": self.lang_manager.tr("forms_detection", "表单检测"),
                "json": self.lang_manager.tr("json_detection", "JSON 检测"),
                "xml": self.lang_manager.tr("xml_detection", "XML 检测"),
                "multipart": self.lang_manager.tr("multipart_detection", "文件上传检测")
            }.items():
                cb = QCheckBox(label)
                cb.setEnabled(False)  # 初始跟随总开关
                cb.toggled.connect(lambda v, k=key: optional_parameters["Data_section_detection"].__setitem__(k, v))
                body_checkboxes[key] = cb
                body_sub_layout.addWidget(cb)
            size_layout = QHBoxLayout()
            size_layout.setContentsMargins(0, 0, 0, 0)
            size_layout.setSpacing(10)

            size_label = QLabel(self.lang_manager.tr("request_body_limit", "请求体检测大小限制 (最大1MB)："))
            size_label.setStyleSheet("color: #CCCCCC;")

            size_input = QDoubleSpinBox()
            size_input.setStyleSheet("""
                QDoubleSpinBox {
                    background-color: #3A3A3A;
                    color: #FFFFFF;
                    border: 1px solid #666666;
                    border-radius: 3px;
                    padding: 4px;
                }
                QDoubleSpinBox::up-button, QDoubleSpinBox::down-button {
                    background-color: #444444;
                    subcontrol-origin: border;
                    width: 12px;
                }
                QDoubleSpinBox::up-button:hover, QDoubleSpinBox::down-button:hover {
                    background-color: #555555;
                }
                QDoubleSpinBox::up-arrow, QDoubleSpinBox::down-arrow {
                    width: 6px;
                    height: 6px;
                }
            """)

            size_input.setRange(0.1, 1.0)
            size_input.setSingleStep(0.1)
            size_input.setValue(optional_parameters["Data_section_detection"]["limit_size"])
            size_input.setSuffix(" MB")
            size_input.setDecimals(1)
            size_input.setEnabled(False)

            size_input.valueChanged.connect(
                lambda v: optional_parameters["Data_section_detection"].__setitem__("limit_size", v)
            )

            size_layout.addWidget(size_label)
            size_layout.addWidget(size_input)
            body_sub_layout.addLayout(size_layout)

            def update_size_input_state(v):
                size_input.setEnabled(v)

            cb_body.toggled.connect(update_size_input_state)
            check_layout.addLayout(body_sub_layout)

            # 同步启用/禁用子项
            def update_body_checks(v):
                optional_parameters["Data_section_detection"]["enabled"] = v
                for cb in body_checkboxes.values():
                    cb.setEnabled(v)

            cb_body.toggled.connect(update_body_checks)

            main_layout.addWidget(check_group)

            # 添加操作按钮
            btn_layout = QHBoxLayout()
            btn_layout.setSpacing(15)

            extract_btn = QPushButton(self.lang_manager.tr("start_analysis", "开始分析"))
            extract_btn.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 8px 16px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #45a049;
                }
                QPushButton:pressed {
                    background-color: #3e8e41;
                }
            """)
            extract_btn.clicked.connect(lambda: self.Analysis_and_selection(dialog, optional_parameters))

            btn_layout.addStretch()
            btn_layout.addWidget(extract_btn)
            main_layout.addLayout(btn_layout)

            if not dialog.exec() == QDialog.DialogCode.Accepted:
                # 用户取消
                self._set_analyze_button_running(False)
                self.status_label.setText(self.lang_manager.tr("canceled", "取消了"))
        else:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "判断不出来文件类型")
            self._set_analyze_button_running(False)
            return

    def toggle_toolbar(self, visible, *_):
        """切换工具栏显示"""
        self.toolbar.setVisible(visible)

    def toggle_statusbar(self, visible, *_):
        """切换状态栏显示"""
        self.statusBar().setVisible(visible)

    def create_tool_bar(self, *_):
        """创建工具栏"""
        self.toolbar = QToolBar(self.lang_manager.tr("main_toolbar", "主工具栏"))
        self.toolbar.setIconSize(QSize(24, 24))

        self.addToolBar(self.toolbar)

        # 开始分析按钮（分析中切换为结束）
        self.analyze_btn = QToolButton()
        self.analyze_btn.setIcon(QIcon("ico/analyze.png"))
        self.analyze_btn.setToolTip(self.lang_manager.tr("start_analysis", "开始分析"))
        self.analyze_btn.clicked.connect(self._on_analyze_button_clicked)
        self.toolbar.addWidget(self.analyze_btn)

        # 停止分析按钮
        self.toolbar_stop_btn = QToolButton()
        self.toolbar_stop_btn.setIcon(QIcon("ico/stop.png"))
        self.toolbar_stop_btn.setToolTip(self.lang_manager.tr("stop_analysis", "停止分析"))
        self.toolbar_stop_btn.clicked.connect(self.stop_analysis)
        self.toolbar_stop_btn.setEnabled(False)
        self.toolbar.addWidget(self.toolbar_stop_btn)

        # 重放请求按钮
        replay_btn = QToolButton()
        replay_btn.setIcon(QIcon("ico/replay.png"))
        replay_btn.setToolTip(self.lang_manager.tr("replay_request", "重放请求"))
        replay_btn.clicked.connect(self.replay_request)
        self.toolbar.addWidget(replay_btn)

        self.toolbar.addSeparator()

        # 导出结果按钮
        export_btn = QToolButton()
        export_btn.setIcon(QIcon("ico/export.png"))
        export_btn.setToolTip(self.lang_manager.tr("export_results", "导出结果"))
        export_btn.clicked.connect(self.export_results)
        self.toolbar.addWidget(export_btn)

        # 添加分隔符
        self.toolbar.addSeparator()

        # 导入框和浏览按钮打包为一组
        file_input_widget = QWidget()
        file_input_layout = QHBoxLayout(file_input_widget)
        file_input_layout.setContentsMargins(0, 0, 0, 0)
        file_input_layout.setSpacing(4)

        self.Import_box = QLineEdit()
        self.Import_box.setPlaceholderText(
            f"CAP/PCAP/LOG/" + self.lang_manager.tr("import_file_placeholder", "CAP/PCAP/LOG/TXT文件...")
        )
        self.Import_box.setReadOnly(True)
        self.Import_box.setMinimumWidth(200)

        self.Import_file_button = QPushButton(
            QIcon("ico/open.png"), self.lang_manager.tr("browse", "浏览...")
        )
        self.Import_file_button.clicked.connect(self.select_file)

        file_input_layout.addWidget(self.Import_box)
        file_input_layout.addWidget(self.Import_file_button)

        self.toolbar.addWidget(file_input_widget)

    def create_main_content(self, *_):
        """创建主内容区域"""
        self.main_content = QWidget()
        main_layout = QHBoxLayout(self.main_content)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # 创建侧边栏
        self.sidebar = QListWidget()

        if self.ui_config.get("language", "zh") == "zh":  # 要是英文会加长
            self.sidebar.setFixedWidth(150)  # 稍微加宽一点
        else:
            self.sidebar.setFixedWidth(200)  # 稍微加宽一点
        self.sidebar.setStyleSheet(self.ui_config['Main_Left_Window_Style'])

        # # 添加侧边栏项目
        sidebar_items = [
            {"icon": "ico/dashboard.png", "text": self.lang_manager.tr("dashboard", "仪表盘"), "tab": "dashboard"},
            {"icon": "ico/analysis.png", "text": self.lang_manager.tr("analysis", "HTTP提取"), "tab": "analysis"},
            {"icon": "ico/stats.png", "text": self.lang_manager.tr("stats", "统计分析"), "tab": "stats"},
            {"icon": "ico/logs.png", "text": self.lang_manager.tr("dns_access", "DNS访问"), "tab": "dns"},
            {"icon": "ico/risk.png", "text": self.lang_manager.tr("risk", "风险分析"), "tab": "risk"},
            {"icon": "ico/replay.png", "text": self.lang_manager.tr("replay", "请求重放"), "tab": "replay"},
            {"icon": "ico/remote_monitoring.png", "text": self.lang_manager.tr("remote_monitoring", "远程监控"),
             "tab": "remote_monitoring"},
            {"icon": "ico/extract.png", "text": self.lang_manager.tr("extract", "文件提取"), "tab": "extract"},
            {"icon": "ico/logs.png", "text": self.lang_manager.tr("log", "Log分析"), "tab": "log"},
            {"icon": "ico/re_matching.png", "text": self.lang_manager.tr("regular_expression_matching", "规则配置"),
             "tab": "re_matching"},
            {"icon": "ico/intelligence.png", "text": self.lang_manager.tr("intelligence", "情报分析"),
             "tab": "intelligence"},
            {"icon": "ico/ai.png", "text": self.lang_manager.tr("ai", "AI分析"), "tab": "ai"},
            {"icon": "ico/report.png", "text": self.lang_manager.tr("report", "报告生成"), "tab": "report"},
            {"icon": "ico/settings.png", "text": self.lang_manager.tr("settings", "设置"), "tab": "settings"}
        ]

        for item in sidebar_items:
            list_item = QListWidgetItem(QIcon(item["icon"]), item["text"])
            list_item.setData(Qt.ItemDataRole.UserRole, item["tab"])
            list_item.setToolTip(item["text"])  # ★ 关键：设置悬浮文本
            list_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            list_item.setForeground(QColor("#B1B1B1"))  # 设置每个item的字体颜色
            list_item.setSizeHint(QSize(70, 50))  # 设置固定高度

            self.sidebar.addItem(list_item)

        # 默认选择仪表盘项
        self.sidebar.setCurrentRow(0)  # 设置为第一个项（仪表盘）
        self.sidebar.itemClicked.connect(self.switch_tab)
        main_layout.addWidget(self.sidebar)

        # 创建主工作区
        self.workspace = QTabWidget()
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(20)
        shadow.setOffset(0, 0)
        shadow.setColor(QColor(0, 0, 0, 20))
        self.workspace.setGraphicsEffect(shadow)
        self.workspace.setTabsClosable(True)
        self.workspace.tabCloseRequested.connect(self.close_tab)
        self.workspace.tabBar().hide()

        self.workspace.setStyleSheet(self.ui_config['Main_Right_Window_Style'])

        # 添加各个功能标签页
        self.tabs = {}
        self.create_dashboard_tab()
        self.create_intelligence_tab()
        self.create_risk_tab()
        self.create_ai_tab()
        self.regular_expression_matching()
        self.create_analysis_tab()
        self.create_stats_tab()
        self.create_dns_tab()
        self.create_replay_tab()
        self.create_extract_tab()
        self.create_log_tab()
        self.remote_monitoring_tab()

        main_layout.addWidget(self.workspace, stretch=1)

    def remote_monitoring_tab(self):
        """远程分析模块"""
        self.remote_monitoring_ui_tab = uic.loadUi("ui/remote_monitoring.ui")
        self.workspace.addTab(
            self.remote_monitoring_ui_tab,
            QIcon("ico/remote_monitoring.png"),
            "远程分析"
        )

        # 默认服务信息
        self.remote_monitoring_ui_tab.lineEdit_3.setText("127.0.0.1")
        self.remote_monitoring_ui_tab.lineEdit_4.setText("8765")

        # 连接状态
        self.remote_monitoring_ui_tab.connection_label = QLabel("未连接 ❌")
        self.remote_monitoring_ui_tab.connection_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.remote_monitoring_ui_tab.connection_label.setStyleSheet(
            "color: red; font-weight: bold; font-size: 14px;"
        )
        layout = QVBoxLayout(self.remote_monitoring_ui_tab.groupBox_4)
        layout.addWidget(self.remote_monitoring_ui_tab.connection_label)
        layout.setContentsMargins(10, 10, 10, 10)

        # 在顶部配置区追加：模式 / Token / 日志路径 / 一键生成
        cfg_box = self.remote_monitoring_ui_tab.groupBox
        cfg_layout = cfg_box.layout()
        if cfg_layout is None:
            cfg_layout = QHBoxLayout(cfg_box)

        self.remote_mode_combo = QComboBox()
        self.remote_mode_combo.addItem("本机连接服务器 (Agent serve)", "client")
        self.remote_mode_combo.addItem("服务器连接本机 (Agent dial)", "listen")

        self.remote_token_input = QLineEdit()
        self.remote_token_input.setPlaceholderText("鉴权 Token")
        self.remote_token_input.setText(secrets.token_urlsafe(12))
        self.remote_token_input.setMinimumWidth(140)

        btn_regen_token = QPushButton("随机Token")
        btn_regen_token.clicked.connect(self._regen_remote_token)

        self.remote_log_path_input = QLineEdit()
        self.remote_log_path_input.setPlaceholderText("服务器日志路径，如 /var/log/nginx/access.log")
        self.remote_log_path_input.setText("/var/log/nginx/access.log")
        self.remote_log_path_input.setMinimumWidth(220)

        btn_gen_agent = QPushButton("一键生成 Agent")
        btn_gen_agent.setToolTip("生成可在服务器上运行的脚本与命令")
        btn_gen_agent.clicked.connect(self._generate_remote_agent)

        # SSH 隧道（可选）
        self.remote_ssh_check = QCheckBox("使用 SSH 隧道（加密）")
        self.remote_ssh_check.setToolTip(
            "通过 SSH 本地/反向端口转发加密传输；Agent 只需监听 127.0.0.1"
        )
        self.remote_ssh_check.toggled.connect(self._on_remote_ssh_toggled)

        self.remote_ssh_user = QLineEdit()
        self.remote_ssh_user.setPlaceholderText("SSH 用户")
        self.remote_ssh_user.setText("root")
        self.remote_ssh_user.setMaximumWidth(100)

        self.remote_ssh_host = QLineEdit()
        self.remote_ssh_host.setPlaceholderText("SSH 主机（可与服务地址相同）")

        self.remote_ssh_port = QLineEdit()
        self.remote_ssh_port.setPlaceholderText("22")
        self.remote_ssh_port.setText("22")
        self.remote_ssh_port.setMaximumWidth(60)

        self.remote_ssh_key = QLineEdit()
        self.remote_ssh_key.setPlaceholderText("私钥路径（可选，默认用 ssh-agent）")

        btn_ssh_key = QPushButton("…")
        btn_ssh_key.setMaximumWidth(28)
        btn_ssh_key.clicked.connect(self._pick_remote_ssh_key)

        self.remote_ssh_local_port = QLineEdit()
        self.remote_ssh_local_port.setPlaceholderText("本地端口")
        self.remote_ssh_local_port.setText("8765")
        self.remote_ssh_local_port.setMaximumWidth(70)

        self.remote_ssh_panel = QWidget()
        ssh_row = QHBoxLayout(self.remote_ssh_panel)
        ssh_row.setContentsMargins(0, 0, 0, 0)
        ssh_row.addWidget(QLabel("SSH用户:"))
        ssh_row.addWidget(self.remote_ssh_user)
        ssh_row.addWidget(QLabel("SSH主机:"))
        ssh_row.addWidget(self.remote_ssh_host, stretch=1)
        ssh_row.addWidget(QLabel("SSH端口:"))
        ssh_row.addWidget(self.remote_ssh_port)
        ssh_row.addWidget(QLabel("密钥:"))
        ssh_row.addWidget(self.remote_ssh_key, stretch=1)
        ssh_row.addWidget(btn_ssh_key)
        ssh_row.addWidget(QLabel("本地端口:"))
        ssh_row.addWidget(self.remote_ssh_local_port)
        self.remote_ssh_panel.setVisible(False)

        hint = QLabel("未开 SSH 时为明文 ws://（仅建议内网/VPN）；勾选 SSH 隧道后流量经加密转发")
        hint.setStyleSheet("color: gray; font-size: 11px;")

        extra = QHBoxLayout()
        extra.addWidget(QLabel("模式:"))
        extra.addWidget(self.remote_mode_combo)
        extra.addWidget(QLabel("Token:"))
        extra.addWidget(self.remote_token_input)
        extra.addWidget(btn_regen_token)
        extra.addWidget(QLabel("日志:"))
        extra.addWidget(self.remote_log_path_input, stretch=1)
        extra.addWidget(btn_gen_agent)

        extra2 = QHBoxLayout()
        extra2.addWidget(self.remote_ssh_check)

        # 数据范围 / 恶意分析（仅：实时 / 全部历史）
        self.remote_hist_combo = QComboBox()
        self.remote_hist_combo.addItem("仅实时新日志（推荐）", "live")
        self.remote_hist_combo.addItem("全部历史 + 实时", "all")
        self.remote_hist_combo.setMinimumWidth(180)
        self.remote_hist_combo.setToolTip(
            "仅实时：只收连接之后的新访问\n全部历史：先推送日志文件已有内容，再继续实时"
        )
        self.remote_malware_check = QCheckBox("恶意流量分析（URI规则，较耗性能）")
        self.remote_malware_check.setChecked(False)
        # 兼容旧 UI 里的 checkBox
        try:
            old_cb = getattr(self.remote_monitoring_ui_tab, "checkBox", None)
            if old_cb is not None:
                old_cb.setText("恶意流量分析（URI规则）")
                old_cb.setChecked(False)
                self.remote_malware_check = old_cb
        except Exception:
            pass

        extra2.addWidget(QLabel("范围:"))
        extra2.addWidget(self.remote_hist_combo)
        extra2.addWidget(self.remote_malware_check)
        extra2.addStretch(1)

        # groupBox 原布局是水平的，改为垂直包一层
        parent_layout = self.remote_monitoring_ui_tab.widget_2.layout()
        if parent_layout is not None:
            main_v = self.remote_monitoring_ui_tab.layout()
            if main_v is not None:
                wrap = QWidget()
                wrap_l = QVBoxLayout(wrap)
                wrap_l.setContentsMargins(8, 0, 8, 4)
                wrap_l.addLayout(extra)
                wrap_l.addLayout(extra2)
                wrap_l.addWidget(self.remote_ssh_panel)
                wrap_l.addWidget(hint)
                main_v.insertWidget(1, wrap)

        # 获取控件
        self.addr_input: QLineEdit = self.remote_monitoring_ui_tab.lineEdit_3
        self.port_input: QLineEdit = self.remote_monitoring_ui_tab.lineEdit_4
        self.log_view: QTextEdit = self.remote_monitoring_ui_tab.textEdit
        self.btn_test: QPushButton = self.remote_monitoring_ui_tab.pushButton
        self.btn_start: QPushButton = self.remote_monitoring_ui_tab.pushButton_2
        self.btn_disconnect: QPushButton = self.remote_monitoring_ui_tab.pushButton_3

        # 状态属性
        self.remote_connected = False
        self.remote_log_thread = None
        self.remote_thread = None
        self.stream_analyzer = None
        self._remote_ssh_proc = None
        self.remote_timer = QTimer()
        self.remote_timer.timeout.connect(self._blink_status)
        self._status_toggle = False

        # 绑定按钮事件（控件就绪后再连模式切换）
        self.btn_test.clicked.connect(self._test_remote_connection)
        self.btn_start.clicked.connect(self.Analyze_logs_remotely)
        self.btn_disconnect.clicked.connect(self._disconnect_remote)
        self.remote_mode_combo.currentIndexChanged.connect(self._on_remote_mode_changed)
        self._on_remote_mode_changed()

        self.tabs["remote_monitoring"] = self.remote_monitoring_ui_tab

    def _remote_mode(self) -> str:
        return self.remote_mode_combo.currentData() or "client"

    def _remote_history_mode(self) -> str:
        """返回 live|all"""
        combo = getattr(self, "remote_hist_combo", None)
        if combo is None:
            return "live"
        return combo.currentData() or "live"

    def _remote_opts_line(self) -> str:
        if self._remote_history_mode() == "all":
            return "OPTS mode=all"
        return "OPTS mode=live"

    def _remote_malware_enabled(self) -> bool:
        cb = getattr(self, "remote_malware_check", None)
        return bool(cb is not None and cb.isChecked())

    def _ssh_tunnel_enabled(self) -> bool:
        return bool(getattr(self, "remote_ssh_check", None) and self.remote_ssh_check.isChecked())

    def _on_remote_ssh_toggled(self, checked: bool):
        self.remote_ssh_panel.setVisible(checked)
        if checked:
            # 默认用「服务地址」作为 SSH 主机
            if not self.remote_ssh_host.text().strip():
                h = self.addr_input.text().strip()
                if h and h not in ("0.0.0.0", "127.0.0.1", "localhost"):
                    self.remote_ssh_host.setText(h)
            lp = self.port_input.text().strip() or "8765"
            if not self.remote_ssh_local_port.text().strip():
                self.remote_ssh_local_port.setText(lp)
            self._append_remote_log(
                "[INFO] 已启用 SSH 隧道：模式A 用 -L 本地转发；模式B 用 -R 反向转发"
            )
        else:
            self._stop_ssh_tunnel()

    def _pick_remote_ssh_key(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "选择 SSH 私钥", "", "All Files (*);;PEM (*.pem);;Key (id_rsa *)"
        )
        if path:
            self.remote_ssh_key.setText(path)

    def _find_ssh_executable(self) -> str:
        import shutil
        exe = shutil.which("ssh")
        if exe:
            return exe
        # Windows OpenSSH 常见路径
        candidates = [
            r"C:\Windows\System32\OpenSSH\ssh.exe",
            r"C:\Program Files\Git\usr\bin\ssh.exe",
        ]
        for c in candidates:
            if os.path.isfile(c):
                return c
        return "ssh"

    def _ssh_tunnel_targets(self):
        """返回 (ssh_user, ssh_host, ssh_port, agent_port, local_port)"""
        ssh_host = self.remote_ssh_host.text().strip() or self.addr_input.text().strip()
        ssh_user = self.remote_ssh_user.text().strip() or "root"
        ssh_port = self.remote_ssh_port.text().strip() or "22"
        agent_port = self.port_input.text().strip() or "8765"
        local_port = self.remote_ssh_local_port.text().strip() or agent_port
        return ssh_user, ssh_host, ssh_port, agent_port, local_port

    def _start_ssh_tunnel(self) -> bool:
        """建立 SSH 端口转发；失败返回 False。"""
        self._stop_ssh_tunnel()
        if not self._ssh_tunnel_enabled():
            return True

        ssh_user, ssh_host, ssh_port, agent_port, local_port = self._ssh_tunnel_targets()
        if not ssh_host or ssh_host in ("0.0.0.0",):
            QMessageBox.warning(None, "错误", "请填写 SSH 主机地址！")
            return False
        if not ssh_port.isdigit() or not agent_port.isdigit() or not local_port.isdigit():
            QMessageBox.warning(None, "错误", "SSH/本地/Agent 端口必须是数字！")
            return False

        mode = self._remote_mode()
        ssh_bin = self._find_ssh_executable()
        # 模式 A：本地 -L 把本机 local_port 转到服务器 127.0.0.1:agent_port
        # 模式 B：反向 -R 把服务器 agent_port 转到本机 127.0.0.1:local_port（本机监听）
        if mode == "listen":
            forward_flag = "-R"
            forward_spec = f"{agent_port}:127.0.0.1:{local_port}"
            tip = f"反向隧道 -R {forward_spec}（服务器 dial ws://127.0.0.1:{agent_port}）"
        else:
            forward_flag = "-L"
            forward_spec = f"{local_port}:127.0.0.1:{agent_port}"
            tip = f"本地隧道 -L {forward_spec}（本机连 ws://127.0.0.1:{local_port}）"

        cmd = [
            ssh_bin,
            "-N",
            forward_flag, forward_spec,
            "-p", ssh_port,
            "-o", "ExitOnForwardFailure=yes",
            "-o", "ServerAliveInterval=30",
            "-o", "ServerAliveCountMax=3",
            "-o", "StrictHostKeyChecking=accept-new",
        ]
        key = self.remote_ssh_key.text().strip()
        if key:
            cmd.extend(["-i", key])
        cmd.append(f"{ssh_user}@{ssh_host}")

        self._append_remote_log(f"[SSH] 启动隧道: {tip}")
        self._append_remote_log(f"[SSH] {' '.join(cmd)}")
        try:
            creationflags = 0
            if sys.platform == "win32":
                creationflags = subprocess.CREATE_NO_WINDOW
            self._remote_ssh_proc = subprocess.Popen(
                cmd,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                creationflags=creationflags,
            )
            # 等转发就绪
            time.sleep(1.2)
            if self._remote_ssh_proc.poll() is not None:
                err = b""
                try:
                    err = self._remote_ssh_proc.stderr.read() or b""
                except Exception:
                    pass
                msg = err.decode("utf-8", errors="ignore").strip() or "SSH 进程已退出"
                self._append_remote_log(f"[ERROR] SSH 隧道失败: {msg}")
                self._remote_ssh_proc = None
                QMessageBox.warning(
                    None, "SSH 隧道失败",
                    f"{msg}\n\n请确认已配置免密登录或填写私钥；并已安装 OpenSSH 客户端。"
                )
                return False
            # 探测本地端口是否在听（模式 A）
            if mode != "listen":
                try:
                    sock = socket.create_connection(("127.0.0.1", int(local_port)), timeout=2)
                    sock.close()
                except Exception as e:
                    self._append_remote_log(f"[WARN] 本地转发端口暂不可达: {e}（若刚启动可稍后重试）")
            self._append_remote_log("[OK] SSH 隧道已建立")
            return True
        except FileNotFoundError:
            self._append_remote_log("[ERROR] 未找到 ssh 命令，请安装 OpenSSH")
            QMessageBox.warning(None, "错误", "未找到 ssh 命令，请安装 OpenSSH 客户端。")
            return False
        except Exception as e:
            self._append_remote_log(f"[ERROR] SSH 隧道异常: {e}")
            QMessageBox.warning(None, "错误", str(e))
            return False

    def _stop_ssh_tunnel(self):
        proc = getattr(self, "_remote_ssh_proc", None)
        if not proc:
            return
        try:
            if proc.poll() is None:
                proc.terminate()
                try:
                    proc.wait(timeout=2)
                except Exception:
                    proc.kill()
            self._append_remote_log("[SSH] 隧道已关闭")
        except Exception as e:
            self._append_remote_log(f"[WARN] 关闭 SSH 隧道: {e}")
        finally:
            self._remote_ssh_proc = None

    def _on_remote_mode_changed(self, *_):
        mode = self._remote_mode()
        if mode == "listen":
            self.btn_test.setText("检查监听端口")
            self.btn_start.setText("开始本机监听并分析")
            self.addr_input.setPlaceholderText("本机监听地址，一般 0.0.0.0")
            if self.addr_input.text().strip() in ("127.0.0.1", "localhost"):
                self.addr_input.setText("0.0.0.0")
            self._append_remote_log("模式：服务器 Agent dial → 本机监听（适合无公网/NAT）")
        else:
            self.btn_test.setText("WebSocket 连接测试")
            self.btn_start.setText("开始动态日志分析")
            self.addr_input.setPlaceholderText("服务器地址")
            if self.addr_input.text().strip() in ("0.0.0.0",):
                self.addr_input.setText("127.0.0.1")
            self._append_remote_log("模式：本机连接服务器 Agent serve")

    def _regen_remote_token(self):
        self.remote_token_input.setText(secrets.token_urlsafe(12))

    def _agent_script_path(self) -> str:
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "tools", "remote_log_agent.py")

    def _agent_history_cli_flags(self) -> str:
        """模式 B dial：历史范围写在服务器命令行；模式 A 由本机 OPTS 控制。"""
        if self._remote_history_mode() == "all":
            return " --from-start"
        return ""  # 默认仅实时新行

    def _build_agent_commands(self) -> tuple:
        """返回 (说明文本, 服务器执行命令, agent脚本源码路径)"""
        token = self.remote_token_input.text().strip()
        port = self.port_input.text().strip() or "8765"
        log_path = self.remote_log_path_input.text().strip() or "/var/log/nginx/access.log"
        host = self.addr_input.text().strip() or "127.0.0.1"
        mode = self._remote_mode()
        agent_path = self._agent_script_path()
        use_ssh = self._ssh_tunnel_enabled()
        _, ssh_host, ssh_port, agent_port, local_port = self._ssh_tunnel_targets() if use_ssh else (
            "", host, "22", port, port
        )
        ssh_user = self.remote_ssh_user.text().strip() if use_ssh else ""
        hist_mode = self._remote_history_mode()
        hist_desc = "全部历史+实时" if hist_mode == "all" else "仅连接后的新访问（推荐）"
        hist_flags = self._agent_history_cli_flags()

        if mode == "listen":
            if use_ssh:
                cmd = (
                    f"python3 remote_log_agent.py dial "
                    f"--url ws://127.0.0.1:{agent_port} "
                    f"--token {token} "
                    f"--file {log_path}{hist_flags}"
                )
                tip = (
                    "【模式 B + SSH】本机勾选 SSH 隧道并「开始本机监听」后，程序会自动建立：\n"
                    f"  ssh -N -R {agent_port}:127.0.0.1:{local_port} -p {ssh_port} {ssh_user}@{ssh_host}\n"
                    f"服务器上 Agent 只需 dial 本机回环（经反向隧道到你的 PC）。\n"
                    f"本机监听建议 127.0.0.1:{local_port}。\n"
                    f"历史范围（当前 GUI 选择）：{hist_desc} —— 已写入下方 dial 参数。\n"
                )
            else:
                dial_host = "YOUR_PC_IP" if host in ("0.0.0.0", "127.0.0.1") else host
                cmd = (
                    f"python3 remote_log_agent.py dial "
                    f"--url ws://{dial_host}:{port} "
                    f"--token {token} "
                    f"--file {log_path}{hist_flags}"
                )
                tip = (
                    "【模式 B】本机先点「开始本机监听并分析」，再在服务器执行下方命令。\n"
                    f"请把 YOUR_PC_IP 换成你电脑的公网/VPN IP（当前监听 {host}:{port}）。\n"
                    f"历史范围（当前 GUI 选择）：{hist_desc} —— 已写入下方 dial 参数。\n"
                    "明文 ws:// 仅建议内网/VPN；公网请勾选「使用 SSH 隧道」。\n"
                )
        else:
            bind_host = "127.0.0.1" if use_ssh else "0.0.0.0"
            # 默认不带 --replay：只推连接后新行；历史由本机连接时发 OPTS 控制
            cmd = (
                f"python3 remote_log_agent.py serve "
                f"--host {bind_host} --port {port} "
                f"--token {token} "
                f"--file {log_path}"
            )
            if use_ssh:
                tip = (
                    "【模式 A + SSH】服务器 Agent 只监听 127.0.0.1（不暴露公网）。\n"
                    f"本机勾选 SSH 隧道后会自动：ssh -N -L {local_port}:127.0.0.1:{agent_port} "
                    f"-p {ssh_port} {ssh_user}@{ssh_host}\n"
                    f"然后本机连接 ws://127.0.0.1:{local_port}（流量经 SSH 加密）。\n"
                    f"历史范围由本机启动分析时下发（当前选择：{hist_desc}）。\n"
                    "请使用最新 remote_log_agent.py（支持 OPTS mode=live|replay|all）。\n"
                )
            else:
                tip = (
                    "【模式 A】先在服务器执行下方命令，再在本机填写服务器 IP/端口并分析。\n"
                    f"本机将连接 ws://{host}:{port}\n"
                    f"历史范围由本机启动分析时下发（当前选择：{hist_desc}）。\n"
                    "默认只收连接后的新访问；需要旧日志时在 GUI 选「全部历史 + 实时」。\n"
                    "请重新上传最新 remote_log_agent.py 到服务器。\n"
                    "明文 ws:// 仅建议内网/VPN；公网请勾选「使用 SSH 隧道」。\n"
                )
        full = (
            f"{tip}\n"
            f"# 1) 上传最新 remote_log_agent.py 到服务器（覆盖旧版）\n"
            f"# 2) pip install websockets -q\n"
            f"# 3) 执行:\n"
            f"{cmd}\n"
        )
        return full, cmd, agent_path

    def _generate_remote_agent(self):
        """一键生成/导出 Agent 脚本与命令"""
        text, cmd, agent_path = self._build_agent_commands()
        dlg = QDialog(self)
        dlg.setWindowTitle("一键生成远程日志 Agent")
        dlg.resize(720, 480)
        v = QVBoxLayout(dlg)
        v.addWidget(QLabel("复制以下命令到服务器运行；也可导出脚本文件："))
        edit = QPlainTextEdit()
        edit.setReadOnly(True)
        edit.setPlainText(text)
        v.addWidget(edit)
        row = QHBoxLayout()
        btn_copy = QPushButton("复制命令")
        btn_save = QPushButton("导出 Agent 脚本 (.py)")
        btn_close = QPushButton("关闭")
        row.addWidget(btn_copy)
        row.addWidget(btn_save)
        row.addStretch(1)
        row.addWidget(btn_close)
        v.addLayout(row)

        def do_copy():
            QApplication.clipboard().setText(cmd)
            self._append_remote_log("[OK] 已复制 Agent 启动命令到剪贴板")

        def do_save():
            default_name = "remote_log_agent.py"
            path, _ = QFileDialog.getSaveFileName(
                dlg, "导出 Agent 脚本", default_name, "Python (*.py);;All (*)"
            )
            if not path:
                return
            try:
                src = agent_path
                if os.path.isfile(src):
                    with open(src, "r", encoding="utf-8") as f:
                        content = f.read()
                else:
                    QMessageBox.warning(dlg, "警告", f"找不到模板脚本:\n{src}")
                    return
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
                self._append_remote_log(f"[OK] 已导出 Agent: {path}")
                QMessageBox.information(
                    dlg, "完成",
                    f"已导出到:\n{path}\n\n在服务器执行:\n{cmd}"
                )
            except Exception as e:
                QMessageBox.critical(dlg, "错误", str(e))

        btn_copy.clicked.connect(do_copy)
        btn_save.clicked.connect(do_save)
        btn_close.clicked.connect(dlg.accept)
        dlg.exec()
        self._append_remote_log("[INFO] " + text.split("\n")[0])

    def Analyze_logs_remotely(self):
        """开始远程日志分析（模式 A 连接 / 模式 B 本机监听）"""
        if self.remote_thread and self.remote_thread.isRunning():
            QMessageBox.information(None, "提示", "远程日志分析已在运行中。")
            return

        mode = self._remote_mode()
        token = self.remote_token_input.text().strip()
        host = self.addr_input.text().strip()
        port = self.port_input.text().strip()
        if not port.isdigit():
            QMessageBox.warning(None, "错误", "请输入有效端口号！")
            return

        if mode == "client" and not self.remote_connected and not self._ssh_tunnel_enabled():
            QMessageBox.warning(None, "提示", "请先通过「WebSocket 连接测试」确认 Agent 可达！")
            return

        # 模式 A：可先建 SSH 本地转发再连 WS
        if mode == "client" and self._ssh_tunnel_enabled():
            if not self._start_ssh_tunnel():
                return

        optional_parameters_log = {
            "URI_Security_Check": self._remote_malware_enabled(),
            "Skip_IP_Geo": not self._remote_malware_enabled(),  # 开恶意分析时顺便做 IP 归属
        }
        self.url_stats = INITIALIZE_DEFAULTDICT()
        self._remote_streaming = True
        self._remote_ui_light_ticks = 0
        self._remote_log_append_n = 0
        self.gui_update_timer.setInterval(3000)
        self.stream_analyzer = remoteDynamicsLogProcessingThread(
            self.url_stats, optional_parameters_log=optional_parameters_log
        )
        self.stream_analyzer.stats_tick.connect(self._on_remote_stats_tick)
        self.stream_analyzer.start()

        opts_line = self._remote_opts_line()
        hist_mode = self._remote_history_mode()
        self._append_remote_log(
            f"[INFO] 范围={'全部历史+实时' if hist_mode == 'all' else '仅实时'}；"
            f"恶意分析={'开' if self._remote_malware_enabled() else '关'}"
        )

        if mode == "listen":
            # SSH 反向隧道时建议只听 127.0.0.1；先监听再开 -R
            if self._ssh_tunnel_enabled():
                _, _, _, _, local_port = self._ssh_tunnel_targets()
                listen_host = "127.0.0.1"
                listen_port = int(local_port)
            else:
                listen_host = host or "0.0.0.0"
                listen_port = int(port)
            self.remote_thread = RemoteLogListener(
                listen_host, listen_port, token,
                line_sink=self.stream_analyzer.enqueue_batch,
            )
            self.remote_connected = True
            self.remote_monitoring_ui_tab.connection_label.setText("监听中 🎧")
            self.remote_monitoring_ui_tab.connection_label.setStyleSheet(
                "color: #0a7; font-weight: bold; font-size: 14px;")
            # 不再 700ms 闪烁，省一点 UI 开销
            self.remote_thread.log_output.connect(self._log_output_remote_monitoring)
            self.remote_thread.connection_status.connect(self._update_connection_status)
            self.remote_thread.start()
            self.remote_log_thread = self.remote_thread
            time.sleep(0.4)  # 等监听起来
            if self._ssh_tunnel_enabled():
                if not self._start_ssh_tunnel():
                    self._disconnect_remote()
                    return
            # 模式 B：历史范围写在服务器 dial 命令里
            self._append_remote_log(
                f"[INFO] 本机监听 {listen_host}:{listen_port}；"
                + ("SSH 反向隧道已启用，服务器 dial 127.0.0.1" if self._ssh_tunnel_enabled()
                   else "请在服务器执行 dial 命令（可选 --from-start 拉全部历史）")
            )
            return

        if self._ssh_tunnel_enabled():
            _, _, _, _, local_port = self._ssh_tunnel_targets()
            address = f"ws://127.0.0.1:{local_port}"
        else:
            address = f"ws://{host}:{port}"
        self.remote_thread = RemoteLogReceiver(
            address, token=token,
            line_sink=self.stream_analyzer.enqueue_batch,
            opts_line=opts_line,
        )

        self.remote_thread.log_output.connect(self._log_output_remote_monitoring)
        self.remote_thread.connection_status.connect(self._update_connection_status)

        self.remote_thread.start()
        self.remote_log_thread = self.remote_thread
        self._append_remote_log(
            f"[INFO] 已开始拉取；{opts_line}（默认仅连接后的新访问）"
        )

    def _disconnect_remote(self):
        """断开远程连接 / 停止监听 / 关闭 SSH 隧道"""
        try:
            if self.remote_thread and hasattr(self.remote_thread, "stop"):
                self.remote_thread.stop()
            if self.remote_thread and self.remote_thread.isRunning():
                self.remote_thread.wait(1500)
        except Exception as e:
            self._append_remote_log(f"[WARN] 断开时异常: {e}")
        self.remote_thread = None
        self.remote_log_thread = None

        # 流式分析线程也要停，否则退出时会出现 QThread Destroyed while still running
        try:
            sa = getattr(self, "stream_analyzer", None)
            if sa is not None:
                if hasattr(sa, "stop"):
                    sa.stop()
                if sa.isRunning():
                    sa.quit()
                    if not sa.wait(2500):
                        sa.terminate()
                        sa.wait(500)
        except Exception:
            pass
        self.stream_analyzer = None
        self._remote_streaming = False
        self.gui_update_timer.setInterval(2000)

        self._stop_ssh_tunnel()
        self.remote_connected = False
        self._update_connection_status(False)
        self._append_remote_log("[INFO] 已断开远程连接")

    def _log_output_remote_monitoring(self, log: str):
        if not hasattr(self, "log_view") or self.log_view is None:
            return
        ts = datetime.datetime.now().strftime("[%H:%M:%S]")
        self._remote_log_append_n = getattr(self, "_remote_log_append_n", 0) + 1
        # 很少做全文截断（setPlainText 很贵）
        if self._remote_log_append_n % 80 == 0:
            try:
                plain = self.log_view.toPlainText()
                if plain.count("\n") > 300:
                    self.log_view.setPlainText("\n".join(plain.splitlines()[-150:]))
            except Exception:
                pass
        self.log_view.append(f"{ts} {log}")

    def _update_connection_status(self, status: bool):
        self.remote_connected = status
        if status:
            self.remote_monitoring_ui_tab.connection_label.setText("已连接 ✅")
            self.remote_monitoring_ui_tab.connection_label.setStyleSheet(
                "color: green; font-weight: bold; font-size: 14px;")
            # 远程流式时不闪烁，减少无意义重绘
            if not getattr(self, "_remote_streaming", False):
                self.remote_timer.start(700)
            else:
                self.remote_timer.stop()
        else:
            self.remote_monitoring_ui_tab.connection_label.setText("已断开 ❌")
            self.remote_monitoring_ui_tab.connection_label.setStyleSheet(
                "color: red; font-weight: bold; font-size: 14px;")
            self.remote_timer.stop()

    def _on_remote_stats_tick(self, snap: dict):
        """后台解析线程发来的轻量快照：只更新仪表盘 + 偶尔刷日志表。"""
        if not isinstance(snap, dict):
            return
        g = snap.get("global") or {}
        try:
            if isinstance(self.url_stats, dict):
                self.url_stats["_global_stats"] = g
        except Exception:
            pass
        if hasattr(self, "stats_cards"):
            self.stats_cards["total"]["value"].setText(str(g.get("request_total", 0)))
            self.stats_cards["unique_url"]["value"].setText(str(g.get("total_unique_uris", 0)))
            self.stats_cards["source_ip"]["value"].setText(str(g.get("total_unique_ips", 0)))
            self.stats_cards["status_code"]["value"].setText(str(g.get("total_unique_status_code", 0)))
            self.stats_cards["danger"]["value"].setText(str(g.get("danger_total", 0)))

        lines = snap.get("lines", 0)
        qsize = snap.get("qsize", 0)
        self.status_label.setText(f"远程监控：已解析 {lines} 行，队列 {qsize}")

        self._remote_ui_light_ticks = getattr(self, "_remote_ui_light_ticks", 0) + 1
        # 约每 3 次快照（~7.5s）刷一次日志表，避免频繁重建 QTable
        if self._remote_ui_light_ticks % 3 == 0:
            self._fill_remote_log_table_light(snap.get("top_rows") or [])

    def _fill_remote_log_table_light(self, top_rows):
        if not hasattr(self, "log_table") or self.log_table is None:
            return
        try:
            self.log_table.setSortingEnabled(False)
            self.log_table.setUpdatesEnabled(False)
            self.log_table.setRowCount(len(top_rows))
            for row, item in enumerate(top_rows):
                self.log_table.setItem(row, 0, QTableWidgetItem(unquote(str(item.get("url", "")))))
                self.log_table.setItem(row, 1, QTableWidgetItem(str(item.get("count", 0))))
                self.log_table.setItem(row, 2, QTableWidgetItem(",".join(map(str, item.get("status") or []))[:80]))
                self.log_table.setItem(row, 3, QTableWidgetItem(",".join(map(str, item.get("ips") or []))[:120]))
                self.log_table.setItem(row, 4, QTableWidgetItem(",".join(map(str, item.get("methods") or []))[:40]))
                # 其余列置空，避免复杂聚合
                for c in range(5, self.log_table.columnCount()):
                    if self.log_table.item(row, c) is None:
                        self.log_table.setItem(row, c, QTableWidgetItem(""))
        finally:
            self.log_table.setUpdatesEnabled(True)
            self.log_table.setSortingEnabled(True)

    def _append_remote_log(self, msg: str):
        if not hasattr(self, "log_view") or self.log_view is None:
            return
        ts = datetime.datetime.now().strftime("[%H:%M:%S]")
        self.log_view.append(f"{ts} {msg}")
        self.log_view.verticalScrollBar().setValue(
            self.log_view.verticalScrollBar().maximum()
        )

    def _test_remote_connection(self):
        """WebSocket + Token 探测（模式 A）；模式 B 检查端口；可选先建 SSH 隧道。"""
        host = self.addr_input.text().strip()
        port = self.port_input.text().strip()
        token = self.remote_token_input.text().strip()

        if not port.isdigit():
            QMessageBox.warning(None, "错误", "请输入有效的端口号！")
            return

        mode = self._remote_mode()

        if self._ssh_tunnel_enabled():
            if not self._start_ssh_tunnel():
                self.remote_connected = False
                self._update_connection_status(False)
                return

        if mode == "listen":
            try:
                if self._ssh_tunnel_enabled():
                    _, _, _, _, local_port = self._ssh_tunnel_targets()
                    bind_host, bind_port = "127.0.0.1", int(local_port)
                else:
                    if not host:
                        QMessageBox.warning(None, "错误", "请输入有效的地址！")
                        return
                    bind_host = "0.0.0.0" if host == "0.0.0.0" else host
                    bind_port = int(port)
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
                s.bind((bind_host, bind_port))
                s.close()
                self.remote_connected = True
                self._update_connection_status(True)
                self._append_remote_log(f"[OK] 端口 {bind_host}:{bind_port} 可监听")
            except Exception as e:
                self.remote_connected = False
                self._update_connection_status(False)
                self._append_remote_log(f"[ERROR] 端口不可用: {e}")
            return

        # 模式 A：连本地转发端口或直连服务器
        if self._ssh_tunnel_enabled():
            _, _, _, _, local_port = self._ssh_tunnel_targets()
            ws_host, ws_port = "127.0.0.1", local_port
        else:
            if not host:
                QMessageBox.warning(None, "错误", "请输入有效的服务器地址！")
                return
            ws_host, ws_port = host, port

        url = f"ws://{ws_host}:{ws_port}"
        if token:
            url = f"{url}?token={quote(token, safe='')}"

        async def _probe():
            async with websockets.connect(url, open_timeout=5, max_size=2 ** 20) as ws:
                if token:
                    await ws.send(f"AUTH {token}")
                try:
                    await asyncio.wait_for(ws.recv(), timeout=1.5)
                except asyncio.TimeoutError:
                    pass
                return True

        try:
            asyncio.run(_probe())
            self.remote_connected = True
            self._update_connection_status(True)
            via = "经 SSH 隧道" if self._ssh_tunnel_enabled() else "直连"
            self._append_remote_log(f"[OK] WebSocket 鉴权连通 {ws_host}:{ws_port}（{via}）")
            self._append_remote_log("[INFO] 测试只验证连通；将自动开始拉取日志并分析…")
            # 测试通过后自动开始收流（仅测通不会显示访问）
            QTimer.singleShot(300, self.Analyze_logs_remotely)
        except Exception as e:
            self.remote_connected = False
            self._update_connection_status(False)
            self._append_remote_log(f"[ERROR] WebSocket 测试失败 {ws_host}:{ws_port} - {e}")

    def _blink_status(self):
        """状态闪烁动画"""
        if not self.remote_connected:
            return
        self._status_toggle = not self._status_toggle
        color = "green" if self._status_toggle else "lime"
        self.remote_monitoring_ui_tab.connection_label.setStyleSheet(
            f"color: {color}; font-weight: bold; font-size: 14px;"
        )

    ############################ 下面是正则规则模块 #####################
    def regular_expression_matching(self):
        """正则表达式匹配规则管理界面 - 使用选项卡布局"""

        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/re_matching.png"), "规则配置")
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(0, 0, 0, 0)  # 由 QSS 控制边距

        # 创建选项卡控件
        tab_widget = QTabWidget()
        self.re_matching_tabs = tab_widget
        # 注意：QSS 已经移动到父控件 "tab" 上，这里不需要再设置

        # =================== 选项卡0: 添加安全规则 ===================
        add_tab = QWidget()
        add_tab_main_layout = QVBoxLayout(add_tab)
        add_tab_main_layout.setContentsMargins(15, 15, 15, 15)
        add_tab_main_layout.setSpacing(15)

        # ===== 主布局：左右结构 =====
        main_layout = QHBoxLayout()
        main_layout.setSpacing(15)

        # ===== 左侧：规则配置 (使用 QFormLayout) =====
        left_group = QGroupBox(self.lang_manager.tr("Rule_Config", "规则配置"))
        left_form_layout = QFormLayout(left_group)
        left_form_layout.setSpacing(10)
        left_form_layout.setLabelAlignment(Qt.AlignmentFlag.AlignRight)  # 标签右对齐

        # --- 修改：添加 self. 并设置默认值 ---
        self.add_risk_label_edit = QLineEdit()
        self.add_risk_label_edit.setPlaceholderText("例如：safety_testing")
        self.add_risk_label_edit.setText("safety_testing")  # 默认值
        left_form_layout.addRow(self.lang_manager.tr("Risk_Label", "风险标识名："), self.add_risk_label_edit)

        # --- 修改：添加 self. ---
        self.add_rule_name_edit = QLineEdit()
        self.add_rule_name_edit.setPlaceholderText("例如：My New SQLi Rule")  # 修改了提示
        left_form_layout.addRow(self.lang_manager.tr("Rule_Name", "规则名称："), self.add_rule_name_edit)

        # --- 修改：添加 self. ---
        self.add_detection_location_edit = QLineEdit()
        self.add_detection_location_edit.setPlaceholderText(
            "URI|forms_key_body|multipart_file_name_body|ALL_headers|xml_value_body|!headers:referer")
        left_form_layout.addRow(self.lang_manager.tr("Detection_Location", "检测目标字段："),
                                self.add_detection_location_edit)

        # --- 修改：添加 self. ---
        self.add_severity_combo = QComboBox()
        self.add_severity_combo.addItems([
            self.lang_manager.tr("Critical_Risk", "严重"),
            self.lang_manager.tr("High_Risk", "高危"),
            self.lang_manager.tr("Medium_Risk", "中危"),
            self.lang_manager.tr("Low_Risk", "低危"),
            self.lang_manager.tr("Info_Risk", "信息"),
        ])
        self.add_severity_combo.setCurrentText(self.lang_manager.tr("Medium_Risk", "中危"))
        left_form_layout.addRow(self.lang_manager.tr("Severity", "严重等级："), self.add_severity_combo)

        # 正则表达式 rules (单独添加，使其可以设置高度)
        left_form_layout.addRow(QLabel(self.lang_manager.tr("Rules", "正则表达式：")))
        # --- 修改：添加 self. ---
        self.add_rules_edit = QTextEdit()
        self.add_rules_edit.setAcceptRichText(False)
        self.add_rules_edit.setPlaceholderText("每条正则表达式占一行")
        self.add_rules_edit.setMinimumHeight(150)  # 给正则输入框一个最小高度
        left_form_layout.addRow(self.add_rules_edit)

        left_group.setMinimumWidth(350)
        left_group.setMaximumWidth(500)  # 增加一个最大宽度

        # ===== 右侧：数据识别测试 =====
        right_group = QGroupBox(self.lang_manager.tr("Data_Test", "数据识别测试"))
        right_layout = QVBoxLayout(right_group)
        right_layout.setSpacing(10)

        # --- 输入栏（HTTP 数据） ---
        right_layout.addWidget(QLabel(self.lang_manager.tr("Input_Field", "输入栏（HTTP数据）:")))
        # --- 修改：添加 self. ---
        self.add_input_text = QTextEdit()
        self.add_input_text.setAcceptRichText(False)
        self.add_input_text.setPlaceholderText(self.lang_manager.tr("Input_hint", "请输入要测试的HTTP数据..."))
        right_layout.addWidget(self.add_input_text, 3)  # 输入区域占3份高度

        # --- 按钮区（测试规则 / 添加规则） ---
        button_layout_tab0 = QHBoxLayout()
        button_layout_tab0.addStretch(1)
        btn_test = QPushButton(QIcon("ico/refresh.png"), self.lang_manager.tr("Test_Rule", "测试规则"))

        # --- 新增：连接信号 ---
        btn_test.clicked.connect(self.test_new_security_rule)

        btn_add = QPushButton(QIcon("ico/add.png"), self.lang_manager.tr("Add_Rule", "添加规则"))
        btn_add.setObjectName("PrimaryButton")

        # --- 新增：连接信号 ---
        btn_add.clicked.connect(self.add_new_security_rule)

        button_layout_tab0.addWidget(btn_test)
        button_layout_tab0.addWidget(btn_add)
        right_layout.addLayout(button_layout_tab0)

        # --- 测试结果区 ---
        right_layout.addWidget(QLabel(self.lang_manager.tr("Test_Result", "测试结果:")))
        # --- 修改：添加 self. ---
        self.add_result_text = QTextEdit()
        self.add_result_text.setPlaceholderText(self.lang_manager.tr("Result_hint", "这里显示规则测试的结果输出..."))
        self.add_result_text.setReadOnly(True)
        right_layout.addWidget(self.add_result_text, 2)  # 结果区域占2份高度

        # ===== 组装主布局 =====
        main_layout.addWidget(left_group, 1)  # 左侧占1份
        main_layout.addWidget(right_group, 2)  # 右侧占2份
        add_tab_main_layout.addLayout(main_layout)

        tab_widget.addTab(add_tab, self.lang_manager.tr("Add_security_rules", "添加安全检测规则"))

        # =================== 选项卡1: 编辑正则库 ===================
        config_tab = QWidget()
        config_layout = QVBoxLayout(config_tab)
        config_layout.setContentsMargins(15, 15, 15, 15)
        config_layout.setSpacing(10)

        # 类别选择
        category_layout = QHBoxLayout()
        category_layout.addWidget(QLabel(self.lang_manager.tr("category", "类别:")))
        self.regex_category = QComboBox()
        self.regex_category.addItems([
            self.lang_manager.tr("security_detection", "安全检测"),
            self.lang_manager.tr("log_format_recognition", "日志格式识别"),
            self.lang_manager.tr("log_parsing", "日志解析")
        ])
        self.regex_category.currentIndexChanged.connect(self.load_regex_config)
        category_layout.addWidget(self.regex_category, 1)

        category_layout.addWidget(QLabel(self.lang_manager.tr("select_rule", "选择规则:")))
        self.rule_name_combo = QComboBox()
        self.rule_name_combo.currentIndexChanged.connect(self.load_selected_rule)
        category_layout.addWidget(self.rule_name_combo, 2)  # 让规则名选择框更宽
        config_layout.addLayout(category_layout)

        # 规则名称编辑
        rule_name_layout = QHBoxLayout()
        rule_name_layout.addWidget(QLabel(self.lang_manager.tr("rule_name", "规则名称:")))
        self.rule_name_edit = QLineEdit()
        rule_name_layout.addWidget(self.rule_name_edit)
        config_layout.addLayout(rule_name_layout)

        # 正则表达式编辑区域
        config_layout.addWidget(QLabel(self.lang_manager.tr("regex_pattern", "正则表达式模式:")))
        self.regex_edit = QTextEdit()
        self.regex_edit.setPlaceholderText(self.lang_manager.tr("regex_placeholder", "在此输入正则表达式..."))
        self.regex_edit.setAcceptRichText(False)
        self.regex_edit.setLineWrapMode(QTextEdit.LineWrapMode.NoWrap)
        self.regex_edit.setMinimumHeight(150)  # 设置最小高度

        # 添加语法高亮 (确保 RegexHighlighter 类已定义)
        try:
            self.highlighter = RegexHighlighter(self.regex_edit.document())
        except NameError:
            print("警告: RegexHighlighter 类未定义，语法高亮将不可用。")
            self.highlighter = None  # 避免崩溃

        config_layout.addWidget(self.regex_edit)

        # 测试区域
        test_group = QGroupBox(self.lang_manager.tr("test_regex", "测试正则表达式"))
        test_layout = QGridLayout(test_group)  # 使用 QGridLayout 更好地对齐
        test_layout.setSpacing(10)

        self.test_input = QTextEdit()
        self.test_input.setPlaceholderText(self.lang_manager.tr("test_input_placeholder", "输入测试文本..."))
        self.test_input.setMaximumHeight(100)  # 限制测试输入框高度

        self.test_output = QTextEdit()
        self.test_output.setReadOnly(True)
        self.test_output.setMaximumHeight(100)  # 限制测试输出框高度

        test_btn = QPushButton(QIcon("ico/refresh.png"), self.lang_manager.tr("test_regex_btn", "测试"))  # 使用刷新图标
        test_btn.clicked.connect(self.test_regex)
        test_btn.setFixedWidth(120)  # 固定测试按钮宽度

        test_layout.addWidget(QLabel("测试输入:"), 0, 0)
        test_layout.addWidget(self.test_input, 0, 1)
        test_layout.addWidget(test_btn, 0, 2)
        test_layout.addWidget(QLabel("匹配结果:"), 1, 0)
        test_layout.addWidget(self.test_output, 1, 1, 1, 2)  # 结果框跨2列

        config_layout.addWidget(test_group)

        # 操作按钮
        button_layout_tab1 = QHBoxLayout()
        button_layout_tab1.setSpacing(10)

        self.save_regex_btn = QPushButton(QIcon("ico/save.png"), self.lang_manager.tr("save_rule", "保存规则"))
        self.save_regex_btn.setObjectName("PrimaryButton")  # 设为主要按钮
        self.save_regex_btn.clicked.connect(self.save_regex_config)

        self.add_regex_btn = QPushButton(QIcon("ico/add.png"), self.lang_manager.tr("add_rule", "添加规则"))
        self.add_regex_btn.clicked.connect(self.add_new_rule)

        self.delete_regex_btn = QPushButton(QIcon("ico/delete.png"), self.lang_manager.tr("delete_rule", "删除规则"))
        self.delete_regex_btn.clicked.connect(self.delete_rule)

        self.reset_regex_btn = QPushButton(QIcon("ico/reset.png"), self.lang_manager.tr("reset", "重置"))
        self.reset_regex_btn.clicked.connect(self.load_regex_config)

        button_layout_tab1.addStretch()
        button_layout_tab1.addWidget(self.reset_regex_btn)
        button_layout_tab1.addWidget(self.delete_regex_btn)
        button_layout_tab1.addWidget(self.add_regex_btn)
        button_layout_tab1.addWidget(self.save_regex_btn)  # 把主要按钮放右边

        config_layout.addLayout(button_layout_tab1)
        tab_widget.addTab(config_tab, self.lang_manager.tr("Edit_regular_library", "编辑正则表达式库"))

        # =================== 选项卡2: 正则表达式库 ===================
        library_tab = QWidget()
        library_layout = QVBoxLayout(library_tab)
        library_layout.setContentsMargins(15, 15, 15, 15)
        library_layout.setSpacing(10)

        # 搜索框
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel(self.lang_manager.tr("search", "搜索:")))

        self.regex_search_input = QLineEdit()
        self.regex_search_input.setPlaceholderText(self.lang_manager.tr("search_placeholder", "搜索规则名称或模式..."))
        self.regex_search_input.textChanged.connect(self.filter_regex_library)
        search_layout.addWidget(self.regex_search_input)

        # 刷新按钮
        refresh_btn = QPushButton(QIcon("ico/refresh.png"), self.lang_manager.tr("refresh", "刷新"))
        refresh_btn.clicked.connect(self.load_regex_library)
        search_layout.addWidget(refresh_btn)
        library_layout.addLayout(search_layout)

        # 规则库表格
        self.regex_table = QTableWidget()
        self.regex_table.setColumnCount(4)
        self.regex_table.setHorizontalHeaderLabels([
            self.lang_manager.tr("category", "类别"),
            self.lang_manager.tr("rule_name", "规则名称"),
            self.lang_manager.tr("pattern", "模式"),
            self.lang_manager.tr("description", "描述")
        ])
        self.regex_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)  # 整行选择
        self.regex_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)  # 禁止编辑
        self.regex_table.doubleClicked.connect(self.load_regex_from_table)
        self.regex_table.setAlternatingRowColors(True)  # 开启斑马纹

        # 设置列宽和调整模式
        header = self.regex_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)  # 类别 (可拖动)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Interactive)  # 规则名称 (可拖动)
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)  # 模式 (自动拉伸)
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)  # 描述 (自动拉伸)

        self.regex_table.setColumnWidth(0, 120)
        self.regex_table.setColumnWidth(1, 180)

        library_layout.addWidget(self.regex_table)

        # 库操作按钮
        library_btn_layout = QHBoxLayout()
        import_btn = QPushButton(QIcon("ico/import.png"), self.lang_manager.tr("import", "导入"))
        import_btn.clicked.connect(self.import_regex_library)

        export_btn = QPushButton(QIcon("ico/export.png"), self.lang_manager.tr("export", "导出"))
        export_btn.clicked.connect(self.export_regex_library)

        library_btn_layout.addStretch()
        library_btn_layout.addWidget(import_btn)
        library_btn_layout.addWidget(export_btn)

        library_layout.addLayout(library_btn_layout)
        tab_widget.addTab(library_tab, self.lang_manager.tr("regex_library", "正则表达式库"))

        # 将选项卡控件添加到主布局
        layout.addWidget(tab_widget)

        # 初始化加载配置 (假设这些方法存在)
        try:
            self.load_regex_config()
            self.load_regex_library()
        except Exception as e:
            print(f"初始化加载失败: {e}")

        # 加入到 tabs
        self.tabs["re_matching"] = tab

    def Construct_security_detection(self):
        """ 构建安全检测数据变成json """
        # 1. 读取UI输入
        category_key = self.add_risk_label_edit.text().strip()
        human_name = self.add_rule_name_edit.text().strip()
        location = self.add_detection_location_edit.text().strip()
        severity_level = self.add_severity_combo.currentText()
        regex_list = [r.strip() for r in self.add_rules_edit.toPlainText().split('\n') if r.strip()]

        # 2. 验证输入
        if not category_key:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                "请输入风险标识名 (例如: safety_testing)!")
            self.add_risk_label_edit.setFocus()
            return

        if not human_name:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                "请输入规则名称 (例如: My SQLi Rule)!")
            self.add_rule_name_edit.setFocus()
            return

        if not location:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "请输入检测目标字段!")
            self.add_detection_location_edit.setFocus()
            return

        if not regex_list:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "请输入至少一条正则表达式!")
            self.add_rules_edit.setFocus()
            return

        return {category_key: {"detection_location": [location], "name": [human_name], "rules": regex_list,
                               "severity": [severity_level]}}

    def test_new_security_rule(self):
        """测试 '添加安全规则' 选项卡中的正则表达式"""
        regex_text = self.add_rules_edit.toPlainText().strip()
        if not regex_text:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("please_enter_regex", "请输入正则表达式!"))
            return

        test_text = self.add_input_text.toPlainText()
        if not test_text:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("please_enter_test_text", "请输入测试文本!"))
            return

        try:

            safety_testing = self.Construct_security_detection()
            if not safety_testing:  # 如果输入的规则有错误退出
                return None

            try:
                scanner = SecurityScanner(JSON_rules=safety_testing)
            except TypeError:
                self.add_result_text.setPlainText(
                    "错误: SecurityScanner 不支持 JSON_rules 初始化。\n"
                    "请检查 SecurityScanner 的 __init__ 方法。"
                )
                return
            except Exception as e:
                self.add_result_text.setPlainText(f"初始化 SecurityScanner 失败: {e}")
                return

            # 判断输入的是全流量请求还是只有一个url
            request_type, wrong = classify_input(test_text)
            if request_type == None:
                QMessageBox.warning(self, self.lang_manager.tr("warning", "解析http错误"), wrong)
                return None

            # 判断检测正则是否是单纯URI
            if request_type == "URL":
                results = scanner.scan_url(test_text)
                results = scanner.pretty_print_results(results)
            else:  # 请求全流量
                optional_parameters = {
                    "URL_Security_Check": True,
                    "Request_Head_Security_Check": True,
                    "Data_section_detection": {
                        "enabled": True,  # ← 是否勾选“请求体安全检测”总开关
                        "binary": True,
                        "limit_size": 0.5,  # 设置最大检测请求体大小 最大1MB 默认是0.5MB
                        "forms": True,
                        "json": True,
                        "xml": True,
                        "multipart": True
                    }
                }
                method, path, headers, body, wrong = parse_http_request(test_text)
                if wrong != None:
                    QMessageBox.warning(self, self.lang_manager.tr("warning", "解析http错误"), wrong)
                    return None
                # 检测uri
                results = scanner.scan_url(path)
                results = "请求URI检测" + scanner.pretty_print_results(results)

                # 请求头检测
                header_results = scanner.scan_headers(headers)

                results += "\n请求头检测" + scanner.pretty_print_results(header_results)
                scan_results = []
                if not ("GET" in method or "HEAD" in method):
                    scan_results = scanner.scan_body(body, headers["content-type"],
                                                     optional_parameters=optional_parameters)

                results += "\n请求体检测" + scanner.pretty_print_results(scan_results)

            self.add_result_text.setPlainText(results)

        except Exception as e:
            self.add_result_text.setPlainText(f"{self.lang_manager.tr('test_error', '测试时发生错误')}: {str(e)}")

    def add_new_security_rule(self):
        """添加新的安全检测规则到 config.yaml"""
        try:
            # 1. 读取UI输入
            category_key = self.add_risk_label_edit.text().strip()
            human_name = self.add_rule_name_edit.text().strip()
            location = self.add_detection_location_edit.text().strip()
            severity_level = self.add_severity_combo.currentText()
            regex_list = [r.strip() for r in self.add_rules_edit.toPlainText().split('\n') if r.strip()]

            # 2. 验证输入
            if not category_key:
                QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                    "请输入风险标识名 (例如: safety_testing)!")
                self.add_risk_label_edit.setFocus()
                return

            if not human_name:
                QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                    "请输入规则名称 (例如: My SQLi Rule)!")
                self.add_rule_name_edit.setFocus()
                return

            if not location:
                QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "请输入检测目标字段!")
                self.add_detection_location_edit.setFocus()
                return

            if not regex_list:
                QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "请输入至少一条正则表达式!")
                self.add_rules_edit.setFocus()
                return

            # 3. 生成 Rule Group ID (简单的 "slugify")
            rule_group_id = re.sub(r'[^a-z0-9_]+', '_', human_name.lower().strip()).strip('_')
            if not rule_group_id:
                rule_group_id = f"rule_{int(time.time())}"  # Fallback

            # 4. 检查配置中是否存在
            if category_key not in self.config:
                self.config[category_key] = {}

            if rule_group_id in self.config[category_key]:
                reply = QMessageBox.question(
                    self,
                    self.lang_manager.tr("warning", "警告"),
                    f"规则ID '{rule_group_id}' (由 '{human_name}' 生成) 已存在。\n您想覆盖它吗？",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                if reply == QMessageBox.StandardButton.No:
                    return

            # 5. 构建新的规则字典
            # 规则列表中的每条规则都共享相同的名称、位置和严重性
            num_rules = len(regex_list)
            new_rule_data = {
                'name': [human_name] * num_rules,
                'detection_location': [location] * num_rules,
                'rules': regex_list,
                'severity': [severity_level] * num_rules
            }

            # 6. 更新 self.config 并保存到 config.yaml
            self.config[category_key][rule_group_id] = new_rule_data

            with open("config.yaml", "w", encoding="utf-8") as f:
                yaml.dump(self.config, f, allow_unicode=True, sort_keys=False)

            # 7. 成功反馈
            QMessageBox.information(
                self,
                self.lang_manager.tr("success", "成功"),
                f"规则 '{human_name}' (ID: {rule_group_id}) 已成功添加到 {category_key}!"
            )
            # 在测试结果框显示添加的内容
            self.add_result_text.setPlainText(
                f"规则 '{human_name}' 已添加。\n{yaml.dump({rule_group_id: new_rule_data}, allow_unicode=True, sort_keys=False)}")

            # 8. 刷新 "编辑" 和 "库" 选项卡
            self.load_regex_config()
            self.load_regex_library()
            self._reload_security_scanner()

            # 9. 清空输入框
            self.add_rule_name_edit.clear()
            self.add_detection_location_edit.clear()
            self.add_rules_edit.clear()


        except Exception as e:
            QMessageBox.critical(self, self.lang_manager.tr("error", "错误"), f"添加规则失败: {str(e)}")
            self.add_result_text.setPlainText(f"添加规则失败: {str(e)}")

    def export_regex_library(self):
        """导出正则表达式库"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            self.lang_manager.tr("export_regex_library", "导出正则表达式库"),
            "regex_library.yaml",
            "YAML Files (*.yaml *.yml);;All Files (*)"
        )

        if file_path:
            try:
                # 确保文件扩展名
                if not file_path.endswith(('.yaml', '.yml')):
                    file_path += '.yaml'

                # 创建导出配置
                export_config = {
                    "log_formats": self.config.get("log_formats", {}),
                    "parsers": self.config.get("parsers", {}),
                    "safety_testing": self.config.get("safety_testing", {})
                }

                with open(file_path, 'w', encoding='utf-8') as f:
                    yaml.dump(export_config, f, allow_unicode=True, sort_keys=False)

                QMessageBox.information(
                    self,
                    self.lang_manager.tr("success", "成功"),
                    self.lang_manager.tr("library_exported", "规则库导出成功!")
                )
            except Exception as e:
                QMessageBox.warning(
                    self,
                    self.lang_manager.tr("error", "错误"),
                    f"{self.lang_manager.tr('export_failed', '导出失败')}: {str(e)}"
                )

    def import_regex_library(self):
        """导入正则表达式库"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            self.lang_manager.tr("import_regex_library", "导入正则表达式库"),
            "",
            "YAML Files (*.yaml *.yml);;All Files (*)"
        )

        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    imported_config = yaml.safe_load(f)

                # 合并配置
                for key in ["log_formats", "parsers", "safety_testing"]:
                    if key in imported_config:
                        if key not in self.config:
                            self.config[key] = {}
                        self.config[key].update(imported_config[key])

                # 保存配置
                with open("config.yaml", "w", encoding="utf-8") as f:
                    yaml.dump(self.config, f, allow_unicode=True, sort_keys=False)

                # 刷新库
                self.load_regex_library()
                self.load_regex_config()
                self._reload_security_scanner()

                QMessageBox.information(
                    self,
                    self.lang_manager.tr("success", "成功"),
                    self.lang_manager.tr("library_imported", "规则库导入成功!")
                )
            except Exception as e:
                QMessageBox.warning(
                    self,
                    self.lang_manager.tr("error", "错误"),
                    f"{self.lang_manager.tr('import_failed', '导入失败')}: {str(e)}"
                )

    def save_regex_config(self, *_):
        """保存编辑正则库"""
        rule_name = self.rule_name_edit.text().strip()
        if not rule_name:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("please_enter_rule_name", "请输入规则名称!"))
            return

        regex_text = self.regex_edit.toPlainText().strip()
        if not regex_text:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("please_enter_regex", "请输入正则表达式!"))
            return

        try:
            category = self.regex_category.currentText()

            if category in ["安全检测", "Security Detection"]:
                # 获取原始规则ID和名称
                original_name = self.rule_name_combo.currentText()
                rule_id = self.display_name_to_id.get(original_name)

                if not rule_id:
                    QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                        self.lang_manager.tr("original_rule_not_found", "找不到原始规则!"))
                    return

                # 获取原始名称索引
                name_index = self.current_config[rule_id]["name"].index(original_name)

                # 更新规则内容
                self.current_config[rule_id]["rules"][name_index] = regex_text

                # 更新名称（如果改变了）
                if rule_name != original_name:
                    self.current_config[rule_id]["name"][name_index] = rule_name
                    # 更新名称映射
                    self.display_name_to_id[rule_name] = rule_id
                    if original_name in self.display_name_to_id:
                        del self.display_name_to_id[original_name]
            else:
                # 其他规则类型的处理保持不变
                rules = [r.strip() for r in regex_text.split("\n") if r.strip()]
                self.current_config[rule_name] = rules[0] if len(rules) == 1 else rules

            # 更新主配置并保存文件
            self.update_and_save_config(category)

            # 完全重新加载配置以确保UI同步
            self.load_regex_config()

            # 设置当前选中的规则
            self.rule_name_combo.setCurrentText(rule_name)

            QMessageBox.information(self, self.lang_manager.tr("done", "完成"),
                                    self.lang_manager.tr("config_saved", "配置已保存!"))
        except Exception as e:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                f"{self.lang_manager.tr('save_config_failed', '保存配置失败')}: {str(e)}")

    def _reload_security_scanner(self):
        """规则变更后热重载 Python 扫描器，避免需重启才生效。"""
        try:
            ok, err = core_processing.reload_scanner("config.yaml")
            if ok:
                self.status_label.setText("安全规则已热重载")
            else:
                self.status_label.setText(f"规则已保存，但热重载失败: {err}")
        except Exception as e:
            self.status_label.setText(f"规则已保存，热重载异常: {e}")

    def update_and_save_config(self, category):
        """更新主配置并保存到文件"""
        if category in ["日志格式识别", "Log Format Recognition"]:
            self.config["log_formats"] = self.current_config
        elif category in ["日志解析", "Log Parsing"]:
            self.config["parsers"] = self.current_config
        elif category in ["安全检测", "Security Detection"]:
            self.config["safety_testing"] = self.current_config

        # 保存到文件
        with open("config.yaml", "w", encoding="utf-8") as f:
            yaml.dump(self.config, f, allow_unicode=True, sort_keys=False)
        self._reload_security_scanner()

    # 添加新规则
    def add_new_rule(self, *_):
        """添加新规则"""
        rule_name, ok = QInputDialog.getText(self,
                                             self.lang_manager.tr("add_new_rule", "添加新规则"),
                                             self.lang_manager.tr("enter_rule_name", "请输入规则名称:"))
        if not ok or not rule_name.strip():
            return

        category = self.regex_category.currentText()
        # 支持中文和英文分类名判断
        is_security = category in ["安全检测", "Security Detection"]

        # 检查名称是否已存在
        if is_security:
            for rule_id, rule_data in self.current_config.items():
                if rule_name in rule_data.get("name", []):
                    QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                        self.lang_manager.tr("rule_exists", "规则名称已存在!"))
                    return
        else:
            if rule_name in self.current_config:
                QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                    self.lang_manager.tr("rule_exists", "规则名称已存在!"))
                return

        try:
            if is_security:
                rule_id = rule_name.lower().replace(" ", "_")
                self.current_config[rule_id] = {
                    "name": [rule_name],
                    "rules": [""],
                    "detection_location": ["URI"],
                    "severity": ["中危"]
                }
                self.display_name_to_id[rule_name] = rule_id
            else:
                self.current_config[rule_name] = ""

            self.load_regex_config()
            self.rule_name_combo.setCurrentText(rule_name)
        except Exception as e:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("add_rule_failed", "添加规则失败") + f": {str(e)}")

    def delete_rule(self, *_):
        """删除规则"""
        rule_name = self.rule_name_combo.currentText()
        if not rule_name:
            return

        reply = QMessageBox.question(
            self,
            self.lang_manager.tr("confirm_delete", "确认删除"),
            self.lang_manager.tr("confirm_delete_rule", f"确定要删除规则 '{rule_name}' 吗?"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            try:
                category = self.regex_category.currentText()
                # 支持中文和英文分类判断
                if category in ["安全检测", "Security Detection"]:
                    key_to_delete = None
                    for k, v in self.current_config.items():
                        if rule_name in v.get("name", []):
                            key_to_delete = k
                            break
                    if key_to_delete:
                        del self.current_config[key_to_delete]
                else:
                    if rule_name in self.current_config:
                        del self.current_config[rule_name]

                # 更新主配置，支持中英文分类名
                if category in ["日志格式识别", "Log Format Recognition"]:
                    self.config["log_formats"] = self.current_config
                elif category in ["日志解析", "Log Parsing"]:
                    self.config["parsers"] = self.current_config
                elif category in ["安全检测", "Security Detection"]:
                    self.config["safety_testing"] = self.current_config

                with open("config.yaml", "w", encoding="utf-8") as f:
                    yaml.dump(self.config, f, allow_unicode=True, sort_keys=False)

                self.load_regex_config()
                self.load_regex_library()
                self._reload_security_scanner()
                QMessageBox.information(self, self.lang_manager.tr("done", "完成"),
                                        self.lang_manager.tr("rule_deleted", "规则已删除!"))
            except Exception as e:
                QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                    self.lang_manager.tr("delete_rule_failed", "删除规则失败") + f": {str(e)}")

    def test_regex(self):
        """测试正则表达式"""
        regex_text = self.regex_edit.toPlainText().strip()
        if not regex_text:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("please_enter_regex", "请输入正则表达式!"))
            return

        test_text = self.test_input.toPlainText()
        if not test_text:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("please_enter_test_text", "请输入测试文本!"))
            return

        try:
            results = []
            lines = regex_text.split('\n')

            for i, line in enumerate(lines, 1):
                line = line.strip()
                if not line or line.startswith('#'):  # 跳过空行和注释
                    continue

                try:
                    # 与 SecurityScanner 一致：忽略大小写
                    pattern = re.compile(line, re.IGNORECASE)
                    matches = list(pattern.finditer(test_text))

                    if matches:
                        results.append(f"✅ 第{i}行: {line}")
                        results.append(f"   匹配到 {len(matches)} 处 (IGNORECASE):")
                        for j, match in enumerate(matches[:5], 1):
                            results.append(
                                f"     {j}. @{match.start()}-{match.end()} {repr(match.group(0)[:120])}"
                            )
                        if len(matches) > 5:
                            results.append(f"    ... 还有 {len(matches) - 5} 个匹配")
                        results.append("")
                    else:
                        results.append(f"❌ 第{i}行: {line} - 无匹配")
                        results.append("")

                except re.error as e:
                    results.append(f"❌ 第{i}行: {line} - 错误: {str(e)}")
                    results.append("")

            if results:
                self.test_output.setPlainText("\n".join(results))
            else:
                self.test_output.setPlainText(self.lang_manager.tr("no_valid_regex", "没有有效的正则表达式进行测试"))

        except Exception as e:
            self.test_output.setPlainText(f"{self.lang_manager.tr('test_error', '测试时发生错误')}: {str(e)}")

    def load_selected_rule(self, value=None):
        """加载选中的规则"""
        rule_name = self.rule_name_combo.currentText()
        if not rule_name or not self.current_config:
            return

        try:
            category = self.regex_category.currentText()
            # 支持中文和英文
            if category in ["安全检测", "Security Detection"]:
                rule_id = self.display_name_to_id.get(rule_name)
                if not rule_id:
                    return
                rule_data = self.current_config[rule_id]

                name_index = rule_data["name"].index(rule_name)

                if name_index < len(rule_data["rules"]):
                    self.rule_name_edit.setText(rule_name)
                    self.regex_edit.setPlainText(rule_data["rules"][name_index])
                else:
                    self.rule_name_edit.setText(rule_name)
                    self.regex_edit.clear()
            else:
                rule_data = self.current_config[rule_name]

                if isinstance(rule_data, dict):
                    self.rule_name_edit.setText(rule_name)
                    self.regex_edit.setPlainText("\n".join(rule_data.get("rules", [])))
                elif isinstance(rule_data, str):
                    self.rule_name_edit.setText(rule_name)
                    self.regex_edit.setPlainText(rule_data)
                elif isinstance(rule_data, list):
                    self.rule_name_edit.setText(rule_name)
                    self.regex_edit.setPlainText("\n".join(rule_data))
        except Exception as e:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                f"{self.lang_manager.tr('load_rule_failed', '加载规则失败')}: {str(e)}"
            )

    # 加载编辑正则库
    def load_regex_config(self, *_):
        """加载正则表达式配置"""

        category = self.regex_category.currentText()
        self.display_name_to_id = {}  # 清空映射

        try:
            # 支持中英文判断
            if category in ["日志格式识别", "Log Format Recognition"]:
                self.current_config = self.config.get("log_formats", {})
                self.rule_name_combo.clear()
                self.rule_name_combo.addItems(list(self.current_config.keys()))

            elif category in ["日志解析", "Log Parsing"]:
                self.current_config = self.config.get("parsers", {})
                self.rule_name_combo.clear()
                self.rule_name_combo.addItems(list(self.current_config.keys()))

            elif category in ["安全检测", "Security Detection"]:
                self.current_config = self.config.get("safety_testing", {})
                self.rule_name_combo.clear()
                self.display_name_to_id = {}

                for key, value in self.current_config.items():
                    for name in value.get("name", []):
                        self.rule_name_combo.addItem(name)
                        self.display_name_to_id[name] = key

            if self.rule_name_combo.count() > 0:
                self.rule_name_combo.setCurrentIndex(0)
                self.load_selected_rule()

        except Exception as e:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                f"{self.lang_manager.tr('load_regex_config_failed', '加载正则表达式配置失败')}: {str(e)}"
            )

    def insert_regex_example(self, pattern):
        """插入常用正则表达式示例"""
        current_text = self.regex_edit.toPlainText()
        if current_text:
            current_text += "\n" + pattern
        else:
            current_text = pattern
        self.regex_edit.setPlainText(current_text)

    def load_regex_from_table(self, index):
        """从表格加载选中的正则表达式到「编辑」页"""
        row = index.row()
        category_item = self.regex_table.item(row, 0)
        name_item = self.regex_table.item(row, 1)
        pattern_item = self.regex_table.item(row, 2)
        if not category_item or not name_item or not pattern_item:
            return

        category = category_item.text()
        rule_name = name_item.text()
        pattern = pattern_item.text()

        # Combo 顺序：0 安全检测 / 1 日志格式识别 / 2 日志解析
        cat_idx = -1
        for i in range(self.regex_category.count()):
            if self.regex_category.itemText(i) == category:
                cat_idx = i
                break
        if cat_idx < 0:
            mapping = {
                "安全检测": 0, "Security Detection": 0,
                "日志格式识别": 1, "Log Format Recognition": 1,
                "日志解析": 2, "Log Parsing": 2,
            }
            cat_idx = mapping.get(category, 0)

        # 切到「编辑正则表达式库」子页（通常 index=1）
        if hasattr(self, "re_matching_tabs") and self.re_matching_tabs:
            for i in range(self.re_matching_tabs.count()):
                title = self.re_matching_tabs.tabText(i)
                if "编辑" in title or "Edit" in title:
                    self.re_matching_tabs.setCurrentIndex(i)
                    break

        self.regex_category.blockSignals(True)
        self.regex_category.setCurrentIndex(max(0, cat_idx))
        self.regex_category.blockSignals(False)

        self.load_regex_config()
        found = False
        for i in range(self.rule_name_combo.count()):
            if self.rule_name_combo.itemText(i) == rule_name:
                self.rule_name_combo.setCurrentIndex(i)
                found = True
                break
        if not found:
            self.rule_name_edit.setText(rule_name)
        self.regex_edit.setPlainText(pattern)

    def filter_regex_library(self):
        """过滤正则表达式库"""
        search_text = self.regex_search_input.text().lower()

        for row in range(self.regex_table.rowCount()):
            match = False
            for col in range(self.regex_table.columnCount()):
                item = self.regex_table.item(row, col)
                if item and search_text in item.text().lower():
                    match = True
                    break

            self.regex_table.setRowHidden(row, not match)

    def load_regex_library(self):
        """加载正则表达式库到表格"""
        try:
            # 清空表格
            self.regex_table.setRowCount(0)

            row = 0
            # 加载日志格式识别规则
            log_formats = self.config.get("log_formats", {})
            for name, pattern in log_formats.items():
                self.regex_table.insertRow(row)
                self.regex_table.setItem(row, 0, QTableWidgetItem(
                    self.lang_manager.tr("log_format_recognition", "日志格式识别")))
                self.regex_table.setItem(row, 1, QTableWidgetItem(name))
                self.regex_table.setItem(row, 2, QTableWidgetItem(str(pattern)))
                self.regex_table.setItem(row, 3,
                                         QTableWidgetItem(self.lang_manager.tr("log_format_rule", "日志格式识别规则")))
                row += 1

            # 加载日志解析规则
            parsers = self.config.get("parsers", {})
            for name, pattern in parsers.items():
                self.regex_table.insertRow(row)
                self.regex_table.setItem(row, 0, QTableWidgetItem(self.lang_manager.tr("log_parsing", "日志解析")))
                self.regex_table.setItem(row, 1, QTableWidgetItem(name))
                self.regex_table.setItem(row, 2, QTableWidgetItem(str(pattern)))
                self.regex_table.setItem(row, 3,
                                         QTableWidgetItem(self.lang_manager.tr("log_parsing_rule", "日志解析规则")))
                row += 1

            # 加载安全检测规则
            safety_testing = self.config.get("safety_testing", {})
            for rule_id, rule_data in safety_testing.items():
                for i, (name, pattern) in enumerate(zip(rule_data.get("name", []), rule_data.get("rules", []))):
                    self.regex_table.insertRow(row)
                    self.regex_table.setItem(row, 0,
                                             QTableWidgetItem(self.lang_manager.tr("security_detection", "安全检测")))
                    self.regex_table.setItem(row, 1, QTableWidgetItem(name))
                    self.regex_table.setItem(row, 2, QTableWidgetItem(pattern))
                    self.regex_table.setItem(row, 3, QTableWidgetItem(
                        f"{self.lang_manager.tr('detection_location', '检测位置')}: {', '.join(rule_data.get('detection_location', []))}"
                    ))
                    row += 1

            # 调整列宽
            self.regex_table.resizeColumnsToContents()

        except Exception as e:
            QMessageBox.warning(self, self.lang_manager.tr("error", "错误"),
                                f"{self.lang_manager.tr('load_regex_library_failed', '加载正则表达式库失败')}: {str(e)}")

    ############################ 上面是正则规则模块 #####################
    def add_shadow_effects(self, *_):
        """为关键组件添加阴影效果"""
        # 为侧边栏添加阴影
        sidebar_shadow = QGraphicsDropShadowEffect()
        sidebar_shadow.setBlurRadius(15)
        sidebar_shadow.setXOffset(5)
        sidebar_shadow.setYOffset(0)
        sidebar_shadow.setColor(QColor(0, 0, 0, 150))
        self.sidebar.setGraphicsEffect(sidebar_shadow)

        # 为工作区添加阴影
        workspace_shadow = QGraphicsDropShadowEffect()
        workspace_shadow.setBlurRadius(20)
        workspace_shadow.setXOffset(0)
        workspace_shadow.setYOffset(0)
        workspace_shadow.setColor(QColor(0, 0, 0, 100))
        self.workspace.setGraphicsEffect(workspace_shadow)

        # 为仪表盘卡片添加阴影
        for card in self.stats_cards.values():
            card_shadow = QGraphicsDropShadowEffect()
            card_shadow.setBlurRadius(10)
            card_shadow.setXOffset(3)
            card_shadow.setYOffset(3)
            card_shadow.setColor(QColor(0, 0, 0, 80))
            card['widget'].setGraphicsEffect(card_shadow)

    def switch_tab(self, item, *_):
        """切换标签页"""
        tab_name = item.data(Qt.ItemDataRole.UserRole)
        if tab_name in self.tabs:
            self.workspace.setCurrentWidget(self.tabs[tab_name])
            # 回到仪表盘时默认看最近操作日志最新条目
            if tab_name == "dashboard":
                QTimer.singleShot(0, self._scroll_recent_to_bottom)

    def close_tab(self, index, *_):
        """关闭标签页"""
        widget = self.workspace.widget(index)
        tab_name = None
        for name, w in self.tabs.items():
            if w == widget:
                tab_name = name
                break
        if tab_name and tab_name not in ["dashboard", "analysis", "stats", "dns", "replay", "log"]:
            self.workspace.removeTab(index)
            del self.tabs[tab_name]

    def create_intelligence_tab(self, *_):
        """情报分析（公开版占位）"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        tip = QLabel("公开版暂时不提供服务")
        tip.setAlignment(Qt.AlignmentFlag.AlignCenter)
        tip.setStyleSheet("color:#666; font-size:14px;")
        layout.addStretch(1)
        layout.addWidget(tip)
        layout.addStretch(1)
        self.tabs["intelligence"] = tab
        self.workspace.addTab(tab, QIcon("ico/intelligence.png"), "情报分析")

    def _intelligence_export_module(self):
        root = os.path.dirname(os.path.abspath(__file__))
        if root not in sys.path:
            sys.path.insert(0, root)
        from Intelligence.globe_export import (  # noqa: WPS433
            build_globe_payload,
            export_globe_from_stats,
            write_globe_html,
        )
        return build_globe_payload, export_globe_from_stats, write_globe_html

    def open_access_globe(self, *args, open_browser=True):
        """从当前 url_stats 生成地球 HTML 并打开系统浏览器。

        注意：QPushButton.clicked 会传入 bool，不能把它当成 open_browser。
        """
        # clicked(False) 时 args[0] 为 False —— 忽略，默认仍打开浏览器
        url_stats = getattr(self, "url_stats", None)
        if not url_stats:
            QMessageBox.information(
                self, "提示",
                "当前没有分析数据。\n请先完成日志/流量分析后再打开访问地球。"
            )
            return
        try:
            build_globe_payload, export_globe_from_stats, _ = self._intelligence_export_module()
            payload = build_globe_payload(url_stats, title="TrafficEye 全球访问地球")
            if not payload.get("countries"):
                QMessageBox.warning(
                    self, "无访问数据",
                    "未能从结果中提取到来源国家。\n请确认分析时已做 IP 归属（Go 分析默认会做）。"
                )
            out = export_globe_from_stats(url_stats, title="TrafficEye 全球访问地球")
            s = payload.get("summary") or {}
            hint = (
                f"地球已生成 · 请求 {s.get('requests', 0)} · "
                f"IP {s.get('unique_ips', 0)} · 国家/地区 {s.get('countries', 0)}"
            )
            if hasattr(self, "stats_hint_label"):
                self.stats_hint_label.setText(hint)
            if open_browser:
                ok = self._open_local_in_browser(out)
                if ok:
                    self.status_label.setText("已在浏览器打开访问地球")
                else:
                    QMessageBox.information(
                        self, "请手动打开",
                        f"已生成地球文件，但自动打开浏览器失败。\n请手动打开：\n{out}"
                    )
            else:
                QMessageBox.information(self, "已导出", f"HTML 已保存到：\n{out}")
        except Exception as e:
            QMessageBox.critical(self, "生成失败", str(e))

    def _open_local_in_browser(self, path: str) -> bool:
        """用系统默认浏览器打开本地 HTML（兼容中文路径）。"""
        abs_path = os.path.abspath(path)
        try:
            if QDesktopServices.openUrl(QUrl.fromLocalFile(abs_path)):
                return True
        except Exception:
            pass
        try:
            if sys.platform == "win32":
                os.startfile(abs_path)  # noqa: S606
                return True
        except Exception:
            pass
        try:
            webbrowser.open(Path(abs_path).as_uri())
            return True
        except Exception:
            return False

    # 兼容旧按钮名
    def open_intelligence_globe(self, *args, open_browser=True):
        return self.open_access_globe(*args, open_browser=open_browser)

    def open_intelligence_globe_demo(self, *_):
        """打开示例地球（无分析数据时预览）。"""
        try:
            _, _, write_globe_html = self._intelligence_export_module()
            from Intelligence.globe_export import build_globe_payload
            demo_stats = {
                "data": {
                    "/demo": {
                        "count": 150,
                        "source_ips": {
                            "1.1.1.1：中国-北京": {"count": 80, "danger": []},
                            "8.8.8.8：美国-加利福尼亚": {"count": 50, "danger": ["SQL注入"]},
                            "203.0.113.9：日本-东京": {"count": 20, "danger": []},
                        },
                    }
                }
            }
            payload = build_globe_payload(demo_stats, title="TrafficEye 全球访问地球（示例）")
            out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
            os.makedirs(out_dir, exist_ok=True)
            out = os.path.join(out_dir, "trafficeye_globe_demo.html")
            write_globe_html(payload, out)
            self._open_local_in_browser(out)
            self.status_label.setText("已打开示例访问地球")
        except Exception as e:
            QMessageBox.critical(self, "打开失败", str(e))

    def create_ai_tab(self, *_):
        """创建 AI Agent 标签页（DeepSeek /anthropic + 工具循环）"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/ai.png"), self.lang_manager.tr("ai_analysis", "AI分析"))
        self.tabs["ai"] = tab

        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(6)

        # —— 顶栏 ——
        top = QWidget()
        top_l = QHBoxLayout(top)
        top_l.setContentsMargins(0, 0, 0, 0)
        top_l.setSpacing(8)

        self.ai_model_label = QLabel("DeepSeek Agent")
        self.ai_model_label.setStyleSheet("color: #007ACC; font-weight: 600;")
        top_l.addWidget(self.ai_model_label)

        self.ai_auto_analyze_check = QCheckBox("分析完成后自动研判")
        top_l.addWidget(self.ai_auto_analyze_check)
        top_l.addStretch()

        self.ai_config_btn = QPushButton("配置")
        self.ai_config_btn.setFixedWidth(64)
        self.ai_config_btn.clicked.connect(self.open_ai_agent_config)
        top_l.addWidget(self.ai_config_btn)

        self.ai_analyze_btn = QPushButton("安全研判")
        self.ai_analyze_btn.setIcon(QIcon("ico/ai.png"))
        self.ai_analyze_btn.clicked.connect(self.start_ai_analysis)
        top_l.addWidget(self.ai_analyze_btn)

        self.ai_stop_btn = QPushButton(self.lang_manager.tr("stop_analysis", "停止"))
        self.ai_stop_btn.setIcon(QIcon("ico/stop.png"))
        self.ai_stop_btn.setEnabled(False)
        self.ai_stop_btn.clicked.connect(self.stop_ai_analysis)
        top_l.addWidget(self.ai_stop_btn)

        layout.addWidget(top)

        # 兼容旧代码对 combo 的引用（不再展示）
        self.ai_model_combo = QComboBox()
        self.ai_model_combo.addItems(["DeepSeek Agent"])
        self.ai_model_combo.hide()
        self.ai_analyze_urls = QCheckBox()
        self.ai_analyze_urls.setChecked(True)
        self.ai_analyze_urls.hide()
        self.ai_analyze_params = QCheckBox()
        self.ai_analyze_params.hide()
        self.ai_analyze_headers = QCheckBox()
        self.ai_analyze_headers.hide()
        self.ai_request_body = QCheckBox()
        self.ai_request_body.hide()
        self.ai_request_Body = self.ai_request_body

        # —— 步骤日志 + 最终结论 ——
        split = QSplitter(Qt.Orientation.Horizontal)

        left = QWidget()
        left_l = QVBoxLayout(left)
        left_l.setContentsMargins(0, 0, 0, 0)
        left_l.setSpacing(4)
        left_l.addWidget(QLabel("Agent 步骤"))
        self.ai_agent_view = QPlainTextEdit()
        self.ai_agent_view.setReadOnly(True)
        self.ai_agent_view.setPlaceholderText("工具调用步骤会显示在这里…")
        self.ai_agent_view.setStyleSheet(AI_CONSOLE_QSS)
        self.ai_agent_view.setLineWrapMode(QPlainTextEdit.LineWrapMode.WidgetWidth)
        self._ai_log_highlighter = AgentLogHighlighter(self.ai_agent_view.document())
        left_l.addWidget(self.ai_agent_view)
        split.addWidget(left)

        right = QWidget()
        right_l = QVBoxLayout(right)
        right_l.setContentsMargins(0, 0, 0, 0)
        right_l.setSpacing(4)
        right_l.addWidget(QLabel("研判结论"))
        self.ai_result_text = QTextEdit()
        self.ai_result_text.setReadOnly(True)
        self.ai_result_text.setPlaceholderText("最终中文研判报告…")
        self.ai_result_text.setStyleSheet(AI_CONSOLE_QSS)
        self._ai_report_highlighter = AgentReportHighlighter(self.ai_result_text.document())
        right_l.addWidget(self.ai_result_text)
        # 兼容旧 update_ai_content
        self.ai_analysis_result_text = self.ai_agent_view
        split.addWidget(right)
        split.setStretchFactor(0, 1)
        split.setStretchFactor(1, 1)
        layout.addWidget(split, stretch=1)

        # —— 自由提问 ——
        ask_row = QWidget()
        ask_l = QHBoxLayout(ask_row)
        ask_l.setContentsMargins(0, 0, 0, 0)
        ask_l.setSpacing(6)
        self.ai_goal_edit = QLineEdit()
        self.ai_goal_edit.setPlaceholderText("向 Agent 提问，例如：高危里有没有命令注入？来源 IP 有哪些？")
        self.ai_goal_edit.returnPressed.connect(self.send_ai_agent_goal)
        ask_l.addWidget(self.ai_goal_edit, 1)
        self.ai_send_btn = QPushButton("发送")
        self.ai_send_btn.clicked.connect(self.send_ai_agent_goal)
        ask_l.addWidget(self.ai_send_btn)
        layout.addWidget(ask_row)

        self.ai_analysis_thread = None
        self._agent_worker = None
        self._pending_agent_review = False
        self._refresh_ai_model_label()

    def _refresh_ai_model_label(self):
        try:
            cfg = ai_config.load_ai_config()
            model = cfg.get("model") or "deepseek-chat"
            base = ai_config.resolve_agent_base_url(cfg)
            self.ai_model_label.setText(f"{model} · Agent")
            self.ai_model_label.setToolTip(f"{base}/v1/messages")
        except Exception:
            pass

    def open_ai_agent_config(self, *_):
        """编辑 config/ai.yaml"""
        cfg = ai_config.load_ai_config()
        dlg = QDialog(self)
        dlg.setWindowTitle("AI Agent 配置")
        dlg.setMinimumWidth(460)
        form = QFormLayout(dlg)

        key_edit = QLineEdit(str(cfg.get("api_key") or ""))
        key_edit.setEchoMode(QLineEdit.EchoMode.Password)
        key_edit.setPlaceholderText("DeepSeek API Key")
        form.addRow("API Key", key_edit)

        base_edit = QLineEdit(str(cfg.get("base_url") or "https://api.deepseek.com/v1"))
        form.addRow("Base URL", base_edit)

        agent_base_edit = QLineEdit(str(cfg.get("agent_base_url") or ""))
        agent_base_edit.setPlaceholderText("空=自动推成 /anthropic")
        form.addRow("Agent URL", agent_base_edit)

        model_edit = QLineEdit(str(cfg.get("model") or "deepseek-chat"))
        form.addRow("模型", model_edit)

        steps_edit = QLineEdit(str(cfg.get("agent_max_steps") or 30))
        form.addRow("最大步数", steps_edit)

        proxy_check = QCheckBox("使用 HTTP 代理")
        proxy_check.setChecked(bool(cfg.get("use_http_proxy")))
        form.addRow(proxy_check)
        proxy_edit = QLineEdit(str(cfg.get("http_proxy") or "127.0.0.1:7897"))
        form.addRow("代理", proxy_edit)

        status_lab = QLabel("")
        status_lab.setWordWrap(True)
        status_lab.setStyleSheet("color: #666; font-size: 12px;")
        form.addRow(status_lab)

        btns = QHBoxLayout()
        test_btn = QPushButton("测试连接")
        ok = QPushButton("保存")
        cancel = QPushButton("取消")
        btns.addWidget(test_btn)
        btns.addStretch()
        btns.addWidget(ok)
        btns.addWidget(cancel)
        form.addRow(btns)

        def _collect_cfg():
            try:
                steps = int(steps_edit.text().strip() or "30")
            except ValueError:
                steps = 30
            return {
                "provider": "deepseek",
                "api_key": key_edit.text().strip(),
                "base_url": base_edit.text().strip(),
                "agent_base_url": agent_base_edit.text().strip(),
                "model": model_edit.text().strip() or "deepseek-chat",
                "agent_max_steps": max(3, min(steps, 60)),
                "use_http_proxy": proxy_check.isChecked(),
                "http_proxy": proxy_edit.text().strip(),
            }

        def _save():
            ai_config.save_ai_config(_collect_cfg())
            self._refresh_ai_model_label()
            dlg.accept()

        class _TestConnThread(QThread):
            done = pyqtSignal(bool, str)

            def __init__(self, conf, parent=None):
                super().__init__(parent)
                self.conf = conf

            def run(self):
                ok_flag, msg = ai_config.test_agent_connection(self.conf)
                self.done.emit(ok_flag, msg)

        def _on_test_done(ok_flag, msg):
            test_btn.setEnabled(True)
            ok.setEnabled(True)
            cancel.setEnabled(True)
            if ok_flag:
                status_lab.setStyleSheet("color: #2E7D32; font-size: 12px;")
                status_lab.setText(msg)
                QMessageBox.information(dlg, "测试连接", msg)
            else:
                status_lab.setStyleSheet("color: #C62828; font-size: 12px;")
                status_lab.setText(msg)
                QMessageBox.warning(dlg, "测试连接失败", msg)

        def _test():
            conf = _collect_cfg()
            if not conf.get("api_key"):
                QMessageBox.warning(dlg, "提示", "请先填写 API Key")
                return
            endpoint = ai_config.resolve_agent_base_url(conf) + "/v1/messages"
            status_lab.setStyleSheet("color: #666; font-size: 12px;")
            status_lab.setText(f"正在测试…\n{endpoint}")
            test_btn.setEnabled(False)
            ok.setEnabled(False)
            cancel.setEnabled(False)
            th = _TestConnThread(conf, dlg)
            dlg._test_thread = th  # 防止被回收
            th.done.connect(_on_test_done)
            th.start()

        test_btn.clicked.connect(_test)
        ok.clicked.connect(_save)
        cancel.clicked.connect(dlg.reject)
        dlg.exec()

    def _resolve_pcap_traffic_path(self, stats_json: str) -> str:
        """由 output_pcap_*.json 解析配对的 output_pcap_traffic_*.txt；精确名优先，其次同秒附近。"""
        if not stats_json:
            return ""
        abs_s = os.path.abspath(stats_json)
        dirname = os.path.dirname(abs_s)
        base = os.path.basename(abs_s)
        if "output_pcap_" not in base or not base.endswith(".json"):
            return ""
        exact = os.path.join(
            dirname,
            base.replace("output_pcap_", "output_pcap_traffic_", 1).replace(".json", ".txt"),
        )
        if os.path.isfile(exact):
            return exact
        # 旧版本 JSON/TXT 可能差 1～2 秒时间戳
        try:
            ts = int(base.replace("output_pcap_", "").replace(".json", ""))
        except ValueError:
            return ""
        best = ""
        best_delta = 10
        try:
            for name in os.listdir(dirname):
                if not (name.startswith("output_pcap_traffic_") and name.endswith(".txt")):
                    continue
                mid = name[len("output_pcap_traffic_"):-len(".txt")]
                try:
                    t2 = int(mid)
                except ValueError:
                    continue
                delta = abs(t2 - ts)
                if delta < best_delta:
                    best_delta = delta
                    best = os.path.join(dirname, name)
        except OSError:
            return ""
        return best if best and os.path.isfile(best) else ""

    def _remember_analysis_artifacts(self, stats_json: str = "", traffic_file: str = "", source_file: str = ""):
        """记录最近一次分析结果路径，供 Agent / 风险页绑定到「当前这份」分析。"""
        if source_file:
            self.last_analysis_source = os.path.abspath(source_file) if os.path.isfile(source_file) else source_file.strip()

        if stats_json:
            abs_s = os.path.abspath(stats_json)
            if os.path.isfile(abs_s):
                self.last_stats_json_path = abs_s
                # 换了新的统计结果：先清旧明文，再按配对找回，避免 Agent 读到上一次文件
                self.last_traffic_path = ""
                if traffic_file and os.path.isfile(os.path.abspath(traffic_file)):
                    self.last_traffic_path = os.path.abspath(traffic_file)
                else:
                    cand = self._resolve_pcap_traffic_path(abs_s)
                    if cand:
                        self.last_traffic_path = os.path.abspath(cand)
                self._update_risk_source_label()
                return

        if traffic_file:
            abs_t = os.path.abspath(traffic_file)
            if os.path.isfile(abs_t):
                self.last_traffic_path = abs_t
            elif abs_t.lower().endswith(".txt"):
                self.last_traffic_path = abs_t
        self._update_risk_source_label()

    def _clear_analysis_artifacts(self):
        """开始新分析前清空绑定，避免风险/Agent 仍指向旧文件。"""
        self.last_stats_json_path = ""
        self.last_traffic_path = ""
        self.last_analysis_source = ""
        self._update_risk_source_label()

    def _clear_stats_view(self):
        """清空统计树与图表（无数据时 update_stats_display 会直接 return，旧界面会残留）。"""
        if hasattr(self, "url_tree") and self.url_tree is not None:
            try:
                self.url_tree.setModel(None)
            except Exception:
                pass
        for attr in ("status_chart_view", "ip_chart_view", "uri_chart_view", "time_chart_view"):
            view = getattr(self, attr, None)
            if view is None:
                continue
            try:
                view.setChart(QChart())
            except Exception:
                pass
        if hasattr(self, "stats_hint_label") and self.stats_hint_label is not None:
            self.stats_hint_label.setText("")

    def _reset_analysis_session(self):
        """每次开始新分析前清空上次会话，保证全流量 / LOG / 表格互不串结果。"""
        self._clear_analysis_artifacts()
        self.url_stats = INITIALIZE_DEFAULTDICT()
        self.memory_optimization = False
        self.start_analysis_timestamp = ""

        if hasattr(self, "analysis_text_edit") and self.analysis_text_edit is not None:
            self.analysis_text_edit.clear()
        if hasattr(self, "status_create_analysis_tab") and self.status_create_analysis_tab is not None:
            self.status_create_analysis_tab.setText("")

        if hasattr(self, "log_table") and self.log_table is not None:
            self.log_table.setRowCount(0)

        if hasattr(self, "_clear_dns_view"):
            try:
                self._clear_dns_view()
            except Exception:
                pass

        self._clear_stats_view()
        self.update_risk_display()
        self.update_dashboard_stats()

    def _update_risk_source_label(self):
        lab = getattr(self, "risk_source_label", None)
        if lab is None:
            return
        src = getattr(self, "last_analysis_source", "") or ""
        stats = getattr(self, "last_stats_json_path", "") or ""
        if src:
            lab.setText(f"当前风险来源：{os.path.basename(src)}")
            lab.setToolTip(src)
        elif stats:
            lab.setText(f"当前风险来源：{os.path.basename(stats)}")
            lab.setToolTip(stats)
        else:
            lab.setText("当前风险来源：尚未绑定分析文件")
            lab.setToolTip("")

    def _agent_session(self) -> SessionData:
        def _stats():
            return extract_url_stats_data(getattr(self, "url_stats", None))

        out_dir = getattr(self, "output_dir", None) or get_output_dir()

        return SessionData(
            url_stats_provider=_stats,
            stats_json_provider=lambda: getattr(self, "last_stats_json_path", "") or "",
            traffic_file_provider=lambda: getattr(self, "last_traffic_path", "") or "",
            output_dir_provider=lambda: out_dir,
        )

    def _has_agent_data(self) -> bool:
        if extract_url_stats_data(getattr(self, "url_stats", None)):
            return True
        for p in (
            getattr(self, "last_stats_json_path", ""),
            getattr(self, "last_traffic_path", ""),
        ):
            if p and os.path.isfile(p):
                return True
        return False

    def _set_ai_running(self, running: bool):
        self.ai_analyze_btn.setEnabled(not running)
        self.ai_send_btn.setEnabled(not running)
        self.ai_stop_btn.setEnabled(running)
        self.ai_goal_edit.setEnabled(not running)

    def start_traffic_agent(self, goal: str, *_):
        """启动 DeepSeek Agent（需已有 url_stats 或落盘产物）"""
        goal = (goal or "").strip()
        if not goal:
            return
        if not self._has_agent_data():
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                "暂无分析数据。请先完成流量/日志分析，再启动 Agent。",
            )
            return

        if self._agent_worker and self._agent_worker.isRunning():
            QMessageBox.information(self, "提示", "Agent 正在运行，请先停止或等待完成。")
            return

        self.ai_agent_view.clear()
        self.ai_result_text.clear()
        self.ai_agent_view.appendPlainText("—— 你 ——")
        self.ai_agent_view.appendPlainText(goal)
        self.ai_agent_view.appendPlainText("")
        self.ai_agent_view.appendPlainText("—— Agent ——")
        self.ai_agent_view.appendPlainText("")
        self._set_ai_running(True)
        self.status_label.setText("AI Agent 研判中…")

        # 切到 AI 页
        try:
            if "ai" in self.tabs:
                self.workspace.setCurrentWidget(self.tabs["ai"])
        except Exception:
            pass

        cfg = ai_config.load_ai_config()
        self._agent_worker = AgentWorker(goal, self._agent_session(), cfg=cfg, parent=self)
        self.ai_analysis_thread = self._agent_worker  # 兼容旧停止逻辑字段名
        self._agent_worker.log.connect(self._on_agent_log)
        self._agent_worker.finished_ok.connect(self._on_agent_ok)
        self._agent_worker.failed.connect(self._on_agent_fail)
        src = getattr(self, "last_analysis_source", "") or self.Import_box.text().strip()
        if src:
            self.ai_agent_view.appendPlainText(f"分析文件: {src}")
        self._agent_worker.start()
        self.add_recent_activity("AI Agent", self.Import_box.text() or "", "研判中…")

    def _on_agent_log(self, msg: str):
        self.ai_agent_view.appendPlainText(msg)

    def _on_agent_ok(self, text: str):
        self._set_ai_running(False)
        self.ai_result_text.setPlainText(text or "")
        self.ai_agent_view.appendPlainText("")
        self.ai_agent_view.appendPlainText("✅ 完成")
        self.status_label.setText("AI Agent 完成")
        self.add_recent_activity("AI Agent", self.Import_box.text() or "", "完成")

    def _on_agent_fail(self, err: str):
        self._set_ai_running(False)
        self.ai_agent_view.appendPlainText(f"❌ {err}")
        self.status_label.setText(f"AI Agent: {err}")
        self.add_recent_activity("AI Agent", self.Import_box.text() or "", err)

    def send_ai_agent_goal(self, *_):
        goal = self.ai_goal_edit.text().strip()
        if not goal:
            return
        self.start_traffic_agent(goal)
        self.ai_goal_edit.clear()

    def start_ai_analysis(self, *_):
        """安全研判：对当前分析结果跑 Agent（不再重跑 Go / Ollama）"""
        self.start_traffic_agent(SECURITY_REVIEW_GOAL)

    def maybe_start_pending_agent(self, *_):
        """分析完成后启动 Agent（由完成回调在确认需要时调用）"""
        self._pending_agent_review = False
        QTimer.singleShot(200, lambda: self.start_traffic_agent(SECURITY_REVIEW_GOAL))

    #######################风险分析#########################
    def create_risk_tab(self, *_):
        """创建风险分析标签页（紧凑：等级数字 + 搜索）"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/risk.png"), "风险分析")
        self.tabs["risk"] = tab

        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 8, 10, 8)
        layout.setSpacing(6)

        # —— 单行工具条：等级统计 + 搜索（贴合白主题，不用彩色胶囊）——
        bar = QWidget()
        bar.setFixedHeight(32)
        tb = QHBoxLayout(bar)
        tb.setContentsMargins(0, 0, 0, 0)
        tb.setSpacing(0)

        def _make_stat(title, value_color, filter_level):
            """标签灰字 + 数值点色，点击可筛选"""
            wrap = QWidget()
            wrap.setCursor(Qt.CursorShape.PointingHandCursor)
            wrap.setToolTip("显示全部" if filter_level == "所有级别" else f"筛选：{title}")
            hl = QHBoxLayout(wrap)
            hl.setContentsMargins(6, 0, 6, 0)
            hl.setSpacing(4)
            lab = QLabel(title)
            lab.setStyleSheet("color: #888888; font-size: 12px; background: transparent;")
            val = QLabel("0")
            val.setStyleSheet(
                f"color: {value_color}; font-size: 13px; font-weight: 600; background: transparent;"
            )
            hl.addWidget(lab)
            hl.addWidget(val)
            wrap.mousePressEvent = lambda e, lv=filter_level: self._set_risk_severity_filter(lv)
            return wrap, val

        def _sep():
            s = QLabel("·")
            s.setStyleSheet("color: #C8C8C8; font-size: 12px; padding: 0 2px; background: transparent;")
            return s

        # 与主题蓝 / 树内等级色一致，仅数字着色
        self.risk_chip_total, self.risk_chip_total_val = _make_stat("总数", "#007ACC", "所有级别")
        self.risk_chip_critical, self.risk_chip_critical_val = _make_stat("严重", "#D32F2F", "严重")
        self.risk_chip_high, self.risk_chip_high_val = _make_stat("高危", "#E65100", "高危")
        self.risk_chip_medium, self.risk_chip_medium_val = _make_stat("中危", "#F9A825", "中危")
        self.risk_chip_low, self.risk_chip_low_val = _make_stat("低危", "#388E3C", "低危")

        for i, w in enumerate((
            self.risk_chip_total, self.risk_chip_critical, self.risk_chip_high,
            self.risk_chip_medium, self.risk_chip_low
        )):
            if i:
                tb.addWidget(_sep())
            tb.addWidget(w)

        tb.addSpacing(8)

        self.risk_filter_edit = QLineEdit()
        self.risk_filter_edit.setPlaceholderText("搜索 URL / IP / 规则…")
        self.risk_filter_edit.setClearButtonEnabled(True)
        self.risk_filter_edit.setFixedHeight(28)
        self.risk_filter_edit.textChanged.connect(self.filter_risk_data)
        tb.addWidget(self.risk_filter_edit, 1)


        # 等级筛选保留为隐藏默认「所有级别」，避免占位；点击数字可筛也可后续加
        self.risk_severity_combo = QComboBox()
        self.risk_severity_combo.addItems(["所有级别", "严重", "高危", "中危", "低危", "信息"])
        self.risk_severity_combo.hide()

        self.refresh_risk_btn = QToolButton()
        self.refresh_risk_btn.setIcon(QIcon("ico/replay.png"))
        self.refresh_risk_btn.setToolTip("刷新")
        self.refresh_risk_btn.setAutoRaise(True)
        self.refresh_risk_btn.clicked.connect(self.update_risk_display)
        tb.addWidget(self.refresh_risk_btn)

        self.export_risk_btn = QToolButton()
        self.export_risk_btn.setIcon(QIcon("ico/export.png"))
        self.export_risk_btn.setToolTip("导出 CSV")
        self.export_risk_btn.setAutoRaise(True)
        self.export_risk_btn.clicked.connect(self.export_risk_data)
        tb.addWidget(self.export_risk_btn)

        risk_fullscreen_btn = QToolButton()
        risk_fullscreen_btn.setIcon(QIcon("ico/fullscreen.png"))
        risk_fullscreen_btn.setToolTip(self.lang_manager.tr("fullscreen", "全屏"))
        risk_fullscreen_btn.setAutoRaise(True)
        risk_fullscreen_btn.clicked.connect(
            lambda: self.fullscreen_manager.enter_fullscreen(
                self.risk_tree, self.lang_manager.tr("risk_statistics", "风险统计")
            )
        )
        tb.addWidget(risk_fullscreen_btn)

        layout.addWidget(bar)

        self.risk_source_label = QLabel("当前风险来源：尚未绑定分析文件")
        self.risk_source_label.setStyleSheet(
            "color: #666666; font-size: 11px; padding: 0 2px; background: transparent;"
        )
        self.risk_source_label.setWordWrap(False)
        layout.addWidget(self.risk_source_label)

        # —— 风险树（占主要空间）——
        self.risk_tree = QTreeView()
        self.risk_tree.setSortingEnabled(False)
        self.risk_tree.setSelectionMode(QTreeView.SelectionMode.ExtendedSelection)
        self.risk_tree.setUniformRowHeights(True)
        self.risk_tree.setAlternatingRowColors(True)
        self.risk_tree.setAnimated(False)
        self.risk_tree.setExpandsOnDoubleClick(True)
        self.risk_tree.setRootIsDecorated(True)
        self.risk_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.risk_tree.customContextMenuRequested.connect(self.show_risk_context_menu)
        layout.addWidget(self.risk_tree, stretch=5)

        self.risk_tree_menu = QMenu(self)
        self.copy_url_action = self.risk_tree_menu.addAction(
            self.lang_manager.tr("copy_url", "复制URL"), self.copy_risk_url
        )
        self.copy_details_action = self.risk_tree_menu.addAction(
            self.lang_manager.tr("copy_details", "复制详情"), self.copy_risk_details
        )
        self.risk_tree_menu.addSeparator()
        self.exclude_action = self.risk_tree_menu.addAction(
            self.lang_manager.tr("exclude_risk", "排除此风险"), self.exclude_selected_risk
        )

        # —— 详情（收矮）——
        self.risk_detail = QTextEdit()
        self.risk_detail.setReadOnly(True)
        self.risk_detail.setMaximumHeight(160)
        self.risk_detail.setPlaceholderText("点击明细条目查看匹配上下文")
        layout.addWidget(self.risk_detail, stretch=1)

        self.risk_model = QStandardItemModel()
        self.filter_proxy_model = QSortFilterProxyModel()
        self.filter_proxy_model.setFilterCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self.filter_proxy_model.setSourceModel(self.risk_model)
        self.risk_tree.setModel(self.filter_proxy_model)

        try:
            self.risk_tree.selectionModel().selectionChanged.disconnect()
        except Exception:
            pass
        self.risk_tree.selectionModel().selectionChanged.connect(self.show_risk_detail)

        return tab

    def show_risk_detail(self, selected, deselected, *_):
        """显示选中的风险详情"""
        indexes = selected.indexes()
        if not indexes:
            return

        index = self.filter_proxy_model.mapToSource(indexes[0])
        item = index.internalPointer()
        if hasattr(item, 'risk_data'):
            risk = item.risk_data
            level_color = self.get_severity_color(risk['level'])
            detail = f"""
            <div style="font-family: 'Segoe UI', Arial, sans-serif; color: #222; background: #FFF; padding: 10px;">
                <h3 style="color: #007ACC; margin: 0 0 10px; border-bottom: 2px solid #E0E0E0; padding-bottom: 6px;">
                    {self.lang_manager.tr("risk_details", "风险详情")}
                    <span style="float:right; color:{level_color}; font-size:14px;">{risk['level']}</span>
                </h3>
                <table style="width:100%; border-collapse:collapse; margin-bottom:12px;">
                    <tr>
                        <td style="width:100px; color:#666; padding:6px 0; font-weight:bold;">来源 IP</td>
                        <td style="padding:6px 0;">{risk.get('ip', '')}</td>
                    </tr>
                    <tr>
                        <td style="color:#666; padding:6px 0; font-weight:bold;">URL</td>
                        <td style="padding:6px 0; word-break:break-all;">{unquote(risk['url'])}</td>
                    </tr>
                    <tr>
                        <td style="color:#666; padding:6px 0; font-weight:bold;">风险类型</td>
                        <td style="padding:6px 0;">{risk['type']}</td>
                    </tr>
                    <tr>
                        <td style="color:#666; padding:6px 0; font-weight:bold;">匹配规则</td>
                        <td style="padding:6px 0;">{risk['rule']}</td>
                    </tr>
                    <tr>
                        <td style="color:#666; padding:6px 0; font-weight:bold;">匹配位置</td>
                        <td style="padding:6px 0;">{risk['position'][0]}-{risk['position'][1]}</td>
                    </tr>
                </table>
                <h4 style="color:#007ACC; margin:8px 0;">匹配上下文</h4>
                <pre style="background:#F8F8F8; color:#333; padding:12px; border:1px solid #DCDCDC;
                            border-radius:4px; white-space:pre-wrap; word-wrap:break-word; margin:0; font-size:13px;">{risk['content']}</pre>
            </div>
            """
            self.risk_detail.setHtml(detail)
        else:
            # 点到分组节点时给出提示
            self.risk_detail.setHtml(
                "<div style='color:#666;padding:12px;'>请展开并选择具体风险条目，查看完整匹配上下文。</div>"
            )

    def update_risk_display(self, *_):
        """更新风险信息显示"""
        data = extract_url_stats_data(getattr(self, "url_stats", None))

        if not data:
            # 清空旧风险，避免仍显示上一次分析结果
            self.risk_model = QStandardItemModel()
            if hasattr(self, "filter_proxy_model"):
                self.filter_proxy_model.setSourceModel(self.risk_model)
            if hasattr(self, "risk_tree"):
                self.risk_tree.setModel(self.filter_proxy_model)
            if hasattr(self, "risk_detail"):
                self.risk_detail.clear()
            self.update_risk_stats()
            self._update_risk_source_label()
            return

        self.risk_model = RiskTreeModel(data, lang_manager=self.lang_manager)
        self.filter_proxy_model.setSourceModel(self.risk_model)
        self.risk_tree.setModel(self.filter_proxy_model)

        selection_model = self.risk_tree.selectionModel()
        if selection_model:
            selection_model.selectionChanged.connect(self.show_risk_detail)

        self.risk_tree.setColumnWidth(0, 180)
        self.risk_tree.setColumnWidth(1, 280)
        self.risk_tree.setColumnWidth(2, 120)
        self.risk_tree.setColumnWidth(3, 80)
        self.risk_tree.setColumnWidth(4, 160)
        self.risk_tree.setColumnWidth(5, 220)

        # 默认展开严重/高危两级分组，便于立刻看到重点
        self.risk_tree.collapseAll()
        root = self.risk_model.root_item
        for i in range(root.child_count()):
            level_item = root.child(i)
            level_name = getattr(level_item, "group_level", "")
            idx = self.risk_model.index(i, 0)
            if level_name in ("严重", "高危"):
                self.risk_tree.expand(self.filter_proxy_model.mapFromSource(idx))
                # 展开其下类型分组
                for j in range(level_item.child_count()):
                    type_idx = self.risk_model.index(j, 0, idx)
                    self.risk_tree.expand(self.filter_proxy_model.mapFromSource(type_idx))

        self.update_risk_stats()
        self.filter_risk_data()
        self._update_risk_source_label()

    def get_severity_color(self, severity):
        """根据严重性级别返回对应的颜色"""
        colors = {
            "严重": "#FF4D4D",
            "高危": "#FF8C4D",
            "中危": "#FFC84D",
            "低危": "#4DCCFF",
            "信息": "#AAAAAA",
            "中": "#FFC84D",
        }
        return colors.get(severity, "#AAAAAA")

    def export_risk_data(self, *_):
        """导出风险数据为CSV文件"""
        if not hasattr(self, 'risk_model') or not isinstance(self.risk_model, RiskTreeModel):
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("no_risk_data_export", "没有可导出的风险数据!"))
            return

        file_name, _ = QFileDialog.getSaveFileName(
            self,
            self.lang_manager.tr("save_risk_data", "保存风险数据"),
            f"risk_analysis_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            self.lang_manager.tr("csv_files", "CSV 文件 (*.csv);;所有文件 (*)")
        )
        if not file_name:
            return

        try:
            with open(file_name, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                writer.writerow([
                    "来源IP",
                    self.lang_manager.tr("url", "URL"),
                    self.lang_manager.tr("risk_type", "风险类型"),
                    self.lang_manager.tr("risk_level", "风险等级"),
                    self.lang_manager.tr("matched_rule", "匹配规则"),
                    self.lang_manager.tr("matched_position", "匹配位置"),
                    self.lang_manager.tr("matched_content", "匹配内容"),
                ])

                # 三级结构：严重级别 → 类型 → 明细
                root = self.risk_model.root_item
                for i in range(root.child_count()):
                    level_item = root.child(i)
                    for j in range(level_item.child_count()):
                        type_item = level_item.child(j)
                        for k in range(type_item.child_count()):
                            risk_item = type_item.child(k)
                            if not hasattr(risk_item, 'risk_data'):
                                continue
                            risk = risk_item.risk_data
                            pos = risk.get('position') or (0, 0)
                            writer.writerow([
                                risk.get('ip', ''),
                                unquote(risk.get('url', '')),
                                risk.get('type', ''),
                                risk.get('level', ''),
                                risk.get('rule', ''),
                                f"{pos[0]}-{pos[1]}" if isinstance(pos, (list, tuple)) and len(pos) >= 2 else "",
                                risk.get('content', ''),
                            ])

            self.status_label.setText(self.lang_manager.tr("risk_export_success", "风险数据已成功导出"))
            self.add_recent_activity(self.lang_manager.tr("risk_export_success", "风险数据已成功导出"), file_name,
                                     self.lang_manager.tr("done", "完成"))
            QMessageBox.information(self, self.lang_manager.tr("done", "完成"), f"风险数据已成功导出到:\n{file_name}")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")
            self.add_recent_activity("导出风险数据", file_name, f"失败: {str(e)}")

    def update_risk_stats(self):
        """只刷新顶部等级数字"""
        zeros = True
        critical = high = medium = low = total = 0
        if hasattr(self, 'risk_model') and isinstance(self.risk_model, RiskTreeModel):
            zeros = False
            stats = self.risk_model.stats
            total = stats.get("total", 0)
            by_sev = stats.get("by_severity", {})
            critical = by_sev.get("严重", 0)
            high = by_sev.get("高危", 0)
            medium = by_sev.get("中危", 0) + by_sev.get("中", 0)
            low = by_sev.get("低危", 0)

        if hasattr(self, 'risk_chip_total_val'):
            self.risk_chip_total_val.setText(str(0 if zeros else total))
            self.risk_chip_critical_val.setText(str(0 if zeros else critical))
            self.risk_chip_high_val.setText(str(0 if zeros else high))
            self.risk_chip_medium_val.setText(str(0 if zeros else medium))
            self.risk_chip_low_val.setText(str(0 if zeros else low))

    def _set_risk_severity_filter(self, level):
        """点击等级数字切换筛选"""
        if not hasattr(self, 'risk_severity_combo'):
            return
        idx = self.risk_severity_combo.findText(level)
        if idx >= 0:
            self.risk_severity_combo.setCurrentIndex(idx)
            self.filter_risk_data()

    def filter_risk_data(self):
        """按等级 + 关键字筛选（适配：严重级别 → 类型 → 明细）"""
        if not hasattr(self, 'risk_model') or not isinstance(self.risk_model, RiskTreeModel):
            return
        if not hasattr(self, 'risk_filter_edit') or not hasattr(self, 'risk_severity_combo'):
            return

        filter_text = (self.risk_filter_edit.text() or "").strip().lower()
        severity_filter = self.risk_severity_combo.currentText()
        proxy = self.filter_proxy_model

        root = self.risk_model.root_item

        for i in range(root.child_count()):
            level_item = root.child(i)
            level_name = getattr(level_item, "group_level", "") or ""
            level_src = self.risk_model.index(i, 0)
            level_proxy = proxy.mapFromSource(level_src)
            if not level_proxy.isValid():
                continue
            level_visible = False

            if severity_filter != "所有级别" and severity_filter not in level_name:
                self.risk_tree.setRowHidden(level_proxy.row(), QModelIndex(), True)
                continue

            for j in range(level_item.child_count()):
                type_item = level_item.child(j)
                type_src = self.risk_model.index(j, 0, level_src)
                type_proxy = proxy.mapFromSource(type_src)
                if not type_proxy.isValid():
                    continue
                type_visible = False

                for k in range(type_item.child_count()):
                    leaf = type_item.child(k)
                    leaf_src = self.risk_model.index(k, 0, type_src)
                    leaf_proxy = proxy.mapFromSource(leaf_src)
                    if not leaf_proxy.isValid():
                        continue

                    match = True
                    if hasattr(leaf, 'risk_data'):
                        risk = leaf.risk_data
                        if severity_filter != "所有级别" and severity_filter not in risk.get('level', ''):
                            match = False
                        if match and filter_text:
                            blob = " ".join([
                                str(risk.get('ip', '')),
                                str(risk.get('ip_short', '')),
                                unquote(str(risk.get('url', ''))),
                                str(risk.get('type', '')),
                                str(risk.get('rule', '')),
                                str(risk.get('content', '')),
                                str(risk.get('level', '')),
                            ]).lower()
                            match = filter_text in blob
                    else:
                        match = not filter_text

                    self.risk_tree.setRowHidden(leaf_proxy.row(), type_proxy, not match)
                    if match:
                        type_visible = True
                        level_visible = True

                self.risk_tree.setRowHidden(type_proxy.row(), level_proxy, not type_visible)

            self.risk_tree.setRowHidden(level_proxy.row(), QModelIndex(), not level_visible)

            if level_visible and (
                filter_text or severity_filter != "所有级别" or level_name in ("严重", "高危")
            ):
                self.risk_tree.expand(level_proxy)
                for j in range(level_item.child_count()):
                    type_src = self.risk_model.index(j, 0, level_src)
                    type_proxy = proxy.mapFromSource(type_src)
                    if type_proxy.isValid() and not self.risk_tree.isRowHidden(type_proxy.row(), level_proxy):
                        self.risk_tree.expand(type_proxy)

    def copy_risk_detail(self):
        """复制风险详情内容"""
        if self.risk_detail.toPlainText():
            QApplication.clipboard().setText(self.risk_detail.toPlainText())

    def exclude_selected_risk(self):
        """排除选中的风险"""
        indexes = self.risk_tree.selectedIndexes()
        if indexes:
            index = self.filter_proxy_model.mapToSource(indexes[0])
            item = index.internalPointer()
            if hasattr(item, 'risk_data'):
                risk = item.risk_data
                # 在实际应用中，您可能需要将排除的风险保存到配置文件中
                QMessageBox.information(self, "排除风险", f"已排除风险: {risk['type']} - {risk['url']}")

    def copy_risk_details(self):
        """复制选中的风险详情"""
        indexes = self.risk_tree.selectedIndexes()
        if indexes:
            index = self.filter_proxy_model.mapToSource(indexes[0])
            item = index.internalPointer()
            if hasattr(item, 'risk_data'):
                risk = item.risk_data
                details = (
                    f"来源IP: {risk.get('ip', '')}\n"
                    f"{self.lang_manager.tr('url', 'URL')}: {unquote(risk.get('url', ''))}\n"
                    f"{self.lang_manager.tr('risk_type', '风险类型')}: {risk.get('type', '')}\n"
                    f"{self.lang_manager.tr('risk_level', '风险等级')}: {risk.get('level', '')}\n"
                    f"{self.lang_manager.tr('matched_rule', '匹配规则')}: {risk.get('rule', '')}\n"
                    f"{self.lang_manager.tr('matched_content', '匹配内容')}: {risk.get('content', '')}"
                )
                QApplication.clipboard().setText(details)

    def copy_risk_url(self):
        """复制选中的风险URL"""
        indexes = self.risk_tree.selectedIndexes()
        if indexes:
            index = self.filter_proxy_model.mapToSource(indexes[0])
            item = index.internalPointer()
            if hasattr(item, 'risk_data'):
                QApplication.clipboard().setText(item.risk_data['url'])

    def show_risk_context_menu(self, position):
        """显示风险表格的右键菜单"""
        indexes = self.risk_tree.selectedIndexes()
        if indexes:
            self.risk_tree_menu.exec(self.risk_tree.viewport().mapToGlobal(position))

    #######################风险分析#########################
    def create_dashboard_tab(self, *_):
        """创建仪表盘标签页 (已适配白色主题)"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/dashboard.png"), "仪表盘")
        self.tabs["dashboard"] = tab

        layout = QVBoxLayout(tab)

        # --- 欢迎面板 (已修复硬编码颜色) ---
        welcome_panel = QGroupBox(self.lang_manager.tr("Welcome_to", "欢迎使用 TrafficEye"))
        welcome_layout = QVBoxLayout(welcome_panel)

        # HTML 样式已修改，以适应白色主题
        welcome_label = QLabel(f"""
            <div style="text-align: center; font-family: 'Segoe UI', sans-serif;">
                <h1 style="
                    font-size: 30px;
                    color: #007ACC; /* 更改为主题蓝色 */
                    font-weight: 600;
                    margin: 0;
                ">
                    TrafficEye Web {self.lang_manager.tr("issuance", issuance)}
                </h1>
                <p style="margin: 6px 0 2px; font-size: 16px; color: #444444;">{self.lang_manager.tr("Introduction_name", "网络流量分析工具")}</p>
                <p style="margin: 0; font-size: 13px; color: #777777;">
                    {self.lang_manager.tr("version", "版本")}: {version} | {self.lang_manager.tr("Last_update", "最后更新")}: {last_updated} | W啥都学
                </p>
            </div>
        """)
        welcome_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        welcome_layout.addWidget(welcome_label)

        # --- 快速操作按钮区域 (已移除 setStyleSheet) ---
        quick_actions = QWidget()
        quick_layout = QHBoxLayout(quick_actions)
        quick_layout.setContentsMargins(0, 0, 0, 0)
        quick_layout.setSpacing(10)  # 按钮间距

        # 分析按钮 (使用 setObjectName)
        self.btn_analyze = QPushButton(QIcon("ico/analyze.png"),
                                       self.lang_manager.tr("auto_analysis", "一键识别自动化分析"))
        self.btn_analyze.clicked.connect(self._on_analyze_button_clicked)
        self.btn_analyze.setObjectName("AnalyzeButton")  # <-- 设置对象名
        self.btn_analyze.setIconSize(QSize(24, 24))
        self._analysis_ui_running = False

        # 重放按钮 (使用 setObjectName)
        btn_replay = QPushButton(QIcon("ico/replay.png"), self.lang_manager.tr("replay", "请求重放"))
        btn_replay.clicked.connect(lambda: self.workspace.setCurrentWidget(self.tabs["replay"]))
        btn_replay.setObjectName("ReplayButton")  # <-- 设置对象名
        btn_replay.setIconSize(QSize(24, 24))

        # 统计按钮 (使用 setObjectName)
        btn_stats = QPushButton(QIcon("ico/stats.png"), self.lang_manager.tr("stats", "查看统计"))
        btn_stats.clicked.connect(lambda: self.workspace.setCurrentWidget(self.tabs["stats"]))
        btn_stats.setObjectName("StatsButton")  # <-- 设置对象名
        btn_stats.setIconSize(QSize(24, 24))

        # 添加按钮
        quick_layout.addWidget(self.btn_analyze)
        quick_layout.addWidget(btn_replay)
        quick_layout.addWidget(btn_stats)

        welcome_layout.addWidget(quick_actions)
        layout.addWidget(welcome_panel)

        # --- 统计概览 (已移除 setStyleSheet) ---
        stats_panel = QGroupBox(self.lang_manager.tr("stats_overview", "统计概览"))
        stats_layout = QHBoxLayout(stats_panel)
        stats_layout.setContentsMargins(15, 15, 15, 15)
        stats_layout.setSpacing(15)

        self.stats_cards = {
            "total": {"widget": QWidget(), "value": QLabel("0")},
            "unique_url": {"widget": QWidget(), "value": QLabel("0")},
            "source_ip": {"widget": QWidget(), "value": QLabel("0")},
            "status_code": {"widget": QWidget(), "value": QLabel("0")},
            "danger": {"widget": QWidget(), "value": QLabel("0")}
        }

        cards = [
            {"key": "total", "icon": "ico/traffic.png", "title": self.lang_manager.tr("total", "总请求数"),
             "color": "#2196F3"},
            {"key": "unique_url", "icon": "ico/url.png", "title": self.lang_manager.tr("unique_url", "唯一URL"),
             "color": "#00BCD4"},
            {"key": "source_ip", "icon": "ico/ip.png", "title": self.lang_manager.tr("source_ip", "来源IP"),
             "color": "#FF9800"},
            {"key": "status_code", "icon": "ico/status.png", "title": self.lang_manager.tr("status_code", "状态码"),
             "color": "#4CAF50"},
            {"key": "danger", "icon": "ico/danger.png", "title": self.lang_manager.tr("danger", "攻击危险"),
             "color": "#F44336"}
        ]

        for card in cards:
            card_data = self.stats_cards[card["key"]]
            card_widget = card_data["widget"]
            # 为卡片设置 objectName 以便在 YAML 中统一定义样式
            card_widget.setObjectName("StatCard")

            card_layout = QVBoxLayout(card_widget)
            card_layout.setContentsMargins(15, 15, 15, 15)
            card_layout.setSpacing(10)

            # 图标和标题
            top_widget = QWidget()
            top_layout = QHBoxLayout(top_widget)
            top_layout.setContentsMargins(0, 0, 0, 0)

            icon = QLabel()
            icon.setPixmap(QIcon(card["icon"]).pixmap(24, 24))

            title = QLabel(card["title"])
            # 保留这里的颜色设置，因为它是动态的
            title.setStyleSheet(
                f"color: {card['color']}; font-weight: bold; font-size: 14px; background-color: transparent;")

            top_layout.addWidget(icon)
            top_layout.addWidget(title)
            top_layout.addStretch()

            # 数值
            value = card_data["value"]
            value.setAlignment(Qt.AlignmentFlag.AlignCenter)
            # 为数值标签设置 objectName
            value.setObjectName("StatValue")

            card_layout.addWidget(top_widget)
            card_layout.addWidget(value)
            stats_layout.addWidget(card_widget)

        layout.addWidget(stats_panel)

        # --- 最近活动 (无样式改动) ---
        recent_panel = QGroupBox(self.lang_manager.tr("recent_ops", "最近操作日志"))
        recent_layout = QVBoxLayout(recent_panel)
        recent_layout.setContentsMargins(0, 15, 0, 0)

        self.recent_table = QTableWidget()
        # --- 修改: 列数改为 5 ---
        self.recent_table.setColumnCount(5)
        self.recent_table.setShowGrid(False)
        self.recent_table.setHorizontalHeaderLabels(
            [self.lang_manager.tr("time", "时间"), self.lang_manager.tr("action", "操作"),
             self.lang_manager.tr("file", "文件"), self.lang_manager.tr("status", "状态"),
             self.lang_manager.tr("view_result", "查看结果")])  # <-- 新增列标题
        self.recent_table.horizontalHeader().setStretchLastSection(True)
        self.recent_table.verticalHeader().setVisible(False)

        self.recent_table.setColumnWidth(0, 200)
        self.recent_table.setColumnWidth(1, 200)
        self.recent_table.setColumnWidth(2, 700)
        self.recent_table.setColumnWidth(3, 200)
        self.recent_table.setColumnWidth(4, 80)
        # --- 新增: 设置第5列的宽度 ---

        self.recent_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.recent_table.customContextMenuRequested.connect(self.show_recent_context_menu)

        # --- 新增: 连接点击信号 ---
        self.recent_table.itemClicked.connect(self.on_recent_activity_clicked)
        # --- 新增: 改善UI体验 ---
        self.recent_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.recent_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)

        recent_layout.addWidget(self.recent_table)
        layout.addWidget(recent_panel)

        self.load_recent_activity()

    def _load_traffic_text_to_ui(self, file_path):
        """
        辅助方法：读取大文本文件的前1000行显示到UI，并提供完整文件链接
        """
        if not file_path or not os.path.exists(file_path):
            self.analysis_text_edit.setPlainText(self.lang_manager.tr("file_not_found", "文件未找到"))
            return


        # 2. 读取文件预览
        max_lines = 1000
        # try:
        #     if self.max_lines_input.text():
        #         max_lines = int(self.max_lines_input.text())
        # except:
        #     pass

        preview_content = ""
        line_count = 0
        is_truncated = False

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    if line_count < max_lines:
                        preview_content += line
                        line_count += 1
                    else:
                        is_truncated = True
                        break

            self.analysis_text_edit.setPlainText(preview_content)

            # 3. 设置状态栏提示和链接
            if is_truncated:
                # 获取绝对路径以确保链接有效
                abs_path = os.path.abspath(file_path)
                file_url = QUrl.fromLocalFile(abs_path).toString()

                html = (
                    f"<span style='color:#E65100;font-size:11px;line-height:14px;'>"
                    f"[{self.lang_manager.tr('tip', '提示')}] "
                    f"输出超过 {max_lines} 行，完整内容已保存 · "
                    f"<a href=\"{file_url}\" style='color:#1565C0;'>"
                    f"{self.lang_manager.tr('click_to_open', '打开完整文件')}</a></span>"
                )
                self.status_create_analysis_tab.setTextFormat(Qt.TextFormat.RichText)
                self.status_create_analysis_tab.setTextInteractionFlags(Qt.TextInteractionFlag.TextBrowserInteraction)
                self.status_create_analysis_tab.setOpenExternalLinks(True)
                self.status_create_analysis_tab.setText(html)
            else:
                self.status_create_analysis_tab.setText(f"已加载: {os.path.basename(file_path)} ({line_count} 行)")

        except Exception as e:
            self.analysis_text_edit.setPlainText(f"读取文件失败: {e}")
    def open_recent_result_by_row(self, row: int):
        """
        根据行号，打开“最近操作”中关联的分析结果文件。
        """
        # 1. 从第一列获取 QTableWidgetItem
        first_item = self.recent_table.item(row, 0)
        if not first_item:
            return

        # 2. 获取存储在 UserRole 中的结果文件路径 (通常是 JSON)
        result_file_path = first_item.data(Qt.ItemDataRole.UserRole)

        # 3. 检查路径是否存在
        if result_file_path and os.path.exists(result_file_path):
            def _load_history():
                try:
                    # 同步「当前文件」到该条历史记录对应的源文件，避免风险页与 Import 框不一致
                    file_item = self.recent_table.item(row, 2)
                    hist_src = file_item.text().strip() if file_item else ""
                    if hist_src:
                        self.Import_box.setText(hist_src)
                        self.last_analysis_source = hist_src

                    if result_file_path.endswith('.json'):
                        with open(result_file_path, 'r', encoding='utf-8') as f:
                            final_stats = json.load(f)

                        self.url_stats = final_stats
                        self._remember_analysis_artifacts(
                            stats_json=result_file_path,
                            source_file=hist_src or getattr(self, "last_analysis_source", ""),
                        )
                        if self.url_stats.get('data'):
                            if self.loading_overlay:
                                self.loading_overlay.set_message(
                                    self.lang_manager.tr("rendering_stats", "正在渲染统计界面...")
                                )
                            self.update_stats_display()
                            self.parse_web_server_log()
                            self.update_dashboard_stats()
                        else:
                            self.update_risk_display()

                        basename = os.path.basename(result_file_path)

                        if "output_pcap_" in basename:
                            traffic_path = self._resolve_pcap_traffic_path(result_file_path)

                            if traffic_path and os.path.exists(traffic_path):
                                self._remember_analysis_artifacts(
                                    traffic_file=traffic_path,
                                    source_file=hist_src or getattr(self, "last_analysis_source", ""),
                                )
                                self._load_traffic_text_to_ui(traffic_path)
                            else:
                                self.analysis_text_edit.setPlainText(
                                    "未找到关联的流量明文文件 (.txt)。\n"
                                    "原因：本次/历史分析未生成 HTTP 明文（旧版本不会写出该文件）。\n"
                                    "请重新执行一次「全流量分析」后即可在此查看。"
                                )
                                self.last_traffic_path = ""


                            # 点击历史记录时再加载配对 DNS
                            dns_link = self._dns_path_for_stats(result_file_path)
                            if dns_link and os.path.isfile(dns_link):
                                self.load_dns_results(dns_path=dns_link, persist=False)
                            else:
                                last_stats = ""
                                if hasattr(self, "ui_config") and isinstance(self.ui_config, dict):
                                    last_stats = self.ui_config.get("last_dns_stats") or ""
                                latest = self._dns_latest_path()
                                if (last_stats and os.path.abspath(last_stats) == os.path.abspath(result_file_path)
                                        and os.path.isfile(latest)):
                                    self.load_dns_results(dns_path=latest, persist=False)
                                else:
                                    self._clear_dns_view()
                        else:
                            self.analysis_text_edit.setPlainText("此记录仅包含统计结果 (JSON)。")
                            self._clear_dns_view()
                    else:
                        self._remember_analysis_artifacts(
                            traffic_file=result_file_path,
                            source_file=hist_src or getattr(self, "last_analysis_source", ""),
                        )
                        self._load_traffic_text_to_ui(result_file_path)
                        self._clear_dns_view()
                        self.update_risk_display()

                    self.status_label.setText(f"已加载历史记录: {os.path.basename(result_file_path)}")
                except Exception as e:
                    QMessageBox.warning(self, "加载失败", f"无法加载历史记录:\n{e}")

            self.run_with_loading(
                _load_history,
                message=self.lang_manager.tr("loading_history", "正在加载历史数据，请稍候...")
            )
            return

        if result_file_path:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                f"文件已丢失: {result_file_path}"
            )
        else:
            QMessageBox.information(
                self,
                self.lang_manager.tr("tip", "提示"),
                self.lang_manager.tr("no_result_file", "该记录没有可加载的结果文件。")
            )

    def on_recent_activity_clicked(self, item: QTableWidgetItem):
        """
        点击“最近操作”表格中的某一行时触发。
        如果该行有关联的分析结果文件，则加载该文件内容到“流量分析”选项卡。
        """
        if not item:
            return

        row = item.row()
        self.open_recent_result_by_row(item.row())

        # 1. 从第一列获取 QTableWidgetItem
        first_item = self.recent_table.item(row, 0)
        if not first_item:
            return

        # 2. 获取存储在 UserRole 中的结果文件路径
        result_file_path = first_item.data(Qt.ItemDataRole.UserRole)

        # 3. 检查路径是否存在
        if result_file_path and os.path.exists(result_file_path):
            try:
                # 4. 读取文件内容
                with open(result_file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # 5. 切换到“流量分析”选项卡
                # (假设“流量分析”是侧边栏的第2项，索引为1)
                self.sidebar.setCurrentRow(1)
                self.workspace.setCurrentWidget(self.tabs["analysis"])

                # 6. 将内容设置到 analysis_text_edit
                self.analysis_text_edit.setPlainText(content)

                # 7. (可选) 重新加载原始的 .pcap/.log 文件路径
                file_item = self.recent_table.item(row, 2)
                if file_item and os.path.exists(file_item.text()):
                    self.Import_box.setText(file_item.text())

                self.status_label.setText(
                    self.lang_manager.tr("loaded_result_file", "已加载结果文件: ") + os.path.basename(result_file_path))

            except Exception as e:
                QMessageBox.warning(self,
                                    self.lang_manager.tr("error", "错误"),
                                    self.lang_manager.tr("failed_to_load_result", "加载分析结果文件失败: ") + f"\n{e}"
                                    )
                self.status_label.setText(self.lang_manager.tr("failed_to_load_result_status", "加载结果失败"))
        else:
            # 如果没有关联文件，或者文件已丢失，则不执行任何操作
            pass



    def show_recent_context_menu(self, position, *_):
        """ 删除该条记录或全部记录 """
        index = self.recent_table.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()

        # "删除该条记录" 操作
        delete_action = menu.addAction(self.lang_manager.tr("delete_record", "删除该条记录"))

        # "删除所有记录" 操作
        delete_all_action = menu.addAction(self.lang_manager.tr("delete_all", "删除所有记录"))

        # 右键菜单展示
        action = menu.exec(self.recent_table.viewport().mapToGlobal(position))

        if action == delete_action:
            row = index.row()

            # 删除表格中的该行
            self.recent_table.removeRow(row)

            # 删除数据源中的对应记录
            self.delete_recent_entry(row)

        elif action == delete_all_action:
            # 删除所有记录
            row_count = self.recent_table.rowCount()
            for row in range(row_count - 1, -1, -1):  # 从最后一行开始删除
                self.recent_table.removeRow(row)
                self.delete_recent_entry(row)

    def delete_recent_entry(self, row_index, *_):
        """ 删除该条记录存储 """
        path = "history/trafficeye_data.json"
        if not os.path.exists(path):
            return
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        recent = data.get("recent", [])
        if 0 <= row_index < len(recent):
            del recent[row_index]
            data["recent"] = recent
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

    def update_dashboard_stats(self, *_):
        """更新仪表盘统计信息"""
        if not hasattr(self, 'stats_cards'):
            return

        g = compute_global_stats(getattr(self, "url_stats", None))
        # 远程扁平结构时回写，后续摘要/报告可读到
        try:
            if isinstance(self.url_stats, dict):
                self.url_stats["_global_stats"] = g
        except Exception:
            pass

        self.stats_cards["total"]["value"].setText(str(g.get("request_total", 0)))
        self.stats_cards["unique_url"]["value"].setText(str(g.get("total_unique_uris", 0)))
        self.stats_cards["source_ip"]["value"].setText(str(g.get("total_unique_ips", 0)))
        self.stats_cards["status_code"]["value"].setText(str(g.get("total_unique_status_code", 0)))
        self.stats_cards["danger"]["value"].setText(str(g.get("danger_total", 0)))

    def create_analysis_tab(self, *_):
        """创建流量分析选项卡 (使用 UI 文件)"""
        try:
            # 1. 加载 UI 文件
            ui_path = os.path.join("ui", "analysis_tab.ui")
            if not os.path.exists(ui_path):
                QMessageBox.critical(self, "错误", f"UI文件未找到: {ui_path}")
                return

            self.analysis_tab_widget = uic.loadUi(ui_path)

            # 添加到主 TabWidget
            self.workspace.addTab(
                self.analysis_tab_widget,
                QIcon("ico/analysis.png"),
                self.lang_manager.tr("analysis", "分析提取的HTTP结果")
            )
            self.tabs["analysis"] = self.analysis_tab_widget

            # 2. 获取控件引用 (绑定到 self)
            self.analysis_text_edit = self.analysis_tab_widget.analysis_text_edit
            self.status_create_analysis_tab = self.analysis_tab_widget.status_create_analysis_tab
            # 底部提示条：尽量紧凑，不占大块区域
            self.status_create_analysis_tab.setWordWrap(False)
            self.status_create_analysis_tab.setMaximumHeight(22)
            self.status_create_analysis_tab.setMinimumHeight(0)
            self.status_create_analysis_tab.setAlignment(
                Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter
            )
            self.status_create_analysis_tab.setStyleSheet(
                "QLabel { font-size: 11px; color: #666; padding: 0 2px; margin: 0; }"
            )
            self.status_create_analysis_tab.setSizePolicy(
                QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed
            )

            # 分割器
            self.analysis_splitter = self.analysis_tab_widget.splitter
            self.analysis_splitter.setSizes([1120, 280])

            # 3. 设置多语言文本
            self.analysis_tab_widget.result_group.setTitle(self.lang_manager.tr("analysis_result", "分析提取的HTTP结果"))
            self.analysis_text_edit.setPlaceholderText(
                self.lang_manager.tr("analysis_placeholder", "分析提取的HTTP结果将显示在这里..."))


    
        except Exception as e:
            print(f"加载流量分析Tab UI失败: {e}")
            import traceback
            traceback.print_exc()

    def load_ui_settings(self):
        """从 ui_config.yaml 加载界面设置"""
        try:
            if os.path.exists(self.ui_config_path):
                with open(self.ui_config_path, "r", encoding="utf-8") as f:
                    ui_config = yaml.safe_load(f)
                    self.ui_config = ui_config

                    self.ui_config['Interface_width'] = ui_config.get('Interface_width', 1000)
                    self.ui_config['Interface_height'] = ui_config.get('Interface_height', 800)

            else:
                self.ui_config['Interface_width'] = 1000
                self.ui_config['Interface_height'] = 800
                self.ui_config['language'] = 'zh'  # 默认为中文
        except Exception as e:
            print("加载界面配置失败：", e)
            self.ui_config = {
                'Interface_width': 1000,
                'Interface_height': 800,
                'language': 'zh'
            }

    def save_ui_settings(self):
        """保存界面设置到配置文件"""
        try:

            with open(self.ui_config_path, "w", encoding="utf-8") as f:
                yaml.dump(self.ui_config, f)
        except Exception as e:
            print("保存界面配置失败：", e)


    def create_stats_tab(self, *_):
        """创建统计信息标签页 - 深色主题"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/stats.png"), self.lang_manager.tr("stats", "统计分析"))
        self.tabs["stats"] = tab

        layout = QVBoxLayout(tab)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)

        self.fullscreen_manager = FullscreenManager(self, lang_manager=self.lang_manager)

        # URL访问统计表格
        url_group = QGroupBox(self.lang_manager.tr("url_access_stats", "URL访问统计"))

        # 创建QTreeWidget代替QTableWidget
        self.url_tree = QTreeView()

        self.url_tree.setSortingEnabled(False)  # 已按访问量预排序；自定义模型点标题排序易乱序
        self.url_tree.setSelectionMode(QTreeWidget.SelectionMode.ExtendedSelection)
        self.url_tree.setSelectionBehavior(QTreeWidget.SelectionBehavior.SelectRows)
        self.url_tree.setUniformRowHeights(True)
        self.url_tree.setAnimated(False)
        self.url_tree.setExpandsOnDoubleClick(True)
        self.url_tree.setAlternatingRowColors(True)
        self.url_tree.setRootIsDecorated(True)
        self.url_tree.setWordWrap(False)
        self.url_tree.setTextElideMode(Qt.TextElideMode.ElideMiddle)
        header = self.url_tree.header()
        header.setStretchLastSection(True)
        header.setDefaultAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)

        url_header = QWidget(url_group)
        url_header.setObjectName("qt_groupbox_titlewidget")
        url_header_layout = QHBoxLayout(url_header)
        url_header_layout.setContentsMargins(0, 0, 0, 0)
        url_header_layout.setSpacing(4)

        # 刷新/导出收进表头，避免单独占一整行
        self.refresh_stats_btn = QToolButton()
        self.refresh_stats_btn.setIcon(QIcon("ico/replay.png"))
        self.refresh_stats_btn.setToolTip(self.lang_manager.tr("refresh_stats", "刷新统计"))
        self.refresh_stats_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.refresh_stats_btn.setAutoRaise(True)
        self.refresh_stats_btn.clicked.connect(self.update_stats_display)

        self.export_stats_btn = QToolButton()
        self.export_stats_btn.setIcon(QIcon("ico/export.png"))
        self.export_stats_btn.setToolTip(self.lang_manager.tr("export_stats", "导出统计"))
        self.export_stats_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.export_stats_btn.setAutoRaise(True)
        self.export_stats_btn.clicked.connect(self.export_stats)

        self.open_globe_btn = QToolButton()
        self.open_globe_btn.setIcon(QIcon("ico/globe.png"))
        self.open_globe_btn.setToolTip("打开访问地球")
        self.open_globe_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.open_globe_btn.setAutoRaise(True)
        # 必须用 lambda：clicked(bool) 否则会把 False 传给 open_browser
        self.open_globe_btn.clicked.connect(lambda *_: self.open_access_globe(open_browser=True))

        self.stats_hint_label = QLabel("")
        self.stats_hint_label.setStyleSheet("color: #888; font-size: 12px;")

        url_header_layout.addWidget(self.refresh_stats_btn)
        url_header_layout.addWidget(self.export_stats_btn)
        url_header_layout.addWidget(self.open_globe_btn)
        url_header_layout.addStretch()
        url_header_layout.addWidget(self.stats_hint_label)

        url_fullscreen_btn = QToolButton()
        url_fullscreen_btn.setIcon(QIcon("ico/fullscreen.png"))
        url_fullscreen_btn.setToolTip(self.lang_manager.tr("fullscreen", "全屏"))
        url_fullscreen_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        url_fullscreen_btn.setAutoRaise(True)
        url_fullscreen_btn.clicked.connect(lambda: self.fullscreen_manager.enter_fullscreen(self.url_tree,
                                                                                            self.lang_manager.tr(
                                                                                                "url_access_stats",
                                                                                                "URL访问统计")))
        url_header_layout.addWidget(url_fullscreen_btn)

        url_group.setLayout(QVBoxLayout())
        url_group.layout().addWidget(url_header)
        url_group.layout().addWidget(self.url_tree)
        layout.addWidget(url_group, stretch=3)

        # 图表部分 - 横向滑动
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        charts_container = QWidget()
        charts_container.setMinimumWidth(1600)  # 设置宽度，确保内容超出可视区域时显示滚动条
        charts_layout = QHBoxLayout(charts_container)
        charts_layout.setSpacing(15)
        charts_layout.setContentsMargins(10, 10, 10, 10)

        def create_chart_group(title, chart_view):
            group = QGroupBox(title)

            chart_view.setRenderHint(QPainter.RenderHint.Antialiasing)

            header = QWidget(group)
            header.setObjectName("qt_groupbox_titlewidget")
            header_layout = QHBoxLayout(header)
            header_layout.setContentsMargins(0, 0, 0, 0)
            header_title = QLabel(self.lang_manager.tr("fullscreen", "全屏"))
            header_layout.addWidget(header_title)
            header_layout.addStretch()

            fullscreen_btn = QToolButton()
            fullscreen_btn.setIcon(QIcon("ico/fullscreen.png"))
            fullscreen_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            fullscreen_btn.clicked.connect(lambda: self.fullscreen_manager.enter_fullscreen(chart_view, title))
            header_layout.addWidget(fullscreen_btn)

            group.setLayout(QVBoxLayout())
            group.layout().addWidget(header)
            group.layout().addWidget(chart_view)
            return group

        self.status_chart_view = QChartView()
        self.ip_chart_view = QChartView()
        self.time_chart_view = QChartView()
        self.uri_chart_view = QChartView()

        charts_layout.addWidget(create_chart_group(
            self.lang_manager.tr("status_code_distribution", "状态码分布"), self.status_chart_view
        ))
        charts_layout.addWidget(create_chart_group(
            self.lang_manager.tr("source_ip_stats", "来源IP统计"), self.ip_chart_view
        ))
        charts_layout.addWidget(create_chart_group(
            self.lang_manager.tr("access_time_trend", "访问时间趋势"), self.time_chart_view
        ))
        charts_layout.addWidget(create_chart_group(
            self.lang_manager.tr("uri_access_stats", "URI访问统计"), self.uri_chart_view
        ))

        scroll_area.setWidget(charts_container)
        layout.addWidget(scroll_area, stretch=2)

        return tab

    ##########################下面是设置模块#######################################
    # 创建设置标签页

    def create_dns_tab(self, *_):
        """全流量 DNS 访问统计（仅 PCAP 分析产出）"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/logs.png"), self.lang_manager.tr("dns_access", "DNS访问"))
        self.tabs["dns"] = tab

        layout = QVBoxLayout(tab)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(8)

        # 工具条：不要固定过矮，避免按钮文字被裁切/被表格盖住
        bar = QWidget()
        bar.setMinimumHeight(36)
        bar.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        hl = QHBoxLayout(bar)
        hl.setContentsMargins(0, 2, 0, 2)
        hl.setSpacing(8)

        self.dns_search = QLineEdit()
        self.dns_search.setPlaceholderText("搜索域名 / IP / 应答")
        self.dns_search.setClearButtonEnabled(True)
        self.dns_search.setMinimumHeight(30)
        self.dns_search.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.dns_search.textChanged.connect(self._filter_dns_table)
        hl.addWidget(self.dns_search, stretch=1)

        btn_refresh = QPushButton("重新抽取")
        btn_refresh.setMinimumHeight(30)
        btn_refresh.setMinimumWidth(84)
        btn_refresh.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_refresh.setStyleSheet("padding: 4px 12px;")
        btn_refresh.clicked.connect(self.refresh_dns_from_pcap)
        hl.addWidget(btn_refresh)

        btn_export = QPushButton("导出 CSV")
        btn_export.setMinimumHeight(30)
        btn_export.setMinimumWidth(84)
        btn_export.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_export.setStyleSheet("padding: 4px 12px;")
        btn_export.clicked.connect(self.export_dns_table)
        hl.addWidget(btn_export)
        layout.addWidget(bar, stretch=0)

        self.dns_table = QTableWidget()
        self.dns_table.setColumnCount(6)
        self.dns_table.setHorizontalHeaderLabels([
            "域名", "查询次数", "类型", "源IP", "应答", "最近时间"
        ])
        self.dns_table.horizontalHeader().setStretchLastSection(False)
        # 域名列优先完整展示：按内容加宽，不够时横向滚动，不截断成 ...
        self.dns_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        for col in (1, 2, 3, 4, 5):
            self.dns_table.horizontalHeader().setSectionResizeMode(col, QHeaderView.ResizeMode.ResizeToContents)
        self.dns_table.setSortingEnabled(True)
        self.dns_table.setAlternatingRowColors(True)
        self.dns_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.dns_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.dns_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.dns_table.setShowGrid(False)
        self.dns_table.verticalHeader().setVisible(False)
        self.dns_table.verticalHeader().setDefaultSectionSize(28)
        self.dns_table.setTextElideMode(Qt.TextElideMode.ElideNone)
        self.dns_table.setWordWrap(False)
        self.dns_table.setHorizontalScrollMode(QTableWidget.ScrollMode.ScrollPerPixel)
        self.dns_table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.dns_table.setStyleSheet("""
            QTableWidget {
                background: #FFFFFF;
                color: #1A1A1A;
                border: 1px solid #E6E8EB;
                gridline-color: transparent;
                font-size: 13px;
                alternate-background-color: #F7F9FC;
            }
            QTableWidget::item {
                padding: 4px 8px;
                color: #1A1A1A;
            }
            QTableWidget::item:selected {
                background: #D6EBFF;
                color: #000000;
            }
            QHeaderView::section {
                background: #F0F2F5;
                color: #333333;
                font-weight: 600;
                font-size: 12px;
                border: none;
                border-bottom: 1px solid #E0E0E0;
                padding: 6px 8px;
            }
        """)
        layout.addWidget(self.dns_table, stretch=1)

        self.dns_empty_label = QLabel("暂无 DNS 数据")
        self.dns_empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.dns_empty_label.setStyleSheet("color: #9AA0A6; font-size: 12px; background: transparent;")
        layout.addWidget(self.dns_empty_label)
        self.dns_empty_label.hide()

        self._last_dns_data = None
        self._last_dns_path = ""
        self._dns_rows_cache = []
        # 不自动恢复：等用户点击历史记录或重新分析后再显示
        self.dns_table.hide()
        self.dns_empty_label.show()

    def _dns_latest_path(self) -> str:
        return os.path.abspath(os.path.join("output", "dns_latest.json"))

    def _dns_path_for_stats(self, stats_json: str) -> str:
        """与全流量统计 JSON 配对的 DNS 文件路径。"""
        if not stats_json:
            return ""
        dirname = os.path.dirname(os.path.abspath(stats_json))
        base = os.path.basename(stats_json)
        if "output_pcap_" in base and base.endswith(".json"):
            return os.path.join(dirname, base.replace("output_pcap_", "dns_for_pcap_", 1))
        return ""

    def _clear_dns_view(self):
        self._last_dns_data = None
        self._last_dns_path = ""
        if hasattr(self, "dns_table"):
            self.dns_table.setRowCount(0)
            self.dns_table.hide()
        if hasattr(self, "dns_empty_label"):
            self.dns_empty_label.setText("暂无 DNS 数据")
            self.dns_empty_label.show()

    def _persist_dns_data(self, data: dict, src_path: str = "", stats_json: str = ""):
        """写入 dns_latest，并按全流量统计文件生成配对 dns_for_pcap_*.json。"""
        try:
            os.makedirs("output", exist_ok=True)
            latest = self._dns_latest_path()
            with open(latest, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            self._last_dns_path = latest

            link = self._dns_path_for_stats(stats_json) if stats_json else ""
            if link:
                with open(link, "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                self._last_dns_path = link

            if hasattr(self, "ui_config") and isinstance(self.ui_config, dict):
                self.ui_config["last_dns_path"] = self._last_dns_path
                if src_path:
                    self.ui_config["last_dns_source"] = os.path.abspath(src_path)
                if stats_json:
                    self.ui_config["last_dns_stats"] = os.path.abspath(stats_json)
                self.save_ui_settings()
        except Exception as e:
            print(f"persist dns failed: {e}")

    def _restore_dns_history(self):
        """兼容旧调用：默认不自动恢复。"""
        return

    def _populate_dns_table(self, queries):
        self._dns_rows_cache = []
        self.dns_table.setSortingEnabled(False)
        self.dns_table.setRowCount(0)
        for q in queries:
            types = q.get("types") or {}
            src = q.get("src_ips") or {}
            ans = q.get("answers") or {}
            type_s = ", ".join(f"{k}:{v}" for k, v in list(types.items())[:6])
            src_s = ", ".join(list(src.keys())[:6])
            ans_s = ", ".join(list(ans.keys())[:8])
            domain = q.get("domain", "")
            count = int(q.get("count", 0) or 0)
            last_t = q.get("last_time", "")
            self._dns_rows_cache.append((domain, count, type_s, src_s, ans_s, last_t))

            r = self.dns_table.rowCount()
            self.dns_table.insertRow(r)
            item0 = QTableWidgetItem(domain)
            item0.setToolTip(domain)
            item0.setForeground(QBrush(QColor("#111111")))
            f0 = item0.font()
            f0.setPointSize(max(f0.pointSize(), 10))
            item0.setFont(f0)
            self.dns_table.setItem(r, 0, item0)
            item1 = QTableWidgetItem()
            item1.setData(Qt.ItemDataRole.DisplayRole, count)
            item1.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            if count >= 50:
                item1.setForeground(QBrush(QColor("#B71C1C")))
            elif count >= 10:
                item1.setForeground(QBrush(QColor("#E65100")))
            else:
                item1.setForeground(QBrush(QColor("#1565C0")))
            self.dns_table.setItem(r, 1, item1)
            for c, v in enumerate((type_s, src_s, ans_s, last_t), start=2):
                it = QTableWidgetItem(v)
                it.setToolTip(v)
                it.setForeground(QBrush(QColor("#333333")))
                self.dns_table.setItem(r, c, it)
        self.dns_table.setSortingEnabled(True)
        if self.dns_table.rowCount() > 0:
            self.dns_table.sortByColumn(1, Qt.SortOrder.DescendingOrder)
        self.dns_table.resizeColumnToContents(0)
        # 域名列至少占可视区域一半，尽量一眼看全
        try:
            min_domain_w = max(280, int(self.dns_table.viewport().width() * 0.45))
            if self.dns_table.columnWidth(0) < min_domain_w:
                self.dns_table.setColumnWidth(0, min_domain_w)
        except Exception:
            pass

        has = self.dns_table.rowCount() > 0
        self.dns_table.setVisible(has)
        self.dns_empty_label.setVisible(not has)

    def _filter_dns_table(self, text=""):
        needle = (text or "").strip().lower()
        if not hasattr(self, "dns_table"):
            return
        for r in range(self.dns_table.rowCount()):
            if not needle:
                self.dns_table.setRowHidden(r, False)
                continue
            blob = " ".join(
                (self.dns_table.item(r, c).text() if self.dns_table.item(r, c) else "")
                for c in range(self.dns_table.columnCount())
            ).lower()
            self.dns_table.setRowHidden(r, needle not in blob)

    def load_dns_results(self, dns_path: str = "", pcap_path: str = "", stats_json: str = "", persist: bool = True):
        """加载 DNS；persist=False 时仅展示（用于打开历史）。"""
        data = None
        src_used = ""
        if dns_path and os.path.isfile(dns_path):
            try:
                with open(dns_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                src_used = dns_path
            except Exception as e:
                self.status_label.setText(f"DNS 结果读取失败: {e}")
        if data is None and pcap_path and os.path.isfile(pcap_path):
            try:
                data = table_import.extract_dns_with_tshark(pcap_path)
                src_used = pcap_path
            except Exception as e:
                self.status_label.setText(f"DNS 回退抽取失败: {e}")

        if not data:
            self._clear_dns_view()
            return

        self._last_dns_data = data
        self._last_dns_path = src_used or dns_path or ""
        self._populate_dns_table(data.get("queries") or [])
        if persist:
            self._persist_dns_data(data, src_used, stats_json=stats_json)
        if hasattr(self, "dns_search"):
            self._filter_dns_table(self.dns_search.text())

    def export_dns_table(self):
        if not self._last_dns_data:
            QMessageBox.information(self, "提示", "没有可导出的 DNS 数据")
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "导出 DNS", f"dns_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            "CSV (*.csv)"
        )
        if not path:
            return
        import csv as _csv
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            w = _csv.writer(f)
            w.writerow(["域名", "查询次数", "类型", "源IP", "应答", "最近时间"])
            for q in self._last_dns_data.get("queries") or []:
                types = q.get("types") or {}
                src = q.get("src_ips") or {}
                ans = q.get("answers") or {}
                w.writerow([
                    q.get("domain", ""),
                    q.get("count", 0),
                    ",".join(f"{k}:{v}" for k, v in types.items()),
                    ",".join(src.keys()),
                    ",".join(ans.keys()),
                    q.get("last_time", ""),
                ])
        QMessageBox.information(self, "完成", f"已导出: {path}")

    def refresh_dns_from_pcap(self):
        file = self.Import_box.text().strip()
        if not file.lower().endswith((".pcap", ".pcapng", ".cap")):
            QMessageBox.warning(self, "提示", "请先选择 PCAP 全流量文件")
            return
        stats = getattr(self, "last_stats_json_path", "") or ""
        self.load_dns_results(pcap_path=file, stats_json=stats, persist=True)
        if self._last_dns_data:
            QMessageBox.information(self, "完成", "DNS 已重新抽取")
            if "dns" in self.tabs:
                self.workspace.setCurrentWidget(self.tabs["dns"])


    def create_settings_tab(self, *_):
        """创建设置标签页"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/settings.png"), "设置")
        self.tabs["settings"] = tab

        layout = QVBoxLayout(tab)

        # 创建选项卡
        tab_widget = QTabWidget()

        # =================== 1. 常规设置 Tab ===================
        general_tab = QWidget()
        general_layout = QVBoxLayout(general_tab)

        # 应用信息
        app_info_group = QGroupBox(self.lang_manager.tr("app_info", "应用信息"))
        app_info_layout = QFormLayout(app_info_group)

        self.app_name = QLabel("TrafficEye Web")
        self.app_version = QLabel(self.lang_manager.tr('version', '版本') + f": {version}")
        self.app_author = QLabel("W啥都学")
        self.last_update = QLabel(self.lang_manager.tr('last_update', '最后更新') + f": {last_updated}")

        app_info_layout.addRow(self.lang_manager.tr("name", "名称") + ":", self.app_name)
        app_info_layout.addRow(self.lang_manager.tr("version", "版本") + ":", self.app_version)
        app_info_layout.addRow(self.lang_manager.tr("author", "作者") + ":", self.app_author)
        app_info_layout.addRow(self.lang_manager.tr("last_update", "最后更新") + ":", self.last_update)

        general_layout.addWidget(app_info_group)

        # 二维码区域
        image_layout = QHBoxLayout()

        wechat_image_label = QLabel()
        pixmap = QPixmap("ico/qrcode_for_gh_e911bdfdbe01_344.png")
        wechat_image_label.setPixmap(pixmap)
        wechat_image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        wechatzf_image_label = QLabel()
        pixmap_zf = QPixmap("ico/wxzf.jpg")
        wechatzf_image_label.setPixmap(pixmap_zf)
        wechatzf_image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        image_layout.addWidget(wechat_image_label)
        image_layout.addSpacing(20)
        image_layout.addWidget(wechatzf_image_label)

        general_layout.addLayout(image_layout)

        # =================== 3. 界面设置 Tab ===================
        ui_settings_tab = QWidget()
        ui_settings_layout = QVBoxLayout(ui_settings_tab)

        ui_settings_group = QGroupBox(self.lang_manager.tr("ui_settings", "界面设置"))
        form_layout = QFormLayout(ui_settings_group)

        self.width_input = QLineEdit(str(self.ui_config['Interface_width']))
        self.width_input.setValidator(QIntValidator(600, 2000))
        form_layout.addRow(self.lang_manager.tr("window_width", "窗口宽度:"), self.width_input)

        self.height_input = QLineEdit(str(self.ui_config['Interface_height']))
        self.height_input.setValidator(QIntValidator(600, 2000))
        form_layout.addRow(self.lang_manager.tr("window_height", "窗口高度:"), self.height_input)

        self.save_ui_settings_btn = QPushButton(self.lang_manager.tr("save_ui_settings", "保存界面设置"))
        self.save_ui_settings_btn.clicked.connect(self.save_ui_settings)
        form_layout.addRow(self.save_ui_settings_btn)

        ui_settings_layout.addWidget(ui_settings_group)

        # =================== 添加所有 Tab ===================
        tab_widget.addTab(general_tab, self.lang_manager.tr("general_tab", "常规"))

        tab_widget.addTab(ui_settings_tab, self.lang_manager.tr("ui_settings_tab", "界面设置"))

        layout.addWidget(tab_widget)

        # self.load_regex_config()

    ##########################上面是设置模块#######################################
    def create_replay_tab(self, *_):
        """创建请求重放选项卡 (使用 UI 文件)"""
        try:
            # 1. 加载 UI 文件
            # 假设你的 ui 文件保存在 ui/replay_tab.ui
            ui_path = os.path.join("ui", "replay_tab.ui")
            if not os.path.exists(ui_path):
                QMessageBox.critical(self, "错误", f"UI文件未找到: {ui_path}")
                return

            # 创建一个容器 Widget 并加载 UI
            self.replay_tab_widget = uic.loadUi(ui_path)

            # 添加到主窗口的 Tab 中
            self.workspace.addTab(
                self.replay_tab_widget,
                QIcon("ico/replay.png"),
                self.lang_manager.tr("replay", "请求重放")
            )
            self.tabs["replay"] = self.replay_tab_widget

            # 2. 获取控件引用 (绑定到 self 以便其他方法调用)
            # 注意：这些名称必须与 .ui 文件中的 objectName 一致
            self.stream_id_input = self.replay_tab_widget.stream_id_input
            self.find_request_button = self.replay_tab_widget.find_request_button

            self.proxy_group = self.replay_tab_widget.proxy_group
            self.http_proxy_input = self.replay_tab_widget.http_proxy_input
            self.https_proxy_input = self.replay_tab_widget.https_proxy_input

            self.replay_button = self.replay_tab_widget.replay_button
            self.clear_replay_button = self.replay_tab_widget.clear_replay_button
            self.stop_replay_button = self.replay_tab_widget.stop_replay_button

            self.request_text_edit = self.replay_tab_widget.request_text_edit
            self.status_create_replay_tab = self.replay_tab_widget.status_create_replay_tab

            # 分割器 (可选：如果你想在代码里设置初始比例)
            self.replay_splitter = self.replay_tab_widget.splitter
            self.replay_splitter.setSizes([200, 600])

            # 3. 设置多语言文本 (如果在 UI 文件里写死了中文，这里最好再重新设置一下以支持切换语言)
            self.replay_tab_widget.stream_id_group.setTitle(self.lang_manager.tr("session_id", "会话ID"))
            self.stream_id_input.setPlaceholderText(self.lang_manager.tr("input_session_id", "输入请求会话ID..."))
            self.find_request_button.setText(self.lang_manager.tr("output_id", "输出ID"))

            self.proxy_group.setTitle(self.lang_manager.tr("proxy_settings", "代理设置"))
            self.http_proxy_input.setPlaceholderText(self.lang_manager.tr("http_proxy", "HTTP代理"))
            self.https_proxy_input.setPlaceholderText(self.lang_manager.tr("https_proxy", "HTTPS代理"))

            self.replay_tab_widget.button_group.setTitle(self.lang_manager.tr("actions", "操作"))
            self.replay_button.setText(self.lang_manager.tr("replay_request", "重放请求"))
            self.clear_replay_button.setText(self.lang_manager.tr("clear_results", "清除结果"))
            self.stop_replay_button.setText(self.lang_manager.tr("stop", "停止"))

            self.replay_tab_widget.bottom_group.setTitle(self.lang_manager.tr("request_info", "请求信息"))
            self.request_text_edit.setPlaceholderText(
                self.lang_manager.tr("request_placeholder", "请求信息将显示在这里..."))

            # 4. 连接信号 (事件绑定)
            # 防重复绑定逻辑
            try:
                self.find_request_button.clicked.disconnect()
            except TypeError:
                pass
            self.find_request_button.clicked.connect(self.find_request)

            self.proxy_group.toggled.connect(self.toggle_proxy_settings)

            self.replay_button.clicked.connect(self.replay_request)
            self.clear_replay_button.clicked.connect(self.clear_replay_results)
            self.stop_replay_button.clicked.connect(self.stop_replay)

        except Exception as e:
            print(f"加载重放Tab UI失败: {e}")
            import traceback
            traceback.print_exc()

    def toggle_proxy_settings(self, enabled, *_):
        """切换代理设置可用状态"""
        self.http_proxy_input.setEnabled(enabled)
        self.https_proxy_input.setEnabled(enabled)

    def create_log_tab(self, *_):
        """创建日志分析标签页"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/logs.png"), "Log分析")
        self.tabs["log"] = tab

        layout = QVBoxLayout(tab)

        # 分析类型按钮面板
        # type_analysis_panel = QGroupBox("分析类型")
        # type_analysis_layout = QVBoxLayout(type_analysis_panel)

        # 第一行按钮
        row1 = QHBoxLayout()
        automatic = QPushButton(QIcon("ico/analyze.png"), "点击WBE LOG分析")
        apache_access = QPushButton(QIcon("ico/Apache.png"), "Apache Access")
        nginx_access = QPushButton(QIcon("ico/Nginx.png"), "Nginx Access")
        json_log = QPushButton(QIcon("ico/json.png"), "JSON Log")

        # 第二行按钮
        row2 = QHBoxLayout()
        f5_healthcheck = QPushButton(QIcon("ico/F5.png"), "F5 HealthCheck")
        haproxy_access = QPushButton(QIcon("ico/Haproxy.png"), "HAProxy Access")
        iis_log = QPushButton(QIcon("ico/IIS.png"), "IIS Log")
        tomcat_access_log = QPushButton(QIcon("ico/Tomcat.png"), "Tomcat Access Log")

        for btn in [automatic, apache_access, nginx_access, json_log,
                    f5_healthcheck, haproxy_access, iis_log, tomcat_access_log]:
            btn.setMinimumHeight(32)

        # 连接分析按钮信号
        automatic.clicked.connect(lambda: self.analyze_logs('auto'))

        # # 添加按钮到对应行
        # for btn in [automatic]:
        #     row1.addWidget(btn)
        #
        # type_analysis_layout.addLayout(row1)
        # type_analysis_layout.addLayout(row2)
        # layout.addWidget(type_analysis_panel)

        # 结果显示区域
        result_panel = QGroupBox(self.lang_manager.tr("analysis_result", "分析结果"))
        result_layout = QVBoxLayout(result_panel)

        # 创建表格控件
        self.log_table = QTableWidget()
        self.log_table.setColumnCount(6)
        self.log_table.setHorizontalHeaderLabels([
            self.lang_manager.tr("url", "URL"),
            self.lang_manager.tr("visit_count", "访问次数"),
            self.lang_manager.tr("status_code", "状态码"),
            self.lang_manager.tr("source_ip", "来源IP"),
            self.lang_manager.tr("method", "方法"),
            self.lang_manager.tr("user_agent", "UA")
        ])
        self.log_table.horizontalHeader().setStretchLastSection(True)
        self.log_table.setSortingEnabled(True)
        self.log_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.log_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)

        # 添加导出按钮
        export_btn = QPushButton(self.lang_manager.tr("export_csv", "导出为CSV"))

        export_btn.clicked.connect(self.export_log_table)

        result_layout.addWidget(self.log_table)
        result_layout.addWidget(export_btn)
        layout.addWidget(result_panel)

    def clear_log_analysis(self, *_):
        """清除日志分析内容"""
        # self.log_file_input.clear()
        self.log_table.setRowCount(0)


    def start_table_analysis(self):
        """天眼类 CSV/TSV 表格导入分析"""
        file_path = self.Import_box.text().strip()
        if not file_path.lower().endswith(('.csv', '.tsv')):
            QMessageBox.warning(self, "警告", "请选择 .csv / .tsv 表格文件")
            self._set_analyze_button_running(False)
            return
        try:
            headers, delim, preview = table_import.read_table_headers(file_path)
            mapping = table_import.auto_map_columns(headers)
            if table_import.mapping_score(mapping) < 3:
                mapping = self._ask_table_column_mapping(headers, mapping)
                if mapping is None:
                    self._set_analyze_button_running(False)
                    return
            # 独立会话：清空上次全流量/LOG 残留
            self._reset_analysis_session()
            self._set_analyze_button_running(True)
            self.progress_bar.show()
            self.progress_bar.setValue(30)
            stats, mapping_used, _ = table_import.import_table_file(file_path, mapping)
            self.progress_bar.setValue(80)
            self.url_stats = stats
            # 兼容 log 表：data 在 url_stats['data']
            if "data" in stats:
                # log_dynamic / update_stats 可能期望顶层 uri；统一成带 data 的结构
                pass
            self.update_stats_display()
            self.update_dashboard_stats()
            self._fill_log_table_from_stats(stats)
            os.makedirs("output", exist_ok=True)
            out = os.path.join("output", f"output_table_{int(time.time())}.json")
            with open(out, "w", encoding="utf-8") as f:
                json.dump(stats, f, ensure_ascii=False, indent=2)
            self._remember_analysis_artifacts(stats_json=out, source_file=file_path)
            self.progress_bar.hide()
            self._set_analyze_button_running(False)
            self.status_label.setText(
                f"表格导入完成：请求 {stats.get('_global_stats', {}).get('request_total', 0)}，"
                f"映射 { {k:v for k,v in mapping_used.items() if v} }"
            )
            if "log" in self.tabs:
                self.workspace.setCurrentWidget(self.tabs["log"])
            self.add_recent_activity("表格导入分析", file_path, "完成")
            self.maybe_start_pending_agent()
        except Exception as e:
            self.progress_bar.hide()
            self._set_analyze_button_running(False)
            QMessageBox.critical(self, "错误", f"表格导入失败: {e}")

    def _ask_table_column_mapping(self, headers, mapping):
        dlg = QDialog(self)
        dlg.setWindowTitle("表格列映射（天眼/CSV）")
        dlg.resize(480, 320)
        form = QFormLayout(dlg)
        combos = {}
        fields = [
            ("ip", "源IP"),
            ("path", "URL/URI"),
            ("method", "方法"),
            ("status", "状态码"),
            ("time", "时间"),
            ("ua", "User-Agent"),
            ("host", "Host/域名"),
        ]
        for key, label in fields:
            cb = QComboBox()
            cb.addItem("(不使用)", "")
            for h in headers:
                cb.addItem(h, h)
            cur = mapping.get(key)
            if cur:
                idx = cb.findData(cur)
                if idx >= 0:
                    cb.setCurrentIndex(idx)
            combos[key] = cb
            form.addRow(label, cb)
        btns = QHBoxLayout()
        ok = QPushButton("确定")
        cancel = QPushButton("取消")
        btns.addWidget(ok)
        btns.addWidget(cancel)
        form.addRow(btns)
        result = {"ok": False}

        def accept():
            result["ok"] = True
            dlg.accept()

        ok.clicked.connect(accept)
        cancel.clicked.connect(dlg.reject)
        if dlg.exec() != QDialog.DialogCode.Accepted or not result["ok"]:
            return None
        out = {}
        for key, cb in combos.items():
            out[key] = cb.currentData() or None
        if not out.get("path") and not out.get("host"):
            QMessageBox.warning(self, "提示", "至少需要映射 URL/URI 或 Host")
            return None
        return out

    def _fill_log_table_from_stats(self, stats):
        """把表格/日志统计填进 Log 分析表。"""
        if not hasattr(self, "log_table") or self.log_table is None:
            return
        data = stats.get("data") or {}
        # 若是扁平旧结构
        if not data and isinstance(stats, dict):
            data = {k: v for k, v in stats.items() if k not in ("_global_stats", "_import_meta", "data", "filter_applied")}
        self.log_table.setSortingEnabled(False)
        self.log_table.setRowCount(0)
        for uri, ustats in list(data.items())[:5000]:
            if not isinstance(ustats, dict):
                continue
            count = ustats.get("count", 0)
            ips = ustats.get("source_ips") or {}
            status_codes = {}
            methods = {}
            uas = {}
            for ip, ip_s in ips.items():
                if not isinstance(ip_s, dict):
                    continue
                for sc, c in (ip_s.get("status_codes") or {}).items():
                    status_codes[sc] = status_codes.get(sc, 0) + c
                for m, c in (ip_s.get("methods") or {}).items():
                    methods[m] = methods.get(m, 0) + c
                for u, c in (ip_s.get("UA") or {}).items():
                    uas[u] = uas.get(u, 0) + c
            row = self.log_table.rowCount()
            self.log_table.insertRow(row)
            self.log_table.setItem(row, 0, QTableWidgetItem(str(uri)))
            self.log_table.setItem(row, 1, QTableWidgetItem(str(count)))
            self.log_table.setItem(row, 2, QTableWidgetItem(",".join(status_codes.keys())[:80]))
            self.log_table.setItem(row, 3, QTableWidgetItem(",".join(list(ips.keys())[:5])))
            self.log_table.setItem(row, 4, QTableWidgetItem(",".join(methods.keys())[:40]))
            self.log_table.setItem(row, 5, QTableWidgetItem(",".join(list(uas.keys())[:2])[:120]))
        self.log_table.setSortingEnabled(True)


    def analyze_logs(self, log_type=None, ai_analysis_starts=None, optional_parameters_log=None, *_):

        self.Analyze_logs_locally(log_type, ai_analysis_starts, optional_parameters_log)

    def Analyze_logs_locally(self, log_type=None, ai_analysis_starts=None, optional_parameters_log=None, *_):
        """分析日志文件"""
        # --- (这是你原有的 Analyze_logs_locally 逻辑) ---
        file_path = self.Import_box.text().strip()
        if not file_path.lower().endswith(('.log', '.txt')):
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                self.lang_manager.tr("please_select_log_file", "请选择正确的.log文件或者.txt文件!"))
            return

        if not file_path:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                self.lang_manager.tr("please_select_log_first", "请先选择日志文件!")
            )
            return

        try:
            # 独立会话：清空上次全流量明文/统计/DNS，避免串到本次 LOG
            self._reset_analysis_session()
            self._analysis_stop_requested = False
            self._set_analyze_button_running(True)
            self.last_analysis_source = file_path
            self._update_risk_source_label()
            self.add_recent_activity(
                self.lang_manager.tr("start_log_analysis", "开始分析LOG日志"),
                file_path,
                self.lang_manager.tr("status_running", "运行中")
            )

            if log_type == 'auto':
                detected_type = log_identification.guess_log_format(file_path)
                if detected_type == 'unknown':
                    # 通用回退，不中断
                    self.status_label.setText("未识别到专用格式，将使用通用解析…")

            self.progress_bar.show()
            self.progress_bar.setValue(0)

            use_go = prefer_go_log_analyzer()
            if use_go:
                self.status_label.setText("正在启动 Go 日志分析服务…")
                self.start_go_backend()
                if not self.go_process and not self._go_service_reachable():
                    self.status_label.setText("Go 服务启动失败，已回退 Python 解析（较慢）")
                    use_go = False
                else:
                    self.status_label.setText(f"使用 Go 分析日志（{self.go_base_url}）")

            if use_go:
                self.worker_thread = GoLogProcessingThread(
                    file_path,
                    log_type,
                    ai_analysis_starts,
                    optional_parameters_log,
                    server_url=self.go_base_url
                )
            else:
                self.worker_thread = python_log_fallback.PythonLogFallbackThread(
                    file_path,
                    ai_analysis_starts,
                    optional_parameters_log or {"URI_Security_Check": False},
                )

            self.worker_thread.finished.connect(self.on_log_processing_finished)
            self.worker_thread.progress_updated.connect(self.update_log_progress)
            self.worker_thread.error.connect(self.on_log_processing_error)
            self.worker_thread.start()

        except Exception as e:
            error_msg = self.lang_manager.tr("log_analysis_error", "分析日志时出错") + f": {str(e)}"
            self.log_table.setItem(0, 0, QTableWidgetItem(error_msg))
            self._set_analyze_button_running(False)

    def _go_service_reachable(self) -> bool:
        try:
            wait_go_ready(self.go_base_url, timeout=1.5)
            return True
        except Exception:
            return False

    def start_go_backend(self):
        """启动并管理Go分析器子进程（跨平台版）"""
        # 已有可达服务（本进程拉起的，或上次未关干净的孤儿）→ 直接复用
        if self._go_service_reachable():
            if self.go_process and self.go_process.poll() is None:
                return
            # 端口上已有服务但不是我们拉的子进程：仍可复用 HTTP
            self.status_label.setText(f"复用已运行的 Go 服务 ({self.go_base_url})")
            return

        if self.go_process and self.go_process.poll() is None:
            # 进程在但端口不通 → 假死，杀掉重启
            self.stop_go_backend()

        try:
            go_exe_path = find_go_executable()
            if not go_exe_path:
                QMessageBox.critical(
                    self, "错误",
                    "后台服务文件未找到。\n请将编译后的 log_identification(.exe) 放到 lib/ 目录。"
                )
                self.go_process = None
                return

            # 端口被僵尸占用时先清掉
            pid = _pid_listening_on_port(int(self.go_port))
            if pid:
                self.status_label.setText(f"清理占用端口 {self.go_port} 的旧进程 PID={pid}…")
                _kill_pid(pid)
                time.sleep(0.4)

            cmd = [go_exe_path, "-port", str(self.go_port)]
            # 冻结包：工作目录用 exe 旁（有 lib/config/output）；开发态用源码根
            if getattr(sys, "frozen", False):
                work_dir = os.path.dirname(os.path.abspath(sys.executable))
            else:
                work_dir = os.path.dirname(os.path.abspath(__file__))
                if not os.path.isfile(os.path.join(work_dir, "config.yaml")):
                    work_dir = os.getcwd()

            lib_dir = os.path.join(work_dir, "lib")
            env = os.environ.copy()
            if os.path.isdir(lib_dir):
                env["PATH"] = lib_dir + os.pathsep + env.get("PATH", "")

            log_path = os.path.join(get_output_dir(), "go_backend.log")
            os.makedirs(os.path.dirname(log_path), exist_ok=True)
            log_fh = open(log_path, "ab")

            popen_kwargs = {
                "args": cmd,
                "cwd": work_dir,
                "env": env,
                "stdout": log_fh,
                "stderr": subprocess.STDOUT,
            }

            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = subprocess.SW_HIDE
                popen_kwargs["startupinfo"] = startupinfo
                popen_kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
            else:
                try:
                    os.chmod(go_exe_path, 0o755)
                except OSError:
                    pass

            self.go_process = subprocess.Popen(**popen_kwargs)
            self._go_log_fh = log_fh

            self.status_label.setText("正在启动Go分析服务...")
            wait_go_ready(self.go_base_url, timeout=15.0)
            self.status_label.setText(f"Go分析服务已就绪 ({self.go_base_url})")

        except TimeoutError as e:
            err_detail = ""
            try:
                log_path = os.path.join(
                    os.path.dirname(os.path.abspath(__file__)), "output", "go_backend.log"
                )
                if os.path.isfile(log_path):
                    with open(log_path, "rb") as f:
                        err_detail = f.read()[-800:].decode("utf-8", errors="ignore")
            except Exception:
                pass
            msg = f"Go后台服务启动超时:\n{e}"
            if err_detail:
                msg += f"\n\n进程日志:\n{err_detail}"
            QMessageBox.critical(self, "错误", msg)
            self.stop_go_backend()
            self.go_process = None
        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法启动Go后台服务: {e}")
            self.go_process = None

    def stop_go_backend(self):
        """停止Go分析器子进程"""
        if self.go_process and self.go_process.poll() is None:
            try:
                self.go_process.terminate()
                self.go_process.wait(timeout=5)
            except Exception:
                try:
                    self.go_process.kill()
                except Exception:
                    pass
            print("Go后台服务已停止。")
        self.go_process = None
        fh = getattr(self, "_go_log_fh", None)
        if fh:
            try:
                fh.close()
            except Exception:
                pass
            self._go_log_fh = None
    def update_log_progress(self, progress):
        """更新日志分析进度"""
        try:
            p = int(progress)
        except Exception:
            p = 0
        self.progress_bar.setValue(p)
        self.progress_bar.show()
        if p < 100:
            self.status_label.setText(f"日志分析中… {p}%")
        else:
            self.status_label.setText("日志分析完成")

    def log_dynamic_updates(self, url_stats, *_):
        # 远程路径改由 stats_tick 驱动；保留兼容
        if getattr(self, "_remote_streaming", False):
            return
        self.url_stats = url_stats
        if not self.gui_update_timer.isActive():
            self.gui_update_timer.start()

    def _perform_gui_update(self):
        if getattr(self, "_remote_streaming", False):
            return  # 远程由 _on_remote_stats_tick 负责
        if not extract_url_stats_data(self.url_stats):
            return
        try:
            self.update_dashboard_stats()
            self.parse_web_server_log()
            self.update_stats_display()
        except Exception as e:
            print(f"_perform_gui_update error: {e}")

    def show_loading_overlay(self, message=None):
        """显示全屏加载中；先让出事件循环，保证用户能看到动画"""
        if message is None:
            message = self.lang_manager.tr("loading_results", "正在加载分析结果，请稍候...")
        # 清理旧实例
        old = getattr(self, "loading_overlay", None)
        if old is not None:
            try:
                old.hide()
                old.deleteLater()
            except RuntimeError:
                pass
            self.loading_overlay = None

        # 挂到 centralWidget，坐标/尺寸才是内容区，不会偏到主窗口左上角
        host = self.centralWidget() or self
        self.loading_overlay = OverlayLoadingWidget(host, lang_manager=self.lang_manager)
        self.loading_overlay.show(message)
        self.status_label.setText(message)
        QApplication.processEvents()
        # 尺寸稳定后再居中一次
        QTimer.singleShot(0, lambda: self.loading_overlay and self.loading_overlay._cover_parent())

    def hide_loading_overlay(self):
        overlay = getattr(self, "loading_overlay", None)
        if overlay is None:
            return
        try:
            overlay.hide()
            overlay.deleteLater()
        except RuntimeError:
            pass
        self.loading_overlay = None

    def run_with_loading(self, work_fn, message=None, delay_ms=50):
        """
        先显示加载层，再延迟执行耗时 UI 刷新。
        避免「show 后立刻同步重活」导致加载动画根本画不出来。
        """
        self.show_loading_overlay(message)

        def _run():
            try:
                work_fn()
            finally:
                self.hide_loading_overlay()

        QTimer.singleShot(delay_ms, _run)

    def on_log_processing_finished(self, ai_analysis_starts=None, result_file_path=None, *_):
        """日志处理完成回调"""
        try:
            self.progress_bar.setValue(100)
            self.progress_bar.hide()
            if getattr(self, "_analysis_stop_requested", False):
                self._analysis_stop_requested = False
                self._set_analyze_button_running(False)
                return
            self._set_analyze_button_running(False)

            if hasattr(self, 'ui_update_timer') and self.ui_update_timer.isActive():
                self.ui_update_timer.stop()
                self.ui_update_timer.setInterval(5000)

            def _load_results():
                try:
                    if not result_file_path or not os.path.exists(result_file_path):
                        self.status_label.setText("分析完成，但结果文件不存在。")
                        return

                    if self.loading_overlay:
                        self.loading_overlay.set_message(
                            self.lang_manager.tr("loading_results", "正在加载分析结果，请稍候...")
                        )

                    with open(result_file_path, 'r', encoding='utf-8') as f:
                        final_stats = json.load(f)

                    self.url_stats = final_stats
                    if not isinstance(self.url_stats, dict) or self.url_stats.get('data') is None:
                        self.status_label.setText("日志结果 JSON 格式不正确，缺少 'data' 键。")
                        return

                    self._remember_analysis_artifacts(
                        stats_json=result_file_path or "",
                        source_file=self.Import_box.text().strip(),
                    )

                    if self.loading_overlay:
                        self.loading_overlay.set_message(
                            self.lang_manager.tr("rendering_stats", "正在渲染统计界面...")
                        )

                    # 分步刷新，单步失败不拖垮整个界面
                    for step_name, step_fn in (
                        ("update_stats_display", self.update_stats_display),
                        ("parse_web_server_log", self.parse_web_server_log),
                        ("update_dashboard_stats", self.update_dashboard_stats),
                        ("_fill_log_table_from_stats", lambda: self._fill_log_table_from_stats(self.url_stats)),
                    ):
                        try:
                            step_fn()
                        except Exception as step_e:
                            print(f"[warn] {step_name} failed: {step_e}")

                    if (
                        ai_analysis_starts
                        or getattr(self, "_pending_agent_review", False)
                        or (hasattr(self, "ai_auto_analyze_check") and self.ai_auto_analyze_check.isChecked())
                    ):
                        try:
                            self.maybe_start_pending_agent()
                        except Exception as agent_e:
                            print(f"[warn] maybe_start_pending_agent failed: {agent_e}")

                    self.add_recent_activity(
                        self.lang_manager.tr("log_analysis_completed", "LOG分析完成"),
                        self.Import_box.text().strip(),
                        self.lang_manager.tr("log_analysis_summary", "分析成功完成"),
                        result_file_path=result_file_path
                    )
                    self.status_label.setText(self.lang_manager.tr("log_analysis_completed", "LOG分析完成"))
                except Exception as e:
                    self.status_label.setText(f"加载分析结果失败: {e}")
                    QMessageBox.warning(self, self.lang_manager.tr("error", "错误"), f"加载分析结果失败:\n{e}")

            self.run_with_loading(
                _load_results,
                message=self.lang_manager.tr("loading_results", "正在加载分析结果，请稍候...")
            )
        except Exception as e:
            self._set_analyze_button_running(False)
            self.progress_bar.hide()
            self.hide_loading_overlay()
            QMessageBox.warning(self, self.lang_manager.tr("error", "错误"), f"日志完成回调异常:\n{e}")

    def on_log_processing_error(self, error_msg, *_):
        """日志处理错误回调"""
        if hasattr(self, 'ui_update_timer') and self.ui_update_timer.isActive():
            self.ui_update_timer.stop()
        self.hide_loading_overlay()
        self.progress_bar.hide()
        try:
            if self.log_table.rowCount() < 1:
                self.log_table.setRowCount(1)
            self.log_table.setItem(0, 0, QTableWidgetItem(
                self.lang_manager.tr("log_analysis_error", "分析日志时出错: ") + error_msg
            ))
        except Exception:
            pass
        self._set_analyze_button_running(False)
        self.status_label.setText(self.lang_manager.tr("log_analysis_error_status", "日志分析出错")+error_msg)

    def parse_web_server_log(self, *_):
        """解析Web服务器日志并填充表格"""
        try:
            data = extract_url_stats_data(self.url_stats)
            sorted_stats = sorted(
                ((u, s) for u, s in data.items() if u != '_global_stats'),
                key=lambda item: item[1].get('count', 0),
                reverse=True
            )
            max_rows = 2000
            display_stats = sorted_stats[:max_rows]

            self.log_table.setSortingEnabled(False)
            self.log_table.setUpdatesEnabled(False)
            try:
                self.log_table.setRowCount(len(display_stats))
                for row, (url, stats) in enumerate(display_stats):
                    self.log_table.setItem(row, 0, QTableWidgetItem(unquote(url)))
                    self.log_table.setItem(row, 1, QTableWidgetItem(str(stats.get('count', 0))))

                    status_codes = defaultdict(int)
                    methods = defaultdict(int)
                    uas = defaultdict(int)
                    ip_items = []
                    for ip, ip_data in list(stats.get('source_ips', {}).items())[:30]:
                        ip_parts = ip.split("：")
                        ip_addr = ip_parts[0]
                        ip_location = ip_parts[1] if len(ip_parts) > 1 else "未知位置"
                        ip_items.append(f"{ip_addr} ({ip_location}): {ip_data.get('count', 0)}次")
                        for code, count in ip_data.get('status_codes', {}).items():
                            status_codes[code] += count
                        for method, count in ip_data.get('methods', {}).items():
                            methods[method] += count
                        for ua, count in list(ip_data.get('UA', {}).items())[:5]:
                            simple_ua = ua.split('(')[0].strip() or (ua[:30] + '...')
                            uas[simple_ua] += count

                    self.log_table.setItem(row, 2, QTableWidgetItem("\n".join(f"{k}: {v}次" for k, v in status_codes.items())))
                    self.log_table.setItem(row, 3, QTableWidgetItem("\n".join(ip_items[:10])))
                    self.log_table.setItem(row, 4, QTableWidgetItem("\n".join(f"{k}: {v}次" for k, v in methods.items())))
                    self.log_table.setItem(row, 5, QTableWidgetItem("\n".join(f"{k}: {v}次" for k, v in list(uas.items())[:8])))

                self.log_table.setColumnWidth(0, 350)
                self.log_table.setColumnWidth(3, 250)
                self.log_table.setColumnWidth(5, 200)
            finally:
                self.log_table.setUpdatesEnabled(True)
                self.log_table.setSortingEnabled(True)

            if len(sorted_stats) > max_rows:
                self.status_label.setText(
                    self.lang_manager.tr(
                        "log_table_truncated",
                        f"日志结果表仅显示 Top {max_rows}/{len(sorted_stats)} 条 URL"
                    )
                )
        except Exception as e:
            self.log_table.setUpdatesEnabled(True)
            self.log_table.setRowCount(1)
            self.log_table.setItem(0, 0, QTableWidgetItem(f"解析日志时出错: {str(e)}"))

    def export_log_table(self, *_):
        """导出日志表格数据为CSV文件"""
        # 检查是否有数据可导出
        if self.log_table.rowCount() == 0:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "当前没有可导出的日志数据！")
            self.add_recent_activity("导出日志", "", "失败：无数据")
            return

        # 获取保存文件路径
        default_name = f"web_log_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出日志数据",
            default_name,
            "CSV文件 (*.csv);;所有文件 (*)"
        )

        if not file_path:  # 用户取消操作
            self.add_recent_activity("导出日志", "", "取消")
            return

        try:
            with open(file_path, 'w', newline='', encoding='utf-8-sig') as f:  # 使用utf-8-sig支持Excel中文
                writer = csv.writer(f)

                # 写入表头（使用中文列名）
                headers = [
                    "URL地址",
                    "总访问次数",
                    "状态码分布",
                    "来源IP及位置",
                    "请求方法",
                    "安全威胁",
                    "用户代理"
                ]
                writer.writerow(headers)

                # 写入表格数据
                for row in range(self.log_table.rowCount()):
                    row_data = []
                    for col in range(self.log_table.columnCount()):
                        item = self.log_table.item(row, col)
                        # 处理换行符，替换为分号以便在CSV中显示
                        text = item.text().replace('\n', '; ') if item else ""
                        row_data.append(text)
                    writer.writerow(row_data)

            # 导出成功反馈
            QMessageBox.information(self, "导出成功", f"日志数据已成功导出到：\n{file_path}")
            self.status_label.setText(f"日志数据已导出: {os.path.basename(file_path)}")
            self.add_recent_activity("导出日志", file_path, self.lang_manager.tr("done", "完成"))

        except PermissionError:
            error_msg = "没有写入权限，请选择其他位置保存"
            QMessageBox.critical(self, "导出失败", error_msg)
            self.add_recent_activity("导出日志", file_path, f"失败：{error_msg}")
        except Exception as e:
            error_msg = f"导出过程中发生错误: {str(e)}"
            QMessageBox.critical(self, "导出失败", error_msg)
            self.add_recent_activity("导出日志", file_path, f"失败：{str(e)}")

    def open_output_folder(self, *_):
        """打开分析结果目录（exe 旁 / 项目根下的 output）。"""
        path = get_output_dir()
        self.output_dir = path
        if hasattr(self, "output_path_label"):
            self.output_path_label.setText(f"结果: {path}")
            self.output_path_label.setToolTip(path)
        ok = QDesktopServices.openUrl(QUrl.fromLocalFile(path))
        if not ok:
            QMessageBox.information(self, "结果目录", path)

    def create_status_bar(self, *_):
        """创建状态栏"""
        status_bar = QStatusBar()
        self.setStatusBar(status_bar)

        # 添加状态标签
        self.status_label = QLabel(self.lang_manager.tr("ready", "就绪"))
        status_bar.addWidget(self.status_label, stretch=1)

        self.output_path_label = QLabel(f"结果: {get_output_dir()}")
        self.output_path_label.setStyleSheet("color:#666;")
        self.output_path_label.setToolTip("点击打开分析结果目录")
        self.output_path_label.setCursor(Qt.CursorShape.PointingHandCursor)
        self.output_path_label.mousePressEvent = lambda *_: self.open_output_folder()
        status_bar.addWidget(self.output_path_label)

        # 添加进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setFixedWidth(200)  # 设置固定宽度

        status_bar.addPermanentWidget(self.progress_bar)
        self.progress_bar.hide()

        # 添加版本信息
        version_label = QLabel("TrafficEye Web v" + version)
        status_bar.addPermanentWidget(version_label)

        # 添加系统时间
        self.time_label = QLabel()
        self.update_time()
        status_bar.addPermanentWidget(self.time_label)

        # 创建定时器更新时间
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_time)
        self.timer.start(1000)

    def update_time(self, *_):
        """更新时间显示"""
        # current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S") # 添加显示时间
        current_time = ""
        self.time_label.setText(current_time)

    def show_about_dialog(self, *_):
        """显示关于对话框"""
        about_box = QMessageBox(self)
        about_box.setWindowTitle("关于 TrafficEye Wbe")
        about_box.setIconPixmap(QIcon("ico/l.png").pixmap(64, 64))
        about_box.setText(f"""
            <h2>TrafficEye</h2>
            <p>{self.lang_manager.tr("version", "版本")}: {version}</p>
            <p>{self.lang_manager.tr("Introduction_name", "网络流量分析工具")}</p>
            <p><a href="https://github.com/CuriousLearnerDev/TrafficEye">访问我们的网站</a></p>
        """)
        about_box.exec()

    def open_documentation(self, *_):
        """打开文档"""
        webbrowser.open("https://github.com/CuriousLearnerDev/TrafficEye")

    def select_file(self, *_):
        """选择流量文件"""
        file, _ = QFileDialog.getOpenFileName(
            self,
            self.lang_manager.tr("select_traffic_file", "选择流量文件"),
            "",
            "PCAP/Log/CSV (*.pcap *.pcapng *.cap *.log *.txt *.csv *.tsv);;CSV (*.csv *.tsv);;All Files (*)"
        )
        if file:
            self.Import_box.setText(file)
            self.status_label.setText(
                f"{self.lang_manager.tr('file_selected', '已选择文件')}: {file}"
            )
            self.add_recent_activity(
                self.lang_manager.tr("open_file", "打开文件"),
                file,
                self.lang_manager.tr("open_file", "打开文件")
            )

    def load_recent_activity(self, *_):
        """ 读取之前的操作 """
        path = "history/trafficeye_data.json"
        if not os.path.exists(path):
            return

        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError:
            print(f"警告: 无法解码 {path}。历史记录可能已损坏。")
            return  # 避免在空文件或损坏文件上崩溃
        except FileNotFoundError:
            print(f"历史文件 {path} 不存在。")
            return

        recent = data.get("recent", [])
        self.recent_table.setRowCount(len(recent))

        for row, entry in enumerate(recent):
            if not entry: continue  # 跳过空条目

            # 1. 填充表格的4列 (使用切片以确保安全)
            for col, text in enumerate(entry[:4]):
                self.recent_table.setItem(row, col, QTableWidgetItem(str(text)))

            # 2. 检查并存储第5个元素 (result_file_path)
            if len(entry) > 4 and entry[4]:
                result_file_path = entry[4]
                first_item = self.recent_table.item(row, 0)
                if first_item:
                    # 3. 存储路径
                    first_item.setData(Qt.ItemDataRole.UserRole, result_file_path)

                    # 4. 创建按钮
                    open_btn = QPushButton(QIcon("ico/open.png"), "打开结果")  # 仅显示图标
                    open_btn.setToolTip(self.lang_manager.tr("open_result", "打开结果"))
                    open_btn.setStyleSheet("padding: 2px; border: 1px solid #555; border-radius: 3px;")

                    # 5. 连接信号
                    open_btn.clicked.connect(lambda checked, r=row: self.open_recent_result_by_row(r))

                    # 6. 将按钮设置到第5列
                    self.recent_table.setCellWidget(row, 4, open_btn)

                    # 7. (可选) 添加整行提示
                    tooltip = self.lang_manager.tr("clickable_result_tooltip", "点击此行或按钮查看分析结果")
                    for col in range(4):
                        item = self.recent_table.item(row, col)
                        if item:
                            item.setToolTip(tooltip)

        # 历史从文件载入后默认看最新（底部），避免记录多时还要手动下滑
        self._scroll_recent_to_bottom()
        QTimer.singleShot(0, self._scroll_recent_to_bottom)
        QTimer.singleShot(100, self._scroll_recent_to_bottom)

    def _scroll_recent_to_bottom(self):
        """最近操作日志滚动到最底部（最新一条）"""
        table = getattr(self, "recent_table", None)
        if table is None:
            return
        last = table.rowCount() - 1
        if last < 0:
            return
        table.scrollToBottom()
        item = table.item(last, 0)
        if item is not None:
            table.scrollToItem(item, QAbstractItemView.ScrollHint.PositionAtBottom)
        # 同步滚动条位置，确保视口真在底部
        bar = table.verticalScrollBar()
        if bar is not None:
            bar.setValue(bar.maximum())

    def add_recent_activity(self, action, filename, status, status_color: QColor = None, result_file_path: str = None,*_):
        """添加最近活动记录，并可选指定状态颜色和结果文件路径"""
        row = self.recent_table.rowCount()
        self.recent_table.insertRow(row)

        time_item = QTableWidgetItem(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        action_item = QTableWidgetItem(action)
        file_item = QTableWidgetItem(filename)
        status_item = QTableWidgetItem(status)

        time_item, action_item, file_item, status_item = session_utils.operational_record_keeping(
            time_item=time_item,
            action_item=action_item,
            file_item=file_item,
            status_item=status_item,
            result_file_path=result_file_path  # <-- 传递新参数
        )

        # 如果调用者传了颜色，就用传入的，否则根据内容判断
        if status_color is None:
            status_lower = status.lower()
            if ("成功" in status_lower) or ("success" in status_lower):
                status_color = QColor("#86efac")  # 柔和绿色
            elif any(x in status_lower for x in ["运行中", "分析中", "进行中", "running", "analyzing", "in progress"]):
                status_color = QColor("#fde68a")  # 柔和黄色
            elif ("失败" in status_lower) or ("fail" in status_lower):
                status_color = QColor("#fca5a5")  # 柔和红色
            else:
                status_color = QColor("#e5e7eb")  # 柔和灰白（避免纯白太亮）

        status_item.setForeground(status_color)

        self.recent_table.setItem(row, 0, time_item)
        self.recent_table.setItem(row, 1, action_item)
        self.recent_table.setItem(row, 2, file_item)
        self.recent_table.setItem(row, 3, status_item)

        # --- 修改: 将路径存储在UserRole中，并添加按钮 ---
        if result_file_path:
            # 1. 存储路径
            time_item.setData(Qt.ItemDataRole.UserRole, result_file_path)

            # 2. 创建按钮
            open_btn = QPushButton(QIcon("ico/open.png"), "打开结果")  # 仅显示图标
            open_btn.setToolTip(self.lang_manager.tr("open_result", "打开结果"))
            open_btn.setStyleSheet("padding: 2px; border: 1px solid #555; border-radius: 3px;")

            # 3. 连接信号 (使用 lambda 捕获当前行号 row)
            open_btn.clicked.connect(lambda checked, r=row: self.open_recent_result_by_row(r))

            # 4. 将按钮设置到第5列 (索引为4)
            self.recent_table.setCellWidget(row, 4, open_btn)

            # 5. (可选) 添加整行提示
            tooltip = self.lang_manager.tr("clickable_result_tooltip", "点击此行或按钮查看分析结果")
            time_item.setToolTip(tooltip)
            action_item.setToolTip(tooltip)
            file_item.setToolTip(f"{filename}\n{tooltip}")
            status_item.setToolTip(tooltip)

        self._scroll_recent_to_bottom()
        QTimer.singleShot(0, self._scroll_recent_to_bottom)

    def get_proxy_settings(self, *_):
        """获取代理设置"""
        if not self.proxy_group.isChecked():
            return None

        proxies = {}
        http_proxy = self.http_proxy_input.text().strip()
        https_proxy = self.https_proxy_input.text().strip()

        if http_proxy:
            proxies['http'] = http_proxy
        if https_proxy:
            proxies['https'] = https_proxy

        return proxies if proxies else None

    # 应用到耗时方法
    # 在 MainWindow 类中替换 start_analysis 方法
    def start_analysis(self, ai_analysis_starts=None, optional_parameters=None, *_):
        """开始流量分析 (调用 Go 后端)"""
        # 独立会话：清空上次全流量/LOG/表格残留
        self._reset_analysis_session()
        self._analysis_stop_requested = False

        # 显示并重置进度条
        self.progress_bar.show()
        self.progress_bar.setValue(0)

        file = self.Import_box.text()
        if not file:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("Please_select_the_traffic_file_first", "请先选择流量文件!"))
            self.progress_bar.hide()
            self._set_analyze_button_running(False)
            return

        self.last_analysis_source = file
        self._update_risk_source_label()

        # 确保后台服务已启动（可复用已在端口上的 Go，不必非本进程子进程）
        self.start_go_backend()
        if not self.go_process and not self._go_service_reachable():
            self.status_label.setText("Go后台服务未启动，无法分析。")
            self.progress_bar.hide()
            self._set_analyze_button_running(False)
            return

        self._set_analyze_button_running(True)

        # 准备参数
        if optional_parameters is None:
            # 如果没有传参（例如直接点击开始分析而非通过对话框），设置默认参数
            optional_parameters = {
                "URL_Security_Check": False,
                "Request_Head_Security_Check": False,
                "Data_section_detection": {
                    "enabled": False,
                    "binary": False,
                    "forms": False,
                    "json": False,
                    "xml": False,
                    "multipart": False
                }
            }

        self.status_label.setText(self.lang_manager.tr("analysis_in_progress", "全流量分析进行中..."))

        # 使用新的 Go 处理线程
        self.analysis_thread = GoPcapProcessingThread(
            file_path=file,
            ai_analysis_starts=ai_analysis_starts,
            optional_parameters=optional_parameters,
            server_url=self.go_base_url
        )

        self.analysis_thread.progress_updated.connect(self.update_log_progress)  # 复用进度条更新
        self.analysis_thread.status_msg.connect(self.update_status_label)
        self.analysis_thread.finished.connect(self.on_pcap_processing_finished)
        self.analysis_thread.error.connect(self.on_pcap_processing_error)

        self.analysis_thread.start()

        self.add_recent_activity(self.lang_manager.tr("start_analysis", "开始分析"), file,
                                 self.lang_manager.tr("In_progress", "进行中"))

    def on_pcap_processing_error(self, error_msg):
        """全流量分析错误回调"""
        self.hide_loading_overlay()
        self.progress_bar.hide()
        self._set_analyze_button_running(False)
        QMessageBox.critical(self, self.lang_manager.tr("error", "错误"), error_msg)
        self.status_label.setText(f"分析失败: {error_msg}")
        self.analysis_text_edit.setPlainText(f"分析过程中发生错误:\n{error_msg}")

        self.add_recent_activity("流量分析失败", self.Import_box.text(), "失败")
    def on_pcap_processing_finished(self, ai_analysis_starts, result_json_path, full_traffic_path, dns_result_path=""):
        """Go 全流量分析完成回调"""
        if getattr(self, "_analysis_stop_requested", False):
            self._analysis_stop_requested = False
            self._set_analyze_button_running(False)
            self.progress_bar.hide()
            return
        self._set_analyze_button_running(False)
        self.progress_bar.hide()

        def _load_results():
            try:
                if self.loading_overlay:
                    self.loading_overlay.set_message(
                        self.lang_manager.tr("loading_results", "正在加载分析结果，请稍候...")
                    )

                with open(result_json_path, 'r', encoding='utf-8') as f:
                    self.url_stats = json.load(f)

                self._remember_analysis_artifacts(
                    stats_json=result_json_path or "",
                    traffic_file=full_traffic_path or "",
                    source_file=self.Import_box.text().strip(),
                )

                if self.loading_overlay:
                    self.loading_overlay.set_message(
                        self.lang_manager.tr("rendering_stats", "正在渲染统计界面...")
                    )

                self.update_stats_display()
                self.update_dashboard_stats()

                try:
                    self.load_dns_results(
                        dns_path=dns_result_path or "",
                        pcap_path=self.Import_box.text().strip(),
                        stats_json=result_json_path or "",
                        persist=True,
                    )
                except Exception as _dns_e:
                    print(f"DNS load error: {_dns_e}")

                max_lines = 1000
                preview_content = ""
                line_count = 0
                is_truncated = False

                if full_traffic_path and os.path.exists(full_traffic_path):
                    self.start_analysis_timestamp = full_traffic_path

                    with open(full_traffic_path, 'r', encoding='utf-8', errors='ignore') as f:
                        for line in f:
                            if line_count < max_lines:
                                preview_content += line
                                line_count += 1
                            else:
                                is_truncated = True
                                break

                    self.analysis_text_edit.setPlainText(preview_content)

                    if is_truncated:
                        abs_path = os.path.abspath(full_traffic_path)
                        file_url = QUrl.fromLocalFile(abs_path).toString()

                        html = (
                            f"<span style='color:#E65100;font-size:11px;line-height:14px;'>"
                            f"[{self.lang_manager.tr('tip', '提示')}] "
                            f"输出超过 {max_lines} 行，完整内容已保存 · "
                            f"<a href=\"{file_url}\" style='color:#1565C0;'>"
                            f"{self.lang_manager.tr('click_to_open', '打开完整文件')}</a></span>"
                        )
                        self.status_create_analysis_tab.setTextFormat(Qt.TextFormat.RichText)
                        self.status_create_analysis_tab.setTextInteractionFlags(
                            Qt.TextInteractionFlag.TextBrowserInteraction)
                        self.status_create_analysis_tab.setOpenExternalLinks(True)
                        self.status_create_analysis_tab.setText(html)
                    else:
                        self.status_create_analysis_tab.setText(f"完整结果已加载 ({line_count} 行)")
                else:
                    # 再尝试按 JSON 文件名配对
                    fallback = self._resolve_pcap_traffic_path(result_json_path or "")
                    if fallback and os.path.exists(fallback):
                        self._remember_analysis_artifacts(traffic_file=fallback)
                        self._load_traffic_text_to_ui(fallback)
                        self.status_create_analysis_tab.setText(
                            f"已从配对文件加载明文: {os.path.basename(fallback)}"
                        )
                    else:
                        self.analysis_text_edit.setPlainText(
                            "分析完成，但未找到流量明文文件。\n"
                            "若包内确有 HTTP，请确认已更新分析引擎后重试。"
                        )

                if (
                    ai_analysis_starts
                    or getattr(self, "_pending_agent_review", False)
                    or (hasattr(self, "ai_auto_analyze_check") and self.ai_auto_analyze_check.isChecked())
                ):
                    self.maybe_start_pending_agent()

                summary = (
                    f"分析完成。请求数：{self.url_stats.get('_global_stats', {}).get('request_total', 0)}，"
                    f"威胁：{self.url_stats.get('_global_stats', {}).get('danger_total', 0)}"
                )
                self.add_recent_activity(
                    self.lang_manager.tr("analysis_complete", "流量分析完成"),
                    self.Import_box.text(),
                    summary,
                    result_file_path=result_json_path
                )
                self.status_label.setText(self.lang_manager.tr("done", "分析完成"))
            except Exception as e:
                QMessageBox.warning(self, "结果处理错误", f"处理分析结果时出错: {e}")
                print(e)

        self.run_with_loading(
            _load_results,
            message=self.lang_manager.tr("loading_results", "正在加载分析结果，请稍候...")
        )
    def update_progress(self, value, *_):
        """更新进度条"""
        if isinstance(value, str):
            # 如果是字符串状态更新
            self.status_label.setText(value)
        elif isinstance(value, int):
            # 如果是数字进度更新
            self.progress_bar.setValue(value)

    def stop_analysis(self, *_):
        """结束当前分析（全流量 / LOG）。"""
        self._analysis_stop_requested = True
        stopped = False
        for name in ("analysis_thread", "worker_thread"):
            th = getattr(self, name, None)
            if th is None:
                continue
            try:
                running = th.isRunning()
            except Exception:
                running = False
            if not running:
                continue
            stopped = True
            try:
                if hasattr(th, "stop") and callable(th.stop):
                    th.stop()
                else:
                    th.requestInterruption()
            except Exception:
                pass
            # 轮询线程会很快退出；超时再强制终止
            if not th.wait(2500):
                try:
                    th.terminate()
                    th.wait(1000)
                except Exception:
                    pass
            try:
                setattr(self, name, None)
            except Exception:
                pass

        self.hide_loading_overlay()
        self.progress_bar.hide()
        self._set_analyze_button_running(False)
        msg = self.lang_manager.tr("end_analysis", "结束分析")
        self.status_label.setText(msg)
        if hasattr(self, "analysis_text_edit") and self.analysis_text_edit is not None:
            try:
                self.analysis_text_edit.append(msg)
            except Exception:
                pass
        if stopped:
            self.add_recent_activity(msg, self.Import_box.text(), msg)
        else:
            # 对话框阶段或已结束：仍恢复按钮文案
            self.status_label.setText(self.lang_manager.tr("canceled", "已取消"))

    def export_results(self, *_):
        """导出分析结果"""
        if not self.analysis_text_edit.toPlainText():
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("no_exportable_results", "没有可导出的结果!"))
            return

        # 生成默认文件名
        default_name = "analysis_result_"
        if self.Import_box.text():
            import os
            base_name = os.path.splitext(os.path.basename(self.Import_box.text()))[0]
            default_name = f"{base_name}_analysis_"
        default_name += datetime.datetime.now().strftime("%Y%m%d_%H%M%S") + ".txt"

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            self.lang_manager.tr("export_analysis_results", "导出分析结果"),
            default_name,
            "Text Files (*.txt);;All Files (*)"
        )

        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.analysis_text_edit.toPlainText())
                self.status_label.setText(
                    self.lang_manager.tr("export_success", "结果已成功导出到: ") + file_path
                )
                self.add_recent_activity(
                    self.lang_manager.tr("export_results", "导出结果"),
                    file_path,
                    self.lang_manager.tr("done", "完成")
                )
                QMessageBox.information(
                    self,
                    self.lang_manager.tr("done", "完成"),
                    self.lang_manager.tr("export_success_msg", "结果已成功导出!")
                )
            except Exception as e:
                QMessageBox.critical(
                    self,
                    self.lang_manager.tr("error", "错误"),
                    self.lang_manager.tr("export_failed", "导出失败: ") + str(e)
                )
                self.add_recent_activity(
                    self.lang_manager.tr("export_results", "导出结果"),
                    file_path,
                    self.lang_manager.tr("failed", "失败") + f": {str(e)}"
                )

    def update_results(self, text, *_):
        """更新分析结果"""
        # try:
        #     max_lines = int(self.max_lines_input.text()) if self.max_lines_input.text() else 1000
        # except ValueError:
        #     max_lines = 1000
        max_lines = 1000
        line_count = self.analysis_text_edit.document().blockCount()

        if line_count < max_lines:
            self.analysis_text_edit.append(text)
        else:
            if not getattr(self, "memory_optimization", False):
                self.memory_optimization = True
                full_content = self.analysis_text_edit.toPlainText()

                # --- 保存全文 ---
                output_dir = get_output_dir()
                self.output_dir = output_dir

                # 确保路径唯一
                self.start_analysis_timestamp = os.path.join(output_dir, self.start_analysis_timestamp)
                self.file_writer.add_content(self.start_analysis_timestamp, full_content)

                # --- 构造可点击提示 ---
                file_url = QUrl.fromLocalFile(self.start_analysis_timestamp).toString()  # ⭐
                html = (
                    f"<span style='color:#E65100;font-size:11px;line-height:14px;'>"
                    f"[{self.lang_manager.tr('tip', '提示')}] "
                    f"已超过 {max_lines} 行，后续仅写入文件 · "
                    f"<a href=\"{file_url}\" style='color:#1565C0;'>"
                    f"{self.lang_manager.tr('click_to_open', '打开文件')}</a></span>"
                )

                lbl: QLabel = self.status_create_analysis_tab
                lbl.setTextFormat(Qt.TextFormat.RichText)  # ⭐
                lbl.setTextInteractionFlags(Qt.TextInteractionFlag.TextBrowserInteraction)  # ⭐
                lbl.setOpenExternalLinks(True)  # ⭐
                lbl.setText(html)
            # 继续把新增行拼到文件里
            self.file_writer.add_content(self.start_analysis_timestamp, text)

    # 在AnalysisThread类的analysis_finished方法中添加自动分析逻辑
    def analysis_finished(self, ai_analysis_storing_data=None, *_):
        """分析完成"""
        self._set_analyze_button_running(False)
        self.analysis_text_edit.append(self.lang_manager.tr("done", "完成"))
        self.progress_bar.hide()

        def _load_results():
            result_file_to_save = None
            if hasattr(self, 'start_analysis_timestamp'):
                result_file_to_save = self.start_analysis_timestamp

                if not getattr(self, "memory_optimization", False):
                    try:
                        full_content = self.analysis_text_edit.toPlainText()
                        with open(result_file_to_save, 'w', encoding='utf-8') as f:
                            f.write(full_content)
                    except Exception as e:
                        print(f"无法保存小型分析结果: {e}")
                        result_file_to_save = None

            if hasattr(self.analysis_thread, 'last_result'):
                self.url_stats = self.analysis_thread.last_result
                if self.loading_overlay:
                    self.loading_overlay.set_message(
                        self.lang_manager.tr("rendering_stats", "正在渲染统计界面...")
                    )
                self.add_recent_activity(
                    self.lang_manager.tr("analysis_complete", "分析完成"),
                    self.Import_box.text(),
                    f"分析成功完成一共：{self.url_stats['_global_stats']['request_total']}请求，唯一URI：{self.url_stats['_global_stats']['total_unique_uris']}，危险数量：{self.url_stats['_global_stats']['danger_total']}",
                    result_file_path=result_file_to_save
                )
                self.update_stats_display()
                self.update_dashboard_stats()
            else:
                self.add_recent_activity(
                    self.lang_manager.tr("analysis_complete", "分析完成"),
                    self.Import_box.text(),
                    f"分析成功完成一共：{self.url_stats['_global_stats']['total_unique_uris']}",
                    result_file_path=result_file_to_save
                )

            self.status_label.setText(self.lang_manager.tr("done", "完成"))

            if (
                ai_analysis_storing_data
                or getattr(self, "_pending_agent_review", False)
                or (hasattr(self, "ai_auto_analyze_check") and self.ai_auto_analyze_check.isChecked())
            ):
                self.maybe_start_pending_agent()

        self.run_with_loading(
            _load_results,
            message=self.lang_manager.tr("loading_results", "正在加载分析结果，请稍候...")
        )

    def update_stats_display(self, *_):
        """更新统计信息显示"""
        # 更新风险显示
        self.update_risk_display()

        data = extract_url_stats_data(self.url_stats)
        if not data:
            return

        self.url_tree.setUpdatesEnabled(False)
        try:
            # 清空旧图表
            self.status_chart_view.setChart(QChart())
            self.ip_chart_view.setChart(QChart())
            self.uri_chart_view.setChart(QChart())
            self.time_chart_view.setChart(QChart())

            # 创建并设置模型（限制 Top N，避免卡顿）
            self.tree_model = StatsTreeModel(data, lang_manager=self.lang_manager)
            self.url_tree.setModel(self.tree_model)

            if hasattr(self, 'stats_hint_label'):
                if self.tree_model.total_urls > self.tree_model.displayed_urls:
                    self.stats_hint_label.setText(
                        self.lang_manager.tr("stats_truncated_hint", "显示 Top {0}/{1}（按访问量）").format(
                            self.tree_model.displayed_urls, self.tree_model.total_urls
                        )
                    )
                else:
                    self.stats_hint_label.setText(
                        self.lang_manager.tr("stats_total_hint", "共 {0} 条 URL").format(
                            self.tree_model.total_urls
                        )
                    )

            if not hasattr(self, 'flag_delegate'):
                self.flag_delegate = CountryFlagDelegate()

            self.url_tree.setItemDelegateForColumn(2, self.flag_delegate)

            self.url_tree.setColumnWidth(0, 360)   # URL
            self.url_tree.setColumnWidth(1, 90)    # 次数
            self.url_tree.setColumnWidth(2, 220)   # IP/归属
            self.url_tree.setColumnWidth(3, 140)   # 状态码
            self.url_tree.setColumnWidth(4, 110)   # 方法
            # UA 列随最后一列拉伸

            # 不 expandAll：子节点（IP）按需展开
            self.url_tree.collapseAll()
        finally:
            self.url_tree.setUpdatesEnabled(True)

        # 更新状态码饼图
        status_counts = defaultdict(int)
        for url_info in data.values():
            if 'source_ips' in url_info:
                for ip_data in url_info['source_ips'].values():
                    if 'status_codes' in ip_data:
                        for status_code, count in ip_data['status_codes'].items():
                            status_counts[status_code] += int(count)

        if status_counts:
            status_series = QPieSeries()
            status_series.hovered.connect(self.on_pie_hover)
            for code, count in status_counts.items():
                slice = status_series.append(f"HTTP {code}", count)
                slice.setLabel(f"HTTP {code}: {count}次")

            status_chart = QChart()
            status_chart.addSeries(status_series)
            status_chart.setTitle(self.lang_manager.tr("status_distribution", "状态码分布"))
            status_chart.legend().setVisible(True)
            status_chart.setAnimationOptions(QChart.AnimationOption.NoAnimation)
            self.status_chart_view.setChart(status_chart)

        # 更新IP统计柱状图
        ip_counts = defaultdict(int)
        time_data = defaultdict(int)

        for stats in data.values():
            if 'source_ips' in stats:
                for ip, ip_data in stats['source_ips'].items():
                    ip_counts[ip] += ip_data.get('count', 0)
                    if 'request_time' in ip_data:
                        for time_str, count in ip_data['request_time'].items():
                            time_data[time_str] += count

        if ip_counts:
            ip_series = QBarSeries()
            ip_series.setLabelsVisible(True)
            ip_series.hovered.connect(self.on_ip_bar_hover)

            ip_set = QBarSet(self.lang_manager.tr("request_count", "请求次数"))
            self.ip_categories = []

            for ip, count in sorted(ip_counts.items(), key=lambda x: x[1], reverse=True)[:30]:
                ip_set.append(count)
                self.ip_categories.append(ip.split("：")[0])

            ip_series.append(ip_set)
            ip_chart = QChart()
            ip_chart.addSeries(ip_series)
            ip_chart.setTitle(self.lang_manager.tr("source_ip_stats", "来源IP统计") + f" (Top 30)")
            ip_chart.setAnimationOptions(QChart.AnimationOption.NoAnimation)

            axis_x = QBarCategoryAxis()
            axis_x.append(self.ip_categories)
            ip_chart.addAxis(axis_x, Qt.AlignmentFlag.AlignBottom)
            ip_series.attachAxis(axis_x)

            axis_y = QValueAxis()
            ip_chart.addAxis(axis_y, Qt.AlignmentFlag.AlignLeft)
            ip_series.attachAxis(axis_y)

            self.ip_chart_view.setChart(ip_chart)

        # 更新URL访问统计柱状图（Top 30）
        url_counts = defaultdict(int)
        for url, stats in data.items():
            if url == '_global_stats':
                continue
            url_counts[urlparse(url).path] += stats.get('count', 0)

        if url_counts:
            url_series = QBarSeries()
            url_series.setLabelsVisible(True)
            url_series.hovered.connect(self.on_url_bar_hover)

            url_set = QBarSet(self.lang_manager.tr("visit_count", "访问次数"))
            self.url_categories = []

            for url, count in sorted(url_counts.items(), key=lambda x: x[1], reverse=True)[:30]:
                url_set.append(count)
                display_url = url[:30] + "..." if len(url) > 30 else url
                self.url_categories.append(display_url)

            url_series.append(url_set)
            url_chart = QChart()
            url_chart.addSeries(url_series)
            url_chart.setTitle(self.lang_manager.tr("url_access_stats", "URL访问统计") + "(Top 30)")
            url_chart.setAnimationOptions(QChart.AnimationOption.NoAnimation)

            url_axis_x = QBarCategoryAxis()
            url_axis_x.append(self.url_categories)
            url_chart.addAxis(url_axis_x, Qt.AlignmentFlag.AlignBottom)
            url_series.attachAxis(url_axis_x)

            url_axis_y = QValueAxis()
            url_chart.addAxis(url_axis_y, Qt.AlignmentFlag.AlignLeft)
            url_series.attachAxis(url_axis_y)

            self.uri_chart_view.setChart(url_chart)

        # 更新访问时间趋势图（时间点过多时抽样，避免卡顿）
        if time_data:
            sorted_times = sorted(time_data.items(), key=lambda x: x[0])
            max_points = 200
            if len(sorted_times) > max_points:
                step = max(1, len(sorted_times) // max_points)
                sampled = sorted_times[::step]
                if sampled[-1] != sorted_times[-1]:
                    sampled.append(sorted_times[-1])
                sorted_times = sampled

            time_series = QLineSeries()
            time_series.setPointsVisible(len(sorted_times) <= 80)
            time_series.hovered.connect(self.on_time_hover)
            self.time_data = []

            for i, (time_str, count) in enumerate(sorted_times):
                time_series.append(i, count)
                self.time_data.append((time_str, count))

            time_chart = QChart()
            time_chart.addSeries(time_series)
            time_chart.setTitle(self.lang_manager.tr("time_trend_chart", "访问时间趋势图（分）"))
            time_chart.setAnimationOptions(QChart.AnimationOption.NoAnimation)

            time_axis_x = QCategoryAxis()
            step = max(1, len(sorted_times) // 10)

            for i, (time_str, _) in enumerate(sorted_times):
                if i % step == 0 or i == len(sorted_times) - 1:
                    try:
                        parts = str(time_str).split()
                        if len(parts) >= 2:
                            time_label = parts[1][:5]
                        else:
                            time_label = str(time_str)[:5] if len(str(time_str)) >= 5 else str(time_str)
                    except Exception:
                        time_label = f"{i:02d}:00"

                    time_axis_x.append(time_label, i)

            time_axis_x.setRange(0, max(0, len(sorted_times) - 1))
            time_chart.addAxis(time_axis_x, Qt.AlignmentFlag.AlignBottom)
            time_series.attachAxis(time_axis_x)

            time_axis_y = QValueAxis()
            time_axis_y.setTitleText(self.lang_manager.tr("request_count", "请求次数"))
            time_axis_y.setMin(0)
            max_count = max((count for _, count in sorted_times), default=0)
            time_axis_y.setMax(max_count * 1.1 if max_count else 1)

            time_chart.addAxis(time_axis_y, Qt.AlignmentFlag.AlignLeft)
            time_series.attachAxis(time_axis_y)
            time_chart.legend().hide()

            self.time_chart_view.setChart(time_chart)
    def on_pie_hover(self, slice, state, *_):
        """处理饼图悬停事件"""
        if state:
            # 显示更详细的信息
            QToolTip.showText(
                QCursor.pos(),
                f"{slice.label()}\n{self.lang_manager.tr('percentage', '占比')}: {slice.percentage() * 100:.1f}%",
                self.status_chart_view
            )

        else:
            QToolTip.hideText()

    def on_ip_bar_hover(self, status, index, barset, *_):
        """处理IP柱状图悬停事件"""
        if status:
            ip = self.ip_categories[index]
            value = barset.at(index)

            # 获取图表位置
            chart = self.ip_chart_view.chart()
            pos_in_chart = self.ip_chart_view.mapFromGlobal(QCursor.pos())
            scene_pos = self.ip_chart_view.mapToScene(pos_in_chart)
            chart_pos = chart.mapFromScene(scene_pos)

            # 计算工具提示位置
            tooltip_pos = self.ip_chart_view.mapToGlobal(
                self.ip_chart_view.mapFromScene(chart.mapToScene(chart_pos))
            )

            QToolTip.showText(
                tooltip_pos,
                f"IP: {ip}\n{self.lang_manager.tr('request_count', '请求次数')}: {int(value)}",
                self.ip_chart_view
            )
        else:
            QToolTip.hideText()

    def on_url_bar_hover(self, status, index, barset, *_):
        """处理URL柱状图悬停事件"""
        if status:
            url = self.url_categories[index]
            value = barset.at(index)

            # 获取完整URL（如果有被截断）
            full_url = self.url_categories[index] if len(self.url_categories[index]) <= 33 else \
                [k for k in self.url_stats['data'].keys() if
                 urlparse(k).path.startswith(self.url_categories[index].split("...")[0])][0]

            chart = self.uri_chart_view.chart()
            pos_in_chart = self.uri_chart_view.mapFromGlobal(QCursor.pos())
            scene_pos = self.uri_chart_view.mapToScene(pos_in_chart)
            chart_pos = chart.mapFromScene(scene_pos)

            tooltip_pos = self.uri_chart_view.mapToGlobal(
                self.uri_chart_view.mapFromScene(chart.mapToScene(chart_pos))
            )

            QToolTip.showText(
                tooltip_pos,
                f'URL: {full_url}\n{self.lang_manager.tr("visit_count", "访问次数")}: {int(value)}',
                self.uri_chart_view
            )
        else:
            QToolTip.hideText()

    def on_time_hover(self, point, state, *_):
        """处理时间趋势图悬停事件"""
        if state:
            index = int(point.x())
            if 0 <= index < len(self.time_data):
                time_str, count = self.time_data[index]
                QToolTip.showText(
                    QCursor.pos(),
                    f"时间: {time_str}\n访问量: {int(count)}",
                    self.time_chart_view
                )
        else:
            QToolTip.hideText()

    def ai_analysis_preparation(self, ai_analysis_storing_data=None, *_):
        """兼容旧调用：改为启动 DeepSeek Agent 安全研判。"""
        self.maybe_start_pending_agent()

    def extract_url_params(self, url, *_):
        """从URL中提取参数"""
        try:
            parsed = urlparse(url)
            params = {}

            if parsed.query:
                for param in parsed.query.split('&'):
                    if '=' in param:
                        key, value = param.split('=', 1)
                        params[key] = value

            return params
        except Exception:
            return {}

    def update_ai_result(self, text, *_):
        """兼容旧流式输出：追加到研判结论区"""
        if hasattr(self, "ai_result_text"):
            self.ai_result_text.moveCursor(QTextCursor.MoveOperation.End)
            self.ai_result_text.insertPlainText(text)
            self.ai_result_text.moveCursor(QTextCursor.MoveOperation.End)

    def update_ai_content(self, text, *_):
        """兼容旧内容区：写入 Agent 步骤日志"""
        if hasattr(self, "ai_agent_view"):
            self.ai_agent_view.appendPlainText(text)

    def ai_analysis_finished(self, ErrorResponse="", *_):
        """兼容旧完成回调"""
        self._set_ai_running(False)
        if ErrorResponse and hasattr(self, "ai_result_text"):
            self.ai_result_text.append(str(ErrorResponse))
        self.status_label.setText(ErrorResponse or "AI 完成")

    def stop_ai_analysis(self, *_):
        """停止 AI Agent"""
        worker = getattr(self, "_agent_worker", None) or getattr(self, "ai_analysis_thread", None)
        if worker is not None:
            if hasattr(worker, "cancel"):
                worker.cancel()
            if hasattr(worker, "isRunning") and worker.isRunning():
                # 不强制 terminate，等待循环响应 cancel
                self.ai_agent_view.appendPlainText("…正在取消")
        self._set_ai_running(False)
        self.status_label.setText("AI Agent 已停止")

    def select_save_path(self, *_):
        """选择保存路径"""
        path = QFileDialog.getExistingDirectory(
            self,
            self.lang_manager.tr("select_save_directory", "选择保存目录")
        )
        if path:
            self.save_path_input.setText(path)
            self.add_recent_activity(
                self.lang_manager.tr("set_save_path", "设置保存路径"),
                path,
                self.lang_manager.tr("done", "完成")
            )

    def extraction_finished(self, *_):
        """提取完成"""
        self.extract_btn.setEnabled(True)
        self.stop_extract_btn.setEnabled(False)

        def _load_results():
            if hasattr(self.extract_thread, 'last_result'):
                self.url_stats = self.extract_thread.last_result
                if self.loading_overlay:
                    self.loading_overlay.set_message(
                        self.lang_manager.tr("rendering_stats", "正在渲染统计界面...")
                    )
                self.add_recent_activity(
                    self.lang_manager.tr("http_file_extraction", "HTTP文件提取"),
                    self.Import_box.text(),
                    self.lang_manager.tr("done", "完成")
                )
                self.update_stats_display()
                self.update_dashboard_stats()
            else:
                self.add_recent_activity(
                    self.lang_manager.tr("http_file_extraction", "HTTP文件提取"),
                    self.Import_box.text(),
                    self.lang_manager.tr("done", "完成")
                )
            self.status_label.setText(self.lang_manager.tr("http_file_extraction_done", "HTTP文件提取完成"))

        self.run_with_loading(
            _load_results,
            message=self.lang_manager.tr("loading_results", "正在加载分析结果，请稍候...")
        )

    def stop_extraction(self, *_):
        """停止提取"""
        if hasattr(self, 'extract_thread') and self.extract_thread.isRunning():
            self.extract_thread.terminate()
            self.extract_thread.wait()

        self.extract_btn.setEnabled(True)
        self.stop_extract_btn.setEnabled(False)
        self.status_label.setText(self.lang_manager.tr("extraction_stopped", "提取已停止"))
        self.add_recent_activity(
            self.lang_manager.tr("stop_extraction", "停止提取"),
            self.Import_box.text(),
            self.lang_manager.tr("stopped", "已停止")
        )

    def clear_extract_results(self, *_):
        """清除提取结果"""
        self.file_table.setRowCount(0)
        self.status_label.setText(self.lang_manager.tr("results_cleared", "结果已清除"))

    def open_saved_file(self, index, *_):
        """打开已保存的文件"""
        row = index.row()
        save_path = self.file_table.item(row, 5).text()

        if not save_path or not os.path.exists(save_path):
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                self.lang_manager.tr("file_not_exist", "文件不存在或未保存!")
            )
            return

        try:
            if sys.platform == "win32":
                os.startfile(save_path)
            elif sys.platform == "darwin":
                subprocess.run(["open", save_path])
            else:
                subprocess.run(["xdg-open", save_path])
        except Exception as e:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("error", "错误"),
                self.lang_manager.tr("cannot_open_file", "无法打开文件: ") + str(e)
            )

    def save_selected_files(self, *_):
        """保存选中文件"""
        selected = self.file_table.selectedItems()
        if not selected:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                self.lang_manager.tr("please_select_files_to_save", "请先选择要保存的文件!")
            )
            return

        # 获取保存路径
        save_path = self.save_path_input.text()
        if not save_path:
            save_path = QFileDialog.getExistingDirectory(
                self,
                self.lang_manager.tr("select_save_directory", "选择保存目录")
            )
            if not save_path:
                return
            self.save_path_input.setText(save_path)

        # 获取选中的行
        rows = set(item.row() for item in selected)
        saved_count = 0

        for row in rows:
            # 模拟保存文件
            self.file_table.item(row, 4).setText(self.lang_manager.tr("saved", "已保存"))
            self.file_table.item(row, 5).setText(os.path.join(save_path, self.file_table.item(row, 0).text()))
            saved_count += 1

        QMessageBox.information(
            self,
            self.lang_manager.tr("done", "完成"),
            self.lang_manager.tr("files_saved_msg", "已保存 {count} 个文件到 {path}").format(
                count=saved_count, path=save_path)
        )
        self.add_recent_activity(
            self.lang_manager.tr("save_selected_files", "保存选中文件"),
            f"{saved_count}个文件",
            self.lang_manager.tr("done", "完成")
        )

    def save_all_files(self, *_):
        """保存所有文件"""
        if self.file_table.rowCount() == 0:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                self.lang_manager.tr("no_files_to_save", "没有可保存的文件!")
            )
            return

        # 获取保存路径
        save_path = self.save_path_input.text()
        if not save_path:
            save_path = QFileDialog.getExistingDirectory(
                self,
                self.lang_manager.tr("select_save_dir", "选择保存目录")
            )
            if not save_path:
                return
            self.save_path_input.setText(save_path)

        saved_count = 0

        for row in range(self.file_table.rowCount()):
            self.file_table.item(row, 4).setText(self.lang_manager.tr("file_saved", "已保存"))
            self.file_table.item(row, 5).setText(
                os.path.join(save_path, self.file_table.item(row, 0).text())
            )
            saved_count += 1

        QMessageBox.information(
            self,
            self.lang_manager.tr("success", self.lang_manager.tr("done", "完成")),
            self.lang_manager.tr("saved_to", "已保存 {count} 个文件到 {path}").format(
                count=saved_count, path=save_path
            )
        )
        self.add_recent_activity(
            self.lang_manager.tr("save_all_files", "保存所有文件"),
            f"{saved_count}个文件",
            self.lang_manager.tr("success", self.lang_manager.tr("done", "完成"))
        )

    def export_file_list(self, *_):
        """导出文件列表"""
        if self.file_table.rowCount() == 0:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "没有可导出的文件列表!")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出文件列表",
            "http_extracted_files.csv",
            "CSV Files (*.csv);;All Files (*)"
        )

        if not file_path:
            return

        try:
            with open(file_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)

                # 写入表头
                headers = []
                for col in range(self.file_table.columnCount()):
                    headers.append(self.file_table.horizontalHeaderItem(col).text())
                writer.writerow(headers)

                # 写入数据
                for row in range(self.file_table.rowCount()):
                    row_data = []
                    for col in range(self.file_table.columnCount()):
                        item = self.file_table.item(row, col)
                        row_data.append(item.text() if item else "")
                    writer.writerow(row_data)

            QMessageBox.information(self, self.lang_manager.tr("done", "完成"), "文件列表导出成功!")
            self.add_recent_activity("导出文件列表", file_path, self.lang_manager.tr("done", "完成"))
        except Exception as e:
            QMessageBox.warning(self, "错误", f"导出失败: {str(e)}")
            self.add_recent_activity("导出文件列表", file_path, f"失败: {str(e)}")

    def start_http_extraction(self, *_):
        """ 自动化调用的开始提取HTTP文件前的选择 """
        # 获取选中的文件类型
        selected_index = self.file_filter_combo.currentIndex()
        file_filter = ["java_class", "java_serialized"]

        # 如果不是"所有文件类型"，则获取具体的文件类型
        if selected_index > 0:
            selected_text = self.file_filter_combo.currentText()
            if selected_text.split()[0] == "所有文件类型":
                file_filter = None
            else:
                # 从"类型 文件"格式中提取类型
                file_filter = [selected_text.split()[0]]
        self.start_http_extraction_file(file_filter)

    def start_http_extraction_alone(self, *_):
        """ 开始提取HTTP文件前的选择 """
        file = self.Import_box.text()
        if not file:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("Please_select_the_traffic_file_first", "请先选择流量文件!"))
            return

        save_path = self.save_path_input.text()
        if not save_path:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("Please_select_a_directory_to_save", "请选择保存目录!"))
            return

        # 获取选中的文件类型
        selected_index = self.file_filter_combo_independent.currentIndex()
        file_filter = ["java_class", "java_serialized"]

        # 如果不是"所有文件类型"，则获取具体的文件类型
        if selected_index > 0:
            selected_text = self.file_filter_combo_independent.currentText()
            if selected_text.split()[0] == "所有文件类型":
                file_filter = None
            else:
                # 从"类型 文件"格式中提取类型
                file_filter = [selected_text.split()[0]]
        self.start_http_extraction_file(file_filter)

    def start_http_extraction_file(self, file_filter, *_):
        """ 开始提取HTTP文件 """
        file = self.Import_box.text()
        if not file:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("Please_select_the_traffic_file_first", "请先选择流量文件!"))
            return

        save_path = self.save_path_input.text()
        if not save_path:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("Please_select_a_directory_to_save", "请选择保存目录!"))
            return

        # 清空之前的结果
        self.file_table.setRowCount(0)

        # 更新UI状态
        self.extract_btn.setEnabled(False)
        self.stop_extract_btn.setEnabled(True)

        # 获取SSL密钥日志文件路径（如果有）
        ssl_keylog = self.ssl_keylog_input.text() if self.ssl_keylog_input.text() else None

        fileextraction = {
            "file_filter": file_filter,  # 是否是要是选择的文件类型
            "save_path": save_path  # 提取要保存的文件
        }

        # 创建并启动提取线程
        self.extract_thread = AnalysisThread(
            file=file,
            uri="",
            keyword="",
            output="",
            request_only=True,
            response_only=False,
            show_body=True,
            sslkeylogfile=ssl_keylog,
            fileextraction=fileextraction,
            lang_manager=self.lang_manager
        )
        self.extract_thread.analysis_similar = "pyshark" if self.extract_pyshark_radio.isChecked() else "tshark"
        self.extract_thread.result_signal_extract.connect(self.add_extracted_file)
        self.extract_thread.finished_signal.connect(self.extraction_finished)
        self.extract_thread.status_label.connect(self.update_status_label)
        self.extract_thread.start()

        self.status_label.setText(self.lang_manager.tr("extracting_http_files", "正在提取HTTP文件..."))
        self.add_recent_activity(
            self.lang_manager.tr("start_extract_http", "开始提取HTTP文件"),
            file,
            self.lang_manager.tr("In_progress", "进行中")
        )

    def add_extracted_file(self, file_info: dict, *_):
        """添加提取的文件到表格"""
        row = self.file_table.rowCount()
        self.file_table.insertRow(row)

        # 提取需要展示的字段
        display_fields = [
            file_info.get("filename", "N/A"),
            file_info.get("filetype", "N/A"),
            f"{file_info.get('size', 0)} B",
            file_info.get("url", "N/A"),
            file_info.get("status", "N/A"),
            file_info.get("save_path", "N/A"),
        ]
        self.add_recent_activity(
            self.lang_manager.tr("http_file_extraction", "HTTP文件提取"),
            self.Import_box.text(),
            self.lang_manager.tr(
                "http_file_extraction_detail",
                f"检测到请求二进制文件 {file_info.get('url', 'N/A')} 文件大小：{file_info.get('size', 0)} B {file_info.get('status', 'N/A')} 详细内容请到【文件提取】查看"
            )
        )

        for col, text in enumerate(display_fields):
            item = QTableWidgetItem(str(text))
            self.file_table.setItem(row, col, item)

    def select_ssl_keylog_file(self, *_):
        """选择SSL密钥日志文件"""
        file, _ = QFileDialog.getOpenFileName(
            self,
            "选择SSL密钥日志文件",
            "",
            "Keylog Files (*.log *.keylog);;All Files (*)"
        )
        if file:
            self.ssl_keylog_input.setText(file)
            self.add_recent_activity("选择SSL密钥文件", file, self.lang_manager.tr("done", "完成"))

    def create_extract_tab(self, *_):
        """创建HTTP文件提取选项卡（美化版，PyQt6适配）"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/extract.png"), "HTTP文件提取")
        self.tabs["extract"] = tab

        # 主布局 - 使用垂直布局
        main_layout = QVBoxLayout(tab)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(15)

        # ==================== 第一部分：分析选项 ====================
        analysis_group = QGroupBox(self.lang_manager.tr("analysis_options", "分析选项"))

        analysis_layout = QGridLayout(analysis_group)
        analysis_layout.setContentsMargins(10, 15, 10, 15)
        analysis_layout.setSpacing(15)

        # 分析引擎选择
        engine_label = QLabel(self.lang_manager.tr("analysis_method", "分析方式:"))
        engine_label.setStyleSheet("font-weight: normal;")
        self.extract_pyshark_radio = QRadioButton("pyshark (Python库)")

        self.tshark_radio = QRadioButton("tshark")
        self.tshark_radio.setChecked(True)

        engine_layout = QHBoxLayout()
        engine_layout.addWidget(self.extract_pyshark_radio)
        engine_layout.addWidget(self.tshark_radio)
        engine_layout.addStretch()

        # analysis_layout.addWidget(engine_label, 0, 0)
        # analysis_layout.addLayout(engine_layout, 0, 1, 1, 2)

        # SSL密钥日志文件
        ssl_label = QLabel(self.lang_manager.tr("optional_ssl_keylog", "可选参数 SSL密钥日志:"))
        self.ssl_keylog_input = QLineEdit()
        self.ssl_keylog_input.setPlaceholderText(
            self.lang_manager.tr("select_ssl_keylog_file", "选择SSL密钥日志文件..."))

        ssl_browse_btn = QPushButton(self.lang_manager.tr("browse", "浏览..."))
        ssl_browse_btn.clicked.connect(self.select_ssl_keylog_file)

        analysis_layout.addWidget(ssl_label, 1, 0)
        analysis_layout.addWidget(self.ssl_keylog_input, 1, 1)
        analysis_layout.addWidget(ssl_browse_btn, 1, 2)

        main_layout.addWidget(analysis_group)

        # ==================== 第二部分：提取选项 ====================
        extract_group = QGroupBox(self.lang_manager.tr("extract_options", "提取选项"))
        extract_layout = QGridLayout(extract_group)
        extract_layout.setContentsMargins(10, 15, 10, 15)
        extract_layout.setSpacing(15)

        # 读取config.yaml文件获取文件类型
        try:
            signatures = self.config.get('signatures', [])
            file_types = sorted(
                {sig.get('type', '') for sig in signatures if sig.get('enabled', False) and sig.get('type', '')})
        except Exception as e:
            print(f"Error loading config.yaml: {e}")
            file_types = []
        # 默认选项
        default_options = ["所有文件类型"]

        # 合并默认选项和文件类型，并确保默认选项在前面
        file_types = default_options + sorted(set(file_types) - set(default_options))

        # 文件类型过滤
        filter_label = QLabel(self.lang_manager.tr("file_type", "文件类型:"))

        self.file_filter_combo_independent = QComboBox()
        self.file_filter_combo_independent.addItem(
            self.lang_manager.tr("recognize_extract_serialized_data", "识别提取常见序列化二数据"))
        if file_types:
            self.file_filter_combo_independent.insertSeparator(1)
            for file_type in file_types:
                self.file_filter_combo_independent.addItem(f"{file_type} {self.lang_manager.tr('file', '文件')}")

        # 设置字体和背景色，避免显示问题

        extract_layout.addWidget(filter_label, 0, 0)
        extract_layout.addWidget(self.file_filter_combo_independent, 0, 1, 1, 2)

        # 保存路径
        save_label = QLabel(self.lang_manager.tr("save_path", "保存路径:"))

        default_save_path = get_output_dir()
        os.makedirs(default_save_path, exist_ok=True)

        self.save_path_input = QLineEdit(default_save_path)
        self.save_path_input.setPlaceholderText(self.lang_manager.tr("select_save_directory", "选择保存目录..."))
        self.save_path_input.setStyleSheet("padding: 5px;")

        save_path_btn = QPushButton(self.lang_manager.tr("browse", "浏览..."))
        save_path_btn.setStyleSheet("padding: 5px 10px;")
        save_path_btn.clicked.connect(self.select_save_path)

        extract_layout.addWidget(save_label, 1, 0)
        extract_layout.addWidget(self.save_path_input, 1, 1)
        extract_layout.addWidget(save_path_btn, 1, 2)

        main_layout.addWidget(extract_group)

        # ==================== 第三部分：控制按钮 ====================
        control_frame = QFrame()
        control_layout = QHBoxLayout(control_frame)
        control_layout.setContentsMargins(0, 0, 0, 0)
        control_layout.setSpacing(10)

        # 添加左侧弹簧使按钮居中
        control_layout.addStretch()

        # 开始提取按钮
        self.extract_btn = QPushButton(self.lang_manager.tr("start_extraction", "开始提取"))
        self.extract_btn.setIcon(QIcon("ico/extract.png"))

        self.extract_btn.clicked.connect(self.start_http_extraction_alone)

        # 停止按钮
        self.stop_extract_btn = QPushButton(self.lang_manager.tr("stop", "停止"))
        self.stop_extract_btn.setIcon(QIcon("ico/stop.png"))
        self.stop_extract_btn.setEnabled(False)

        self.stop_extract_btn.clicked.connect(self.stop_extraction)

        # 清除按钮
        self.clear_extract_btn = QPushButton(self.lang_manager.tr("clear_results", "清除结果"))
        self.clear_extract_btn.setIcon(QIcon("ico/clear.png"))

        self.clear_extract_btn.clicked.connect(self.clear_extract_results)

        # 添加按钮到布局
        control_layout.addWidget(self.extract_btn)
        control_layout.addWidget(self.stop_extract_btn)
        control_layout.addWidget(self.clear_extract_btn)

        # 添加右侧弹簧使按钮居中
        control_layout.addStretch()

        main_layout.addWidget(control_frame)

        # ==================== 第四部分：结果表格 ====================
        result_group = QGroupBox(self.lang_manager.tr("extraction_results", "提取结果"))

        result_layout = QVBoxLayout(result_group)
        result_layout.setContentsMargins(5, 15, 5, 5)
        result_layout.setSpacing(10)

        # 文件列表表格
        self.file_table = QTableWidget()
        self.file_table.setColumnCount(6)
        self.file_table.setHorizontalHeaderLabels([
            self.lang_manager.tr("filename", "文件名"),
            self.lang_manager.tr("filetype", "类型"),
            self.lang_manager.tr("filesize", "大小"),
            self.lang_manager.tr("url", "URL"),
            self.lang_manager.tr("status", "状态"),
            self.lang_manager.tr("save_path", "保存路径")
        ])
        self.file_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.file_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.file_table.horizontalHeader().setStretchLastSection(True)
        self.file_table.setSortingEnabled(True)
        self.file_table.doubleClicked.connect(self.open_saved_file)
        # 去掉行号显示
        self.file_table.verticalHeader().setVisible(False)

        # 设置列宽和列策略 (PyQt6使用新的枚举值)
        header = self.file_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)  # 文件名
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)  # 类型
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)  # 大小
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)  # URL
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)  # 状态
        header.setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)  # 保存路径

        result_layout.addWidget(self.file_table)
        main_layout.addWidget(result_group)

        # ==================== 第五部分：操作按钮 ====================
        action_frame = QFrame()
        action_layout = QHBoxLayout(action_frame)
        action_layout.setContentsMargins(0, 0, 0, 0)
        action_layout.setSpacing(10)

        # 添加左侧弹簧使按钮居中
        action_layout.addStretch()

        # 保存选中按钮
        self.save_selected_btn = QPushButton(self.lang_manager.tr("save_selected", "保存选中"))
        self.save_selected_btn.setIcon(QIcon("ico/save.png"))

        self.save_selected_btn.clicked.connect(self.save_selected_files)

        # 保存所有按钮
        self.save_all_btn = QPushButton(self.lang_manager.tr("save_all", "保存全部"))
        self.save_all_btn.setIcon(QIcon("ico/save_all.png"))

        self.save_all_btn.clicked.connect(self.save_all_files)

        # 导出列表按钮
        self.export_list_btn = QPushButton(self.lang_manager.tr("export_list", "导出列表"))
        self.export_list_btn.setIcon(QIcon("ico/export.png"))

        self.export_list_btn.clicked.connect(self.export_file_list)

        # 添加按钮到布局
        action_layout.addWidget(self.save_selected_btn)
        action_layout.addWidget(self.save_all_btn)
        action_layout.addWidget(self.export_list_btn)

        # 添加右侧弹簧使按钮居中
        action_layout.addStretch()

        main_layout.addWidget(action_frame)

    def export_stats(self, *_):
        """导出统计信息为CSV文件"""
        # 弹出文件保存对话框
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "保存文件",
            f"stats_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            "CSV 文件 (*.csv);;所有文件 (*)"
        )

        if not file_name:
            return  # 用户取消保存

        try:
            with open(file_name, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                # 写入CSV表头
                writer.writerow(['URL', '总访问次数', 'IP地址', 'IP位置', 'IP访问次数',
                                 '状态码', '请求方法', 'User Agent', '威胁类型'])

                for url, stats in self.url_stats['data'].items():
                    # 获取URL级别的信息
                    total_count = stats.get('count', 0)
                    danger_info = "; ".join(stats.get('danger', {}).keys()) if 'danger' in stats else "无"

                    # 如果没有源IP信息，写入一行基本信息
                    if 'source_ips' not in stats or not stats['source_ips']:
                        writer.writerow([url, total_count, "", "", "", "", "", "", danger_info])
                        continue

                    # 遍历每个源IP的详细信息
                    for ip, ip_data in stats['source_ips'].items():
                        # 解析IP和位置信息
                        ip_parts = ip.split("：")
                        ip_address = ip_parts[0] if len(ip_parts) > 0 else ip
                        ip_location = ip_parts[1] if len(ip_parts) > 1 else "未知"

                        # 获取IP级别的信息
                        ip_count = ip_data.get('count', 0)
                        status_codes = "; ".join(f"{k}:{v}" for k, v in ip_data.get('status_codes', {}).items())
                        methods = "; ".join(f"{k}:{v}" for k, v in ip_data.get('methods', {}).items())
                        uas = "; ".join(f"{k[:50]}...:{v}" if len(k) > 50 else f"{k}:{v}"
                                        for k, v in ip_data.get('UA', {}).items())

                        # 写入一行数据
                        writer.writerow([
                            url,
                            total_count,
                            ip_address,
                            ip_location,
                            ip_count,
                            status_codes,
                            methods,
                            uas,
                            danger_info
                        ])

            # 更新界面状态
            self.status_label.setText("统计信息已成功导出")
            self.add_recent_activity(self.lang_manager.tr("export_stats", "导出统计"), file_name,
                                     self.lang_manager.tr("done", "完成"))
            QMessageBox.information(self, self.lang_manager.tr("done", "完成"), "统计信息已成功导出！")
            # 设置 QMessageBox 样式
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")
            self.add_recent_activity(self.lang_manager.tr("export_stats", "导出统计"), file_name, f"失败: {str(e)}")

    def find_request(self, *_):
        """查找指定ID的请求"""

        file = self.Import_box.text()
        if not file:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("Please_select_the_traffic_file_first", "请先选择流量文件!"))
            return

        self.memory_optimization = False  # 当值变成了真会叫数据写入到硬盘里面
        self.status_create_replay_tab.setText("")
        base_name = os.path.basename(file)  # 取文件名，比如 gsl.cap
        folder_name = os.path.splitext(base_name)[0]  # 去掉扩展名，比如 gsl
        self.start_analysis_timestamp = folder_name + "_" + datetime.datetime.now().strftime("%Y%m%d_%H%M%S") + ".txt"

        # 清除之前的结果
        self.request_text_edit.clear()

        # 启动分析线程
        self.replay_thread = AnalysisThread(
            file=file,
            uri="",
            keyword="",
            output="",
            request_only=True,
            response_only=False,
            show_body=True,
            lang_manager=self.lang_manager
        )
        self.replay_thread.analysis_similar = "tshark"
        self.replay_thread.result_signal.connect(self.update_replay_request)
        self.replay_thread.finished_signal.connect(self.replay_finished)

        self.stop_replay_button.setEnabled(True)
        self.status_label.setText(self.lang_manager.tr("searching_requests", "正在查找请求..."))
        self.replay_thread.start()

        self.add_recent_activity(self.lang_manager.tr("search_request", "查找请求"), file,
                                 self.lang_manager.tr("In_progress", "进行中"))

    def replay_request(self, *_):
        """重放请求"""
        global proxies
        proxies = self.get_proxy_settings()

        # 清除之前的结果
        self.request_text_edit.clear()

        file = self.Import_box.text()
        if not file:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"),
                                self.lang_manager.tr("Please_select_the_traffic_file_first", "请先选择流量文件!"))
            return

        stream_id = self.stream_id_input.text().strip()
        if not stream_id:
            QMessageBox.warning(
                self,
                self.lang_manager.tr("warning", "警告"),
                self.lang_manager.tr("please_enter_request_id", "请输入要查找的请求ID!")
            )

            return

        # 启动分析线程
        self.replay_thread = AnalysisThread(
            file=file,
            uri="",
            keyword="",
            output="",
            request_only=True,
            response_only=False,
            show_body=True,
            request_stream_id=stream_id,
            lang_manager=self.lang_manager
        )

        self.replay_thread.analysis_similar = "tshark"
        self.replay_thread.result_signal.connect(self.update_replay_request)
        self.replay_thread.finished_signal.connect(self.replay_finished)

        self.find_request_button.setEnabled(False)
        self.stop_replay_button.setEnabled(True)
        self.status_label.setText(self.lang_manager.tr("replaying_request", "正在重放请求..."))
        self.replay_thread.start()

        self.add_recent_activity(self.lang_manager.tr("replay_request", "重放请求"), file,
                                 self.lang_manager.tr("In_progress", "进行中"))

    ################报告生成###########
    # 创建报告生成标签页
    def create_report_tab(self, *_):
        """创建专业报告生成标签页"""
        tab = QWidget()
        self.workspace.addTab(tab, QIcon("ico/report.png"), "报告生成")
        self.tabs["report"] = tab

        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(15)

        # 报告配置区域
        config_group = QGroupBox(self.lang_manager.tr("report_config", "报告配置"))

        config_layout = QFormLayout(config_group)
        config_layout.setContentsMargins(10, 15, 10, 10)
        config_layout.setSpacing(10)

        # 报告基本信息
        self.report_title = QLineEdit(self.lang_manager.tr("default_title", "网络安全分析报告"))
        self.report_author = QLineEdit()
        self.report_company = QLineEdit()
        self.report_date = QLineEdit(datetime.datetime.now().strftime("%Y-%m-%d"))

        # 报告选项
        options_group = QGroupBox(self.lang_manager.tr("report_options", "报告内容选项"))
        options_layout = QVBoxLayout(options_group)

        self.include_exec_summary = QCheckBox(self.lang_manager.tr("include_exec_summary", "包含执行摘要"))
        self.include_exec_summary.setChecked(True)
        self.include_stats = QCheckBox(self.lang_manager.tr("include_stats", "包含详细统计信息"))
        self.include_stats.setChecked(True)
        self.include_charts = QCheckBox(self.lang_manager.tr("include_charts", "包含图表分析"))
        self.include_charts.setChecked(True)
        self.include_details = QCheckBox(self.lang_manager.tr("include_details", "包含请求详情"))
        self.include_details.setChecked(False)
        self.include_threats = QCheckBox(self.lang_manager.tr("include_threats", "包含威胁分析"))
        self.include_threats.setChecked(True)
        self.include_recommendations = QCheckBox(self.lang_manager.tr("include_recommendations", "包含安全建议"))
        self.include_recommendations.setChecked(True)

        options_layout.addWidget(self.include_exec_summary)
        options_layout.addWidget(self.include_stats)
        options_layout.addWidget(self.include_charts)
        options_layout.addWidget(self.include_details)
        options_layout.addWidget(self.include_threats)
        options_layout.addWidget(self.include_recommendations)
        self.include_charts.setText("包含等级分布图")

        # 添加配置项
        config_layout.addRow(self.lang_manager.tr("title", "报告标题:"), self.report_title)
        config_layout.addRow(self.lang_manager.tr("author", "作者") + ":", self.report_author)
        config_layout.addRow(self.lang_manager.tr("company", "公司/组织") + ":", self.report_company)
        config_layout.addRow(self.lang_manager.tr("date", "日期") + ":", self.report_date)
        config_layout.addRow(options_group)

        # 报告预览区域
        preview_group = QGroupBox(self.lang_manager.tr("report_preview", "报告预览"))
        preview_layout = QVBoxLayout(preview_group)
        preview_layout.setContentsMargins(5, 15, 5, 5)

        self.report_preview = QTextEdit()
        self.report_preview.setReadOnly(True)
        self._last_report_html = ""
        self._last_report_data = None

        preview_layout.addWidget(self.report_preview)

        # 操作按钮区域
        button_group = QWidget()
        button_layout = QHBoxLayout(button_group)
        button_layout.setContentsMargins(0, 0, 0, 0)
        button_layout.setSpacing(10)

        self.generate_btn = QPushButton(self.lang_manager.tr("generate_preview", "生成报告预览"))
        self.generate_btn.setIcon(QIcon("ico/preview.png"))
        self.generate_btn.clicked.connect(self.generate_report_preview)

        self.export_pdf_btn = QPushButton(self.lang_manager.tr("export_pdf", "导出PDF"))
        self.export_pdf_btn.setIcon(QIcon("ico/pdf.png"))
        self.export_pdf_btn.clicked.connect(self.export_report_pdf)

        self.export_html_btn = QPushButton(self.lang_manager.tr("export_html", "导出HTML"))
        self.export_html_btn.setIcon(QIcon("ico/html.png"))
        self.export_html_btn.clicked.connect(self.export_report_html)

        self.export_docx_btn = QPushButton(self.lang_manager.tr("export_docx", "导出Word"))
        self.export_docx_btn.setIcon(QIcon("ico/word.png"))
        self.export_docx_btn.clicked.connect(self.export_report_docx)

        button_layout.addWidget(self.generate_btn)
        button_layout.addWidget(self.export_html_btn)
        button_layout.addWidget(self.export_pdf_btn)
        button_layout.addWidget(self.export_docx_btn)

        layout.addWidget(config_group)
        layout.addWidget(preview_group, stretch=1)
        layout.addWidget(button_group)

    def prepare_report_data(self, *_):
        """准备专业报告数据（单次遍历 + 正确威胁聚合）"""
        if not hasattr(self, 'url_stats') or not isinstance(self.url_stats, dict):
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "没有可用的分析数据!")
            return None
        data = self.url_stats.get('data') or {}
        if not data:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "没有可用的分析数据!")
            return None

        g = self.url_stats.get('_global_stats') or {}
        want_stats = self.include_stats.isChecked()
        want_details = self.include_details.isChecked()
        want_threats = self.include_threats.isChecked()
        want_charts = self.include_charts.isChecked()
        want_summary = self.include_exec_summary.isChecked()
        want_recs = self.include_recommendations.isChecked()

        # —— 单次扫描：方法 / IP / URL ——
        methods = defaultdict(int)
        ip_counts = defaultdict(int)
        url_rows = []
        for url, udata in data.items():
            if url == '_global_stats' or not isinstance(udata, dict):
                continue
            row = {
                'url': unquote(url),
                'count': int(udata.get('count') or 0),
                'danger_count': int(udata.get('danger_count') or 0),
                'ips': [],
            }
            for ip, ip_data in (udata.get('source_ips') or {}).items():
                if not isinstance(ip_data, dict):
                    continue
                cnt = int(ip_data.get('count') or 0)
                ip_counts[ip] += cnt
                for method, c in (ip_data.get('methods') or {}).items():
                    methods[method] += int(c or 0)
                if want_details:
                    parts = str(ip).split('：', 1)
                    row['ips'].append({
                        'ip': parts[0],
                        'location': parts[1] if len(parts) > 1 else '未知',
                        'count': cnt,
                    })
            if want_details and row['ips']:
                row['ips'].sort(key=lambda x: -x['count'])
            url_rows.append(row)

        top_ips = [
            {
                'ip': k.split('：')[0],
                'location': k.split('：')[1] if '：' in k else '未知',
                'count': v,
            }
            for k, v in sorted(ip_counts.items(), key=lambda x: -x[1])[:8]
        ]
        top_urls = [
            {'url': r['url'], 'count': r['count'], 'danger': r['danger_count'] > 0}
            for r in sorted(url_rows, key=lambda x: -x['count'])[:8]
        ]

        # —— 威胁：复用 collect_risks ——
        risks = collect_risks(data) if (want_threats or want_summary or want_charts or want_recs) else []
        by_sev = defaultdict(int)
        by_type = defaultdict(int)
        for r in risks:
            by_sev[r.get('level') or '未知'] += 1
            by_type[r.get('type') or '未知'] += 1

        sev_chart = {
            'critical': by_sev.get('严重', 0),
            'high': by_sev.get('高危', 0),
            'medium': by_sev.get('中危', 0) + by_sev.get('中', 0),
            'low': by_sev.get('低危', 0),
            'info': by_sev.get('信息', 0),
        }
        threat_total = len(risks)
        highish = sev_chart['critical'] + sev_chart['high']
        if sev_chart['critical'] > 0:
            risk_level = '高'
        elif highish > 0:
            risk_level = '中'
        elif threat_total > 0:
            risk_level = '低'
        else:
            risk_level = '低'

        total_requests = int(g.get('request_total') or sum(r['count'] for r in url_rows) or 0)
        unique_urls = int(g.get('total_unique_uris') or len(url_rows) or 0)
        unique_ips = int(g.get('total_unique_ips') or len(ip_counts) or 0)

        report_data = {
            'title': self.report_title.text() or '网络安全分析报告',
            'author': self.report_author.text(),
            'company': self.report_company.text(),
            'date': self.report_date.text() or datetime.datetime.now().strftime('%Y-%m-%d'),
            'filename': os.path.basename(self.Import_box.text()) if self.Import_box.text() else '未指定',
            'generated_at': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'version': version,
            'document_id': f"RPT-{datetime.datetime.now().strftime('%Y%m%d-%H%M%S')}",
            'include_details': want_details,
            'include_charts': want_charts,
        }

        if want_summary:
            findings = []
            if threat_total == 0:
                findings.append('规则引擎未命中安全风险；建议结合业务日志抽样复核。')
            else:
                if sev_chart['critical']:
                    findings.append(f"发现严重风险 {sev_chart['critical']} 条，建议立即处置。")
                if sev_chart['high']:
                    findings.append(f"发现高危风险 {sev_chart['high']} 条，建议优先排查。")
                if sev_chart['medium']:
                    findings.append(f"发现中危风险 {sev_chart['medium']} 条，建议安排核查。")
                top_types = sorted(by_type.items(), key=lambda x: -x[1])[:5]
                if top_types:
                    findings.append(
                        '主要规则类型：' + '、'.join(f'{n}({c})' for n, c in top_types)
                    )
            if top_ips:
                findings.append(
                    f"访问量最高来源 IP：{top_ips[0]['ip']}（{top_ips[0]['location']}，{top_ips[0]['count']} 次）"
                )
            if top_urls:
                findings.append(f"访问量最高 URL：{top_urls[0]['url']}（{top_urls[0]['count']} 次）")

            report_data['exec_summary'] = {
                'overview': (
                    f"本报告基于文件「{report_data['filename']}」的分析结果，"
                    f"共统计 {total_requests} 次请求、{unique_urls} 个唯一 URL、{unique_ips} 个来源 IP，"
                    f"规则命中威胁 {threat_total} 条。"
                ),
                'key_findings': findings,
                'risk_level': risk_level,
            }

        if want_stats or want_charts:
            report_data['stats'] = {
                'total_requests': total_requests,
                'unique_urls': unique_urls,
                'unique_ips': unique_ips,
                'status_codes': g.get('total_unique_status_code', []),
                'threats': int(g.get('danger_total') or threat_total),
                'request_methods': dict(sorted(methods.items(), key=lambda x: -x[1])),
                'top_ips': top_ips,
                'top_urls': top_urls,
                'severity_chart': sev_chart,
            }

        if want_threats:
            level_rank = {'严重': 0, '高危': 1, '中危': 2, '低危': 3, '信息': 4}
            class_map = {
                '严重': 'high', '高危': 'high', '中危': 'medium',
                '低危': 'low', '信息': 'low', '中': 'medium',
            }
            details = []
            for r in sorted(risks, key=lambda x: (level_rank.get(x.get('level'), 9), x.get('url', '')))[:200]:
                details.append({
                    'level': r.get('level', ''),
                    'type': r.get('type', ''),
                    'rule': r.get('rule', ''),
                    'ip': r.get('ip_short') or r.get('ip', ''),
                    'url': r.get('url_display') or unquote(r.get('url', '')),
                    'severity': self._get_severity_level(r.get('level', '')),
                    'severity_class': class_map.get(r.get('level'), 'low'),
                    'content': (r.get('content') or '')[:200],
                })
            type_stats = [
                {'name': n, 'count': c}
                for n, c in sorted(by_type.items(), key=lambda x: -x[1])
            ]
            report_data['threats'] = {
                'total': threat_total,
                'type_count': len(by_type),
                'types': [t['name'] for t in type_stats],
                'type_stats': type_stats,
                'by_severity': {
                    'critical': sev_chart['critical'],
                    'high': sev_chart['high'],
                    'medium': sev_chart['medium'],
                    'low': sev_chart['low'],
                    'info': sev_chart['info'],
                },
                'details': details,
            }

        if want_details:
            sorted_urls = sorted(url_rows, key=lambda x: -x['count'])
            report_data['url_stats'] = {
                'total': len(sorted_urls),
                'top_requests': sorted_urls[:50],
                'suspicious': [u for u in sorted_urls if u['danger_count'] > 0][:80],
            }

        if want_recs:
            threats_block = report_data.get('threats') or {
                'total': threat_total,
                'type_stats': [{'name': n, 'count': c} for n, c in by_type.items()],
                'by_severity': {
                    'critical': sev_chart['critical'],
                    'high': sev_chart['high'],
                    'medium': sev_chart['medium'],
                    'low': sev_chart['low'],
                    'info': sev_chart['info'],
                },
                'details': [],
            }
            report_data['recommendations'] = self._generate_recommendations(threats_block)

        return report_data

    def _get_severity_level(self, severity_str):
        """将严重程度文本转换为数字级别"""
        text = str(severity_str or '')
        if '严重' in text or 'critical' in text.lower():
            return 4
        if '高危' in text or 'high' in text.lower():
            return 3
        if '中危' in text or text == '中' or 'medium' in text.lower():
            return 2
        if '低危' in text or 'low' in text.lower():
            return 1
        if '信息' in text or 'info' in text.lower():
            return 0
        return 0

    def _calculate_request_methods(self):
        """计算请求方法分布"""
        methods = defaultdict(int)
        data = (self.url_stats or {}).get('data') or {}
        for url, udata in data.items():
            if url == '_global_stats' or not isinstance(udata, dict):
                continue
            for ip_data in (udata.get('source_ips') or {}).values():
                if not isinstance(ip_data, dict):
                    continue
                for method, count in (ip_data.get('methods') or {}).items():
                    methods[method] += int(count or 0)
        return dict(methods)

    def _get_top_ips(self, top_n=5):
        """获取Top N IP地址"""
        ip_counts = defaultdict(int)
        data = (self.url_stats or {}).get('data') or {}
        for url, udata in data.items():
            if url == '_global_stats' or not isinstance(udata, dict):
                continue
            for ip, ip_data in (udata.get('source_ips') or {}).items():
                if isinstance(ip_data, dict):
                    ip_counts[ip] += int(ip_data.get('count') or 0)
        sorted_ips = sorted(ip_counts.items(), key=lambda x: x[1], reverse=True)
        return [
            {
                'ip': ip.split('：')[0],
                'location': ip.split('：')[1] if '：' in ip else '未知',
                'count': count,
            }
            for ip, count in sorted_ips[:top_n]
        ]

    def _get_top_urls(self, top_n=5):
        """获取Top N URL"""
        urls = []
        data = (self.url_stats or {}).get('data') or {}
        for url, udata in data.items():
            if url == '_global_stats' or not isinstance(udata, dict):
                continue
            urls.append({
                'url': unquote(url),
                'count': int(udata.get('count') or 0),
                'danger': bool(udata.get('danger_count') or 0),
            })
        return sorted(urls, key=lambda x: x['count'], reverse=True)[:top_n]

    def _generate_recommendations(self, threats_data):
        """根据真实威胁类型生成安全建议"""
        recommendations = []
        type_names = ' '.join(
            t.get('name', '') for t in (threats_data.get('type_stats') or [])
        ).lower()
        type_names += ' ' + ' '.join(str(x) for x in (threats_data.get('types') or [])).lower()
        detail_blob = ' '.join(
            f"{d.get('type','')} {d.get('rule','')} {d.get('url','')}"
            for d in (threats_data.get('details') or [])
        ).lower()
        blob = type_names + ' ' + detail_blob

        by_sev = threats_data.get('by_severity') or {}
        if int(by_sev.get('critical') or 0) + int(by_sev.get('high') or 0) > 0:
            recommendations.append({
                'title': '优先处置高危/严重项',
                'description': '当前存在严重或高危命中，建议按来源 IP 与 URL 逐条复核：',
                'details': [
                    '对严重/高危条目核对匹配上下文，确认是否为真实攻击或误报',
                    '必要时临时限制异常来源 IP 或加固相关接口鉴权',
                    '将确认的攻击样本补充到 WAF / 检测规则',
                ],
                'priority': '高',
            })

        if any(k in blob for k in ('sql', 'injection', '注入', 'database')):
            recommendations.append({
                'title': '注入类风险防护',
                'description': '检测到疑似注入/命令相关规则命中，建议：',
                'details': [
                    '接口统一使用参数化查询 / 预编译语句',
                    '对用户输入做白名单校验，避免拼接执行',
                    '配置 WAF 规则拦截常见注入特征',
                ],
                'priority': '高',
            })

        if any(k in blob for k in ('traversal', '目录', 'path', 'lfi', '../')):
            recommendations.append({
                'title': '路径遍历与文件访问控制',
                'description': '检测到目录穿越/敏感路径相关命中，建议：',
                'details': [
                    '禁止对外暴露备份、配置、源码目录',
                    '规范化路径解析，拒绝 .. 与绝对路径穿越',
                    '静态资源目录与业务上传目录分离并限制权限',
                ],
                'priority': '高',
            })

        if any(k in blob for k in ('xss', 'script', '跨站')):
            recommendations.append({
                'title': 'XSS 与输出编码',
                'description': '检测到跨站脚本相关命中，建议：',
                'details': [
                    '输出到 HTML/JS 的内容按上下文编码',
                    '设置合适的 Content-Security-Policy',
                    '对富文本输入做严格过滤',
                ],
                'priority': '中',
            })

        if any(k in blob for k in ('shell', 'command', 'cmd', '命令')):
            recommendations.append({
                'title': '命令执行面收敛',
                'description': '检测到系统命令/Shell 相关命中，建议：',
                'details': [
                    '避免 Web 进程直接调用系统命令',
                    '若必须调用，使用固定参数白名单',
                    '加强应用账号权限与审计日志',
                ],
                'priority': '高',
            })

        recommendations.append({
            'title': '常规安全加固',
            'description': '基于本次分析的一般性建议：',
            'details': [
                '保持系统与应用组件更新',
                '实施最小权限与访问审计',
                '定期复核访问日志与异常来源 IP',
                '重要接口启用鉴权、限流与告警',
            ],
            'priority': '中',
        })
        return recommendations

    def _report_templates_dir(self):
        base = os.path.dirname(os.path.abspath(__file__))
        return os.path.join(base, 'templates')

    def _ensure_report_html(self):
        """确保已有缓存的 Jinja HTML；没有则生成。"""
        if getattr(self, '_last_report_html', None):
            return self._last_report_html
        self.generate_report_preview()
        return getattr(self, '_last_report_html', '') or ''

    def generate_report_preview(self, *_):
        """生成专业报告预览"""
        report_data = self.prepare_report_data()
        if not report_data:
            return

        try:
            env = Environment(
                loader=FileSystemLoader(self._report_templates_dir()),
                autoescape=True,
                trim_blocks=True,
                lstrip_blocks=True,
            )

            def _format_date(x):
                try:
                    return datetime.datetime.strptime(str(x), '%Y-%m-%d').strftime('%Y年%m月%d日')
                except Exception:
                    return str(x)

            def _severity_class(x):
                mapping = {4: 'high', 3: 'high', 2: 'medium', 1: 'low', 0: 'low'}
                try:
                    return mapping.get(int(x), 'low')
                except Exception:
                    return 'low'

            env.filters['format_date'] = _format_date
            env.filters['severity_class'] = _severity_class

            template = env.get_template('professional_report.html')
            html_content = template.render(**report_data)

            self._last_report_html = html_content
            self._last_report_data = report_data
            self.report_preview.setHtml(html_content)
            self.status_label.setText("专业报告预览已生成")
            self.add_recent_activity("生成报告预览", "", self.lang_manager.tr("done", "完成"))

        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成报告预览失败: {str(e)}")
            self.add_recent_activity("生成报告预览", "", f"失败: {str(e)}")

    def export_report_pdf(self, *_):
        """导出PDF报告"""
        html_content = self._ensure_report_html()
        if not html_content:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "请先生成报告预览!")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出PDF报告",
            f"{self.report_title.text() or 'report'}.pdf",
            "PDF文件 (*.pdf);;所有文件 (*)"
        )
        if not file_path:
            return

        try:
            options = {
                'encoding': 'UTF-8',
                'page-size': 'A4',
                'margin-top': '12mm',
                'margin-right': '12mm',
                'margin-bottom': '12mm',
                'margin-left': '12mm',
                'quiet': '',
                'enable-local-file-access': None,
                'footer-center': '[page]/[topage]',
                'footer-font-size': '8',
            }
            pdfkit.from_string(html_content, file_path, options=options)
            QMessageBox.information(self, self.lang_manager.tr("done", "完成"), "PDF报告导出成功!")
            self.status_label.setText(f"PDF报告已导出: {file_path}")
            self.add_recent_activity("导出PDF报告", file_path, self.lang_manager.tr("done", "完成"))
        except Exception as e:
            QMessageBox.critical(
                self, "错误",
                f"导出PDF失败: {str(e)}\n\n需安装 wkhtmltopdf 并加入 PATH；也可先用「导出HTML」。"
            )
            self.add_recent_activity("导出PDF报告", file_path, f"失败: {str(e)}")

    def export_report_html(self, *_):
        """导出HTML报告（使用原始 Jinja 渲染结果）"""
        html_content = self._ensure_report_html()
        if not html_content:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "请先生成报告预览!")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出HTML报告",
            f"{self.report_title.text() or 'report'}.html",
            "HTML文件 (*.html);;所有文件 (*)"
        )
        if not file_path:
            return

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            QMessageBox.information(self, self.lang_manager.tr("done", "完成"), "HTML报告导出成功!")
            self.status_label.setText(f"HTML报告已导出: {file_path}")
            self.add_recent_activity("导出HTML报告", file_path, self.lang_manager.tr("done", "完成"))
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出HTML失败: {str(e)}")
            self.add_recent_activity("导出HTML报告", file_path, f"失败: {str(e)}")

    def export_report_docx(self, *_):
        """导出Word报告（写入关键结构化内容）"""
        html_content = self._ensure_report_html()
        data = getattr(self, '_last_report_data', None)
        if not html_content or not data:
            QMessageBox.warning(self, self.lang_manager.tr("warning", "警告"), "请先生成报告预览!")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出Word报告",
            f"{self.report_title.text() or 'report'}.docx",
            "Word文件 (*.docx);;所有文件 (*)"
        )
        if not file_path:
            return

        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            doc.add_heading(data.get('title') or '网络安全分析报告', level=0)
            doc.add_paragraph(f"作者: {data.get('author') or '未指定'}")
            doc.add_paragraph(f"公司/组织: {data.get('company') or '未指定'}")
            doc.add_paragraph(f"日期: {data.get('date') or ''}")
            doc.add_paragraph(f"文档ID: {data.get('document_id') or ''}")
            doc.add_paragraph(f"分析文件: {data.get('filename') or ''}")

            if data.get('exec_summary'):
                doc.add_heading('1. 执行摘要', level=1)
                es = data['exec_summary']
                doc.add_paragraph(es.get('overview') or '')
                doc.add_paragraph(f"总体风险等级: {es.get('risk_level') or ''}")
                for fnd in es.get('key_findings') or []:
                    doc.add_paragraph(fnd, style='List Bullet')

            if data.get('stats'):
                st = data['stats']
                doc.add_heading('2. 统计分析', level=1)
                doc.add_paragraph(
                    f"总请求 {st.get('total_requests', 0)} · 唯一URL {st.get('unique_urls', 0)} · "
                    f"唯一IP {st.get('unique_ips', 0)} · 威胁 {st.get('threats', 0)}"
                )

            if data.get('threats'):
                th = data['threats']
                doc.add_heading('3. 威胁分析', level=1)
                doc.add_paragraph(f"共 {th.get('total', 0)} 条风险，{th.get('type_count', 0)} 类规则")
                for t in (th.get('details') or [])[:80]:
                    doc.add_paragraph(
                        f"[{t.get('level')}] {t.get('type')} | {t.get('rule')} | "
                        f"{t.get('ip')} | {t.get('url')}",
                        style='List Bullet',
                    )

            if data.get('recommendations'):
                doc.add_heading('4. 安全建议', level=1)
                for rec in data['recommendations']:
                    doc.add_heading(f"{rec.get('title')}（优先级: {rec.get('priority')}）", level=2)
                    doc.add_paragraph(rec.get('description') or '')
                    for d in rec.get('details') or []:
                        doc.add_paragraph(d, style='List Bullet')

            doc.add_paragraph(f"生成时间: {data.get('generated_at') or ''} | 版本: {data.get('version') or ''}")
            doc.save(file_path)

            QMessageBox.information(self, self.lang_manager.tr("done", "完成"), "Word报告导出成功!")
            self.status_label.setText(f"Word报告已导出: {file_path}")
            self.add_recent_activity("导出Word报告", file_path, self.lang_manager.tr("done", "完成"))
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出Word失败: {str(e)}")
            self.add_recent_activity("导出Word报告", file_path, f"失败: {str(e)}")

    ################报告生成###########

    def memory_optimization_invoke(self, text_edit, status_label, text):
        """安全的内存优化文本追加方法"""
        text_edit.append(text)

    def update_replay_request(self, text, *_):
        """更新请求显示"""
        self.memory_optimization_invoke(self.request_text_edit, self.status_create_replay_tab, text=text)

    def clear_replay_results(self, *_):
        """清除重放结果"""
        self.request_text_edit.clear()
        if hasattr(self, 'current_request'):
            del self.current_request
        self.status_label.setText(self.lang_manager.tr("results_cleared", "结果已清除"))

    def stop_replay(self, *_):
        """停止查找请求"""
        if self.replay_thread and self.replay_thread.isRunning():
            self.replay_thread.terminate()
            self.replay_thread.wait()
            self.request_text_edit.append("\n操作已停止!")
            self.status_label.setText("操作已停止")
            self.add_recent_activity("停止操作", self.Import_box.text(), "已停止")

        self.stop_replay_button.setEnabled(False)

    def replay_finished(self, *_):
        """查找请求完成"""
        self.stop_replay_button.setEnabled(False)
        self.status_label.setText("操作完成")
        self.add_recent_activity("请求操作", self.Import_box.text(), self.lang_manager.tr("done", "完成"))

    def cleanup_resources(self, *_):
        """清理资源"""
        if hasattr(self, '_status_chart'):
            self._status_chart.deleteLater()

        if hasattr(self, 'url_table_model'):
            self.url_table_model.deleteLater()

        # 清理图表视图
        for chart_view in [self.status_chart_view, self.ip_chart_view,
                           self.time_chart_view, self.uri_chart_view]:
            if chart_view.chart():
                chart_view.chart().deleteLater()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        overlay = getattr(self, "loading_overlay", None)
        if overlay is not None and overlay.isVisible():
            try:
                overlay._cover_parent()
            except RuntimeError:
                pass

    def _stop_qthread(self, th, timeout_ms=2000):
        """安全停止 QThread，避免 Destroyed while still running。"""
        if th is None:
            return
        try:
            if hasattr(th, "stop") and callable(th.stop):
                try:
                    th.stop()
                except Exception:
                    pass
            if hasattr(th, "cancel") and callable(th.cancel):
                try:
                    th.cancel()
                except Exception:
                    pass
            if hasattr(th, "isRunning") and th.isRunning():
                if hasattr(th, "quit"):
                    th.quit()
                if not th.wait(timeout_ms):
                    th.terminate()
                    th.wait(500)
        except Exception:
            pass

    def closeEvent(self, event):
        """窗口关闭事件，弹出确认提示"""

        reply = QMessageBox.question(
            self,
            self.lang_manager.tr("confirm_exit", "确认退出"),
            self.lang_manager.tr("exit_confirmation", "你确定要退出程序吗？"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No  # 默认选择“No”
        )

        if reply == QMessageBox.StandardButton.Yes:
            # 用户确认退出，执行清理逻辑
            try:
                self._disconnect_remote()
            except Exception:
                pass

            self.stop_go_backend()

            try:
                self.file_writer.stop()
            except Exception:
                pass
            self.cleanup_resources()

            # 停止并等待各后台线程
            self._stop_qthread(getattr(self, "analysis_thread", None))
            self._stop_qthread(getattr(self, "replay_thread", None))
            self._stop_qthread(getattr(self, "_agent_worker", None) or getattr(self, "ai_analysis_thread", None))
            self._stop_qthread(getattr(self, "worker_thread", None))
            self._stop_qthread(getattr(self, "extract_thread", None))
            self._stop_qthread(getattr(self, "remote_thread", None))
            self._stop_qthread(getattr(self, "stream_analyzer", None))

            # 停止所有子进程
            for process in multiprocessing.active_children():
                try:
                    process.terminate()
                    process.join(2)
                except Exception:
                    pass

            event.accept()

        else:
            # 用户取消关闭
            event.ignore()


if __name__ == "__main__":
    # 确保multiprocessing在Windows下能正确工作
    multiprocessing.freeze_support()
    app = QApplication(sys.argv)
    app.setStyle('Fusion')

    font = QFont()
    font.setFamily("Segoe UI")
    font.setPointSize(10)
    app.setFont(font)

    window = MainWindow()
    window.show()
    sys.exit(app.exec())
